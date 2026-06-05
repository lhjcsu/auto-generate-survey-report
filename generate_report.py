#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
岩土工程勘察报告 — 自动生成工具 v4.0
============================================================
从华宁勘察软件导出的 Excel 数据 + 正式报告 .docx 模板，
一键生成完整的岩土工程勘察报告。

用法:
    # 方式一: 使用 JSON 配置文件 (推荐)
    python generate_report.py --config project_config.json

    # 方式二: 指定项目目录 (向后兼容 v3)
    python generate_report.py --project <项目目录>

    # 可选参数
    --template  指定 .docx 模板路径
    --output    指定输出文件路径
    --layers    指定地层名称映射 JSON
    --dry-run   仅加载数据，不生成报告
    --verbose   输出详细日志

依赖:
    pip install python-docx xlrd openpyxl
"""
from __future__ import annotations

import argparse
import copy
import glob
import json
import logging
import math
import os
import re
import sys
from collections import defaultdict
from typing import Any, Dict, List, Optional, Tuple

# 地震参数查表 (GB 18306-2015 附录 C.15)
try:
    from seismic_lookup import get_seismic_params as _lookup_seismic
except ImportError:
    _lookup_seismic = None  # type: ignore

# 第三方依赖: 延迟导入以支持 --help 在未安装依赖时仍可运行
try:
    import openpyxl
    import xlrd
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    # 仅在真正需要使用时才报错
    openpyxl = None  # type: ignore[assignment]
    xlrd = None  # type: ignore[assignment]
    Document = None  # type: ignore[assignment,misc]

from corrosion_eval import evaluate_corrosion

# 波速 (方案一: 估算 / 方案二: 实测数据解析)
try:
    from wave_velocity_estimate import (
        estimate_vs,
        compute_equivalent_vs,
        compute_cover_thickness,
        classify_site,
        evaluate_site_class_from_layers,
    )
except ImportError:
    estimate_vs = compute_equivalent_vs = compute_cover_thickness = None  # type: ignore
    classify_site = evaluate_site_class_from_layers = None  # type: ignore

try:
    from wave_velocity_parser import load_wave_velocity_data
except ImportError:
    load_wave_velocity_data = None  # type: ignore



def _check_dependencies() -> None:
    """检查第三方依赖是否已安装"""
    missing = []
    if openpyxl is None:
        missing.append("openpyxl")
    if xlrd is None:
        missing.append("xlrd")
    if Document is None:
        missing.append("python-docx")
    if missing:
        print(f"错误: 缺少依赖包: {', '.join(missing)}")
        print(f"请运行: pip install {' '.join(missing)}")
        sys.exit(1)

__version__ = "4.0"

# ============================================================
# 日志配置
# ============================================================

logger = logging.getLogger("survey_report")


def setup_logging(verbose: bool = False) -> None:
    """配置日志输出格式和级别"""
    level = logging.DEBUG if verbose else logging.INFO
    handler = logging.StreamHandler(sys.stdout)
    handler.setFormatter(logging.Formatter("%(message)s"))
    logger.setLevel(level)
    logger.addHandler(handler)


# ============================================================
# 默认表格索引 (可通过配置文件覆盖)
# ============================================================

DEFAULT_TABLE_INDICES: Dict[str, int] = {
    "buildings": 2,          # T2  建筑物特征表
    "workload": 3,           # T3  工作量统计表
    "phys_spt_start": 4,     # T4~T15 物理力学 & 原位测试
    "phys_spt_end": 15,
    "water_level": 16,       # T16 水位统计表
    "bearing_capacity": 18,  # T18 承载力建议值
    "water_sample": 19,      # T19 水样分析表
    "water_corrosion": 20,   # T20 水腐蚀性评价表
    "salt_sample": 21,       # T21 易溶盐分析表
    "salt_corrosion": 22,    # T22 土腐蚀性评价表
    "foundation_type": 23,   # T23 基础类型建议表
    "pile_params": 24,       # T24 桩基参数建议表
}

# 物理力学指标: 行号 → 指标键名
PHYS_ROW_MAP: Dict[int, str] = {
    1: "W", 2: "gamma", 3: "e0",
    4: "WL", 5: "WP", 6: "IP", 7: "IL",
}

# 物理力学统计列: 统计类型 → 列号
STAT_COL_MAP: Dict[str, int] = {
    "min": 1, "max": 2, "avg": 3,
    "n": 4, "std": 5, "cv": 6, "std_val": 7,
}

# 统计关键词 → 内部键名
STAT_KEYWORD_MAP: Dict[str, str] = {
    "最小值": "min", "最大值": "max", "数据个数": "n",
    "平均值": "avg", "标准差": "std", "变异系数": "cv", "标准值": "std_val",
}


# ============================================================
# 工具函数
# ============================================================

def safe_float(v: Any, default: Optional[float] = None) -> Optional[float]:
    """安全转换为 float，失败返回 default"""
    if v is None:
        return default
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if s in ("", "-", "—", "/"):
        return default
    try:
        return float(s)
    except (ValueError, TypeError):
        return default


def safe_str(v: Any) -> str:
    """安全转换为去首尾空格的字符串"""
    if v is None:
        return ""
    return str(v).strip()


def fmt_val(v: Any, fmt: str = ".2f") -> str:
    """格式化浮点数，None 返回空串"""
    if v is None:
        return ""
    try:
        return f"{float(v):{fmt}}"
    except (ValueError, TypeError):
        return ""


def fmt_val_int(v: Any) -> str:
    """格式化整数或保留一位小数"""
    if v is None:
        return ""
    try:
        f = float(v)
        return str(int(f)) if f == int(f) else f"{f:.1f}"
    except (ValueError, TypeError):
        return ""


def find_file(directory: str, pattern: str) -> Optional[str]:
    """在目录及其子目录中查找匹配文件 (忽略大小写)"""
    if not os.path.isdir(directory):
        return None
    pattern_lower = pattern.lower()
    # 先搜当前目录
    for f in os.listdir(directory):
        if pattern_lower in f.lower():
            return os.path.join(directory, f)
    # 递归搜索
    for root, _dirs, files in os.walk(directory):
        for f in files:
            if pattern_lower in f.lower():
                return os.path.join(root, f)
    return None


def set_cell(table: Any, row: int, col: int, val: Any) -> None:
    """设置 docx 表格单元格文字，保留原有格式"""
    if row >= len(table.rows):
        return
    if col >= len(table.rows[row].cells):
        return
    cell = table.rows[row].cells[col]
    text = str(val) if val is not None else ""
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def set_para_text(para: Any, text: str) -> None:
    """设置段落文字，保留首个 run 的格式"""
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def replace_in_para(para: Any, old: str, new: str) -> None:
    """在段落的所有 run 中执行文本替换"""
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)


def layer_sort_key(layer_id: str) -> Tuple[int, int]:
    """地层排序键: 按数字排序，支持 '10-1' 格式"""
    parts = str(layer_id).replace("层", "").replace("第", "").split("-")
    major = int(parts[0]) if parts[0].lstrip("-").isdigit() else 999
    minor = int(parts[1]) if len(parts) > 1 and parts[1].isdigit() else 0
    return (major, minor)


# ============================================================
# 项目配置管理
# ============================================================

class ProjectConfig:
    """项目配置 — 支持 JSON 配置文件和目录自动检测两种方式"""

    def __init__(
        self,
        config_path: Optional[str] = None,
        project_dir: Optional[str] = None,
        template_override: Optional[str] = None,
        output_override: Optional[str] = None,
        layers_override: Optional[str] = None,
    ):
        self.raw: Dict[str, Any] = {}
        self.project_name: str = ""
        self.base_dir: str = ""
        self.excel_dir: str = ""
        self.template_path: str = ""
        self.output_path: str = ""
        self.layers_file: Optional[str] = layers_override

        # 表格索引 (可被配置文件覆盖)
        self.table_indices: Dict[str, int] = dict(DEFAULT_TABLE_INDICES)

        # 地层配置
        self.layers: List[Dict[str, Any]] = []

        # 全局替换规则
        self.replacements: List[Tuple[str, str]] = []

        # 含水层/液化配置
        self.water_aquifer_layers: List[str] = []
        self.liquefaction_layers: Dict[str, Dict] = {}

        # 段落模板
        self.workload_template: str = ""
        self.water_paragraph_template: str = ""

        if config_path:
            self._load_from_config(config_path)
        elif project_dir:
            self._load_from_directory(project_dir)
        else:
            raise ValueError("必须提供 config_path 或 project_dir")

        # CLI 覆盖
        if template_override:
            self.template_path = template_override
        if output_override:
            self.output_path = output_override

    def _load_from_config(self, config_path: str) -> None:
        """从 JSON 配置文件加载"""
        if not os.path.isfile(config_path):
            raise FileNotFoundError(f"配置文件不存在: {config_path}")

        with open(config_path, "r", encoding="utf-8") as f:
            self.raw = json.load(f)

        config_dir = os.path.dirname(os.path.abspath(config_path))
        self.project_name = self.raw.get("project_name", "未知项目")

        # base_dir: 优先使用配置中的绝对路径，否则相对于配置文件
        base = self.raw.get("base_dir", "")
        if base and os.path.isabs(base):
            self.base_dir = base
        elif base:
            self.base_dir = os.path.join(config_dir, base)
        else:
            self.base_dir = config_dir

        # excel_dir
        excel_sub = self.raw.get("excel_dir", "")
        if excel_sub:
            self.excel_dir = os.path.join(self.base_dir, excel_sub)
        else:
            self.excel_dir = self._find_excel_dir()

        # 模板
        tpl = self.raw.get("template", "")
        if tpl and os.path.isabs(tpl):
            self.template_path = tpl
        elif tpl:
            self.template_path = os.path.join(self.base_dir, tpl)
        else:
            self.template_path = self._find_template()

        # 输出
        suffix = self.raw.get("output_suffix", "_正式报告")
        tpl_name = os.path.splitext(os.path.basename(self.template_path))[0]
        self.output_path = os.path.join(self.base_dir, f"{tpl_name}{suffix}.docx")

        # 表格索引覆盖
        if "table_indices" in self.raw:
            self.table_indices.update(self.raw["table_indices"])

        # 地层配置
        self.layers = self.raw.get("layers", [])

        # 全局替换
        self.replacements = [tuple(r) for r in self.raw.get("replacements", [])]

        # 含水层/液化
        self.water_aquifer_layers = self.raw.get("water_aquifer_layers", [])
        self.liquefaction_layers = self.raw.get("liquefaction_layers", {})

        # 段落模板
        wp = self.raw.get("workload_paragraph", {})
        self.workload_template = wp.get("template", "")
        self.water_paragraph_template = self.raw.get("water_paragraph_template", "")

        logger.info(f"  项目: {self.project_name}")
        logger.info(f"  目录: {self.base_dir}")
        logger.info(f"  Excel: {self.excel_dir}")
        logger.info(f"  模板: {self.template_path}")
        logger.info(f"  输出: {self.output_path}")

    def _load_from_directory(self, project_dir: str) -> None:
        """从项目目录自动检测 (向后兼容 v3)"""
        if not os.path.isdir(project_dir):
            raise FileNotFoundError(f"项目目录不存在: {project_dir}")

        self.base_dir = project_dir
        self.project_name = os.path.basename(project_dir)
        self.excel_dir = self._find_excel_dir()
        self.template_path = self._find_template()

        tpl_name = os.path.splitext(os.path.basename(self.template_path))[0]
        self.output_path = os.path.join(self.base_dir, f"{tpl_name}_正式报告.docx")

    def _find_excel_dir(self) -> str:
        """自动查找 Excel 数据目录"""
        candidates = [
            os.path.join(self.base_dir, "已有资料", "excel表格"),
            os.path.join(self.base_dir, "excel表格"),
            os.path.join(self.base_dir, "已有资料"),
            self.base_dir,
        ]
        for d in candidates:
            if not os.path.isdir(d):
                continue
            xls_files = (
                glob.glob(os.path.join(d, "*.XLS"))
                + glob.glob(os.path.join(d, "*.xls"))
                + glob.glob(os.path.join(d, "*.xlsx"))
            )
            if len(xls_files) >= 3:
                return d
        # 递归查找
        for root, _dirs, files in os.walk(self.base_dir):
            xls_count = sum(1 for f in files if f.lower().endswith((".xls", ".xlsx")))
            if xls_count >= 3:
                return root
        return self.base_dir

    def _find_template(self) -> str:
        """查找 .docx 模板文件"""
        for f in os.listdir(self.base_dir):
            fp = os.path.join(self.base_dir, f)
            if (
                f.endswith(".docx")
                and "正式报告" not in f
                and "自动生成" not in f
                and "temp" not in f.lower()
            ):
                return fp
        # 递归查找
        for root, _dirs, files in os.walk(self.base_dir):
            for f in files:
                if (
                    f.endswith(".docx")
                    and "正式报告" not in f
                    and "自动生成" not in f
                ):
                    return os.path.join(root, f)
        raise FileNotFoundError(f"未找到 .docx 模板文件: {self.base_dir}")

    def get_layer_config(self, layer_id: str) -> Optional[Dict[str, Any]]:
        """获取指定地层的配置"""
        for layer in self.layers:
            if str(layer.get("id", "")) == str(layer_id):
                return layer
        return None

    def get_layer_description(self, layer_id: str) -> Optional[str]:
        """获取地层描述模板"""
        cfg = self.get_layer_config(layer_id)
        return cfg.get("description") if cfg else None

    def get_bearing_values(self, layer_id: str) -> Optional[Dict[str, str]]:
        """获取承载力参数"""
        cfg = self.get_layer_config(layer_id)
        return cfg.get("bearing") if cfg else None

    def get_pile_values(self, layer_id: str) -> Optional[Dict[str, str]]:
        """获取桩基参数"""
        cfg = self.get_layer_config(layer_id)
        return cfg.get("pile") if cfg else None

    def get_workload_config(self) -> Dict[str, Any]:
        """获取工作量段落配置"""
        return self.raw.get("workload_paragraph", {})

    def get_date_replacements(self) -> List[Tuple[str, str]]:
        """获取日期替换规则"""
        return [tuple(r) for r in self.raw.get("date_replacements", [])]

    def get_technical_standards(self) -> Dict[str, Any]:
        """获取技术标准配置"""
        return self.raw.get("technical_standards", {})

    def get_site_evaluation(self) -> Dict[str, Any]:
        """获取场地稳定性及适宜性评价配置"""
        return self.raw.get("site_evaluation", {})

    def get_foundation_evaluation(self) -> Dict[str, Any]:
        """获取地基评价配置 (4.5.7 §2 均匀性 / §5 软弱下卧层 / §6 变形参数)"""
        return self.raw.get("foundation_evaluation", {})

    def get_project_overview(self) -> Dict[str, Any]:
        """获取拟建工程概况配置 (第一章)"""
        return self.raw.get("project_overview", {})

    def get_site_conditions(self) -> Dict[str, Any]:
        """获取场地条件配置 (第五章: 地形地貌/地下水/地震/不良地质)"""
        return self.raw.get("site_conditions", {})

    def get_analysis_evaluation(self) -> Dict[str, Any]:
        """获取岩土工程分析评价配置 (第六章各子节段落)"""
        return self.raw.get("analysis_evaluation", {})

    def get_conclusion_suggestions(self) -> Dict[str, Any]:
        """获取结论与建议配置 (第七章: conclusion/suggestions 段落数组)"""
        return self.raw.get("conclusion_suggestions", {})

    def get_hn_db_dir(self) -> str:
        """获取华宁数据库目录 (绝对路径或相对于 base_dir)"""
        d = self.raw.get("hn_db_dir", "")
        if d and os.path.isabs(d):
            return d
        elif d:
            return os.path.join(self.base_dir, d)
        return ""

    def get_geo_age_names(self) -> Dict[str, str]:
        """获取地质年代名称覆盖配置 {group_key: display_name}"""
        return self.raw.get("geo_age_names", {})

    def get_hn_project_code(self) -> str:
        """获取华宁项目代码"""
        return self.raw.get("hn_project_code", "")

    def get_sj_dir(self) -> str:
        """获取 sj 目录路径 (简报 docx 所在位置)

        优先级:
        1. 显式配置 ``sj_dir``
        2. ``hn_db_dir`` 目录 (通常简报也在同一目录)
        3. 在 ``base_dir`` 下搜索 sj / SJ 子目录
        """
        d = self.raw.get("sj_dir", "")
        if d:
            return d if os.path.isabs(d) else os.path.join(self.base_dir, d)

        # 回退到 hn_db_dir
        hn = self.get_hn_db_dir()
        if hn and os.path.isdir(hn):
            return hn

        # 在 base_dir 下搜索 sj/SJ
        for sub in ("sj", "SJ"):
            candidate = os.path.join(self.base_dir, sub)
            if os.path.isdir(candidate):
                return candidate

        return ""

    def get_seismic_params(self) -> Dict[str, str]:
        """
        获取地震动参数 (自动查表 + config覆盖)。

        config 中 seismic 节:
            { "district": "环翠区", "town": "怡园街道" }
        自动从 GB 18306-2015 附录C.15 查出 pga/period/intensity/group。
        也可在 config 中直接覆盖任意字段 (如 {"pga": "0.15"})。
        """
        seismic_cfg = self.raw.get("seismic", {})
        if not seismic_cfg:
            return {}

        district = seismic_cfg.get("district", "")
        town = seismic_cfg.get("town", "")

        # 先查表
        result: Dict[str, str] = {}
        if _lookup_seismic and district and town:
            looked = _lookup_seismic(district, town)
            if looked:
                result.update(looked)
                logger.info(f"  地震参数查表: {district} {town} → "
                            f"PGA={looked['pga']}g, T={looked['period']}s, "
                            f"烈度={looked['intensity']}度, 分组={looked['group']}")
            else:
                logger.warning(f"  地震参数查表: {district} {town} 未找到, "
                               f"请检查 seismic.district 和 seismic.town 配置")

        # config 中的显式值覆盖查表结果 (允许部分覆盖)
        for key in ("pga", "period", "intensity", "group"):
            if key in seismic_cfg:
                result[key] = str(seismic_cfg[key])

        return result

    def get_dxf_path(self) -> str:
        """获取 DXF 总平面图路径

        优先级:
        1. config 中 ``dxf_path`` 显式指定
        2. ``base_dir``/图/ 下搜索 .dxf 文件 (取含 "总图" 或 "平面图" 的第一个)
        """
        dxf = self.raw.get("dxf_path", "")
        if dxf:
            return dxf if os.path.isabs(dxf) else os.path.join(self.base_dir, dxf)

        # 在 图/ 子目录搜索
        tu_dir = os.path.join(self.base_dir, "图")
        if os.path.isdir(tu_dir):
            candidates = []
            for f in os.listdir(tu_dir):
                if f.lower().endswith(".dxf"):
                    candidates.append(os.path.join(tu_dir, f))
            # 优先含"总图"或"平面图"关键字
            preferred = [
                c for c in candidates
                if "总图" in os.path.basename(c) or "平面图" in os.path.basename(c)
            ]
            if preferred:
                return preferred[0]
            elif candidates:
                return candidates[0]

        return ""

    def get_building_info_path(self) -> str:
        """获取建筑信息 JSON 路径

        优先级:
        1. config 中 ``building_info_path`` 显式指定
        2. DXF 同目录下 ``*_建筑信息提取.json``
        3. ``base_dir`` 下搜索 ``*_建筑信息提取.json``
        """
        bi = self.raw.get("building_info_path", "")
        if bi:
            return bi if os.path.isabs(bi) else os.path.join(self.base_dir, bi)

        # DXF 同目录搜索
        dxf_path = self.get_dxf_path()
        if dxf_path:
            dxf_dir = os.path.dirname(dxf_path)
            dxf_base = os.path.splitext(os.path.basename(dxf_path))[0]
            candidate = os.path.join(dxf_dir, f"{dxf_base}_建筑信息提取.json")
            if os.path.isfile(candidate):
                return candidate

        # base_dir 递归搜索
        for root, _dirs, files in os.walk(self.base_dir):
            for f in files:
                if "建筑信息提取" in f and f.endswith(".json"):
                    return os.path.join(root, f)

        return ""


# ============================================================
# 数据加载器
# ============================================================

class SurveyDataLoader:
    """从华宁导出 Excel 中加载所有勘察数据"""

    def __init__(self, excel_dir: str):
        self.excel_dir = excel_dir

    # ---- 通用读取 ----

    @staticmethod
    def _read_sheet(path: str) -> List[List[Any]]:
        """读取 xls 或 xlsx 文件的首个工作表"""
        if path.lower().endswith(".xlsx"):
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            ws = wb.worksheets[0]
            rows = [list(row) for row in ws.iter_rows(values_only=True)]
            wb.close()
            return rows
        else:
            wb = xlrd.open_workbook(path)
            ws = wb.sheet_by_index(0)
            return [
                [ws.cell(r, c).value for c in range(ws.ncols)]
                for r in range(ws.nrows)
            ]

    # ---- 勘探点一览表 ----

    def load_boreholes(self) -> List[Dict[str, Any]]:
        """勘探点一览表 → 钻孔列表"""
        path = find_file(self.excel_dir, "勘探点一览表")
        if not path:
            raise FileNotFoundError("未找到 勘探点一览表*.XLS")

        rows = self._read_sheet(path)
        boreholes: List[Dict[str, Any]] = []

        for r in range(6, len(rows)):
            seq = safe_str(rows[r][0]) if len(rows[r]) > 0 else ""
            if not seq:
                continue
            if seq in ("合计", "总 计") or seq.startswith("="):
                break
            if len(rows[r]) < 14:
                continue

            boreholes.append({
                "seq": seq,
                "id": safe_str(rows[r][1]),
                "type": safe_str(rows[r][2]),
                "elevation": safe_float(rows[r][3]),
                "depth": safe_float(rows[r][4]),
                "x": safe_float(rows[r][5]),
                "y": safe_float(rows[r][6]),
                "wt_depth": safe_float(rows[r][7]),
                "wt_elv": safe_float(rows[r][8]),
                "undisturbed": int(safe_float(rows[r][9], 0)),
                "disturbed": int(safe_float(rows[r][10], 0)),
                "rock": int(safe_float(rows[r][11], 0)),
                "spt": int(safe_float(rows[r][12], 0)),
                "n63": int(safe_float(rows[r][13], 0)),
            })

        logger.info(f"  勘探点: {len(boreholes)} 个")
        return boreholes

    @staticmethod
    def classify_boreholes(
        boreholes: List[Dict[str, Any]],
    ) -> Dict[str, Any]:
        """智能分类钻孔类型并汇总统计"""
        types: Dict[str, int] = defaultdict(int)
        for bh in boreholes:
            types[bh["type"]] += 1

        result: Dict[str, Any] = {
            "total": len(boreholes),
            "total_depth": sum(bh["depth"] or 0 for bh in boreholes),
        }
        result["qutu"] = types.get("取土孔", 0)
        result["biaoguan"] = types.get("标贯孔", 0)
        result["yiban"] = types.get("一般性钻孔", 0)
        result["zhongtan"] = types.get("重探孔", types.get("重探", 0))
        result["bosk"] = types.get("波速孔", 0)
        result["ctrl"] = result["qutu"] + result["zhongtan"] + result["bosk"]
        result["general"] = result["total"] - result["ctrl"]
        result["undisturbed"] = sum(bh["undisturbed"] for bh in boreholes)
        result["disturbed"] = sum(bh["disturbed"] for bh in boreholes)
        result["rock"] = sum(bh["rock"] for bh in boreholes)
        result["spt_total"] = sum(bh["spt"] for bh in boreholes)
        result["n63_total"] = sum(bh["n63"] for bh in boreholes)
        result["all_types"] = dict(types)

        # 水位统计
        wtd = [bh["wt_depth"] for bh in boreholes if bh["wt_depth"]]
        wte = [bh["wt_elv"] for bh in boreholes if bh["wt_elv"]]
        if wtd:
            result["wt_depth_min"] = min(wtd)
            result["wt_depth_max"] = max(wtd)
            result["wt_elv_min"] = min(wte)
            result["wt_elv_max"] = max(wte)

        # 高程统计
        elvs = [bh["elevation"] for bh in boreholes if bh["elevation"]]
        if elvs:
            result["elv_min"] = min(elvs)
            result["elv_max"] = max(elvs)

        return result

    # ---- 地层厚度统计 ----

    def load_layer_stats(self) -> Dict[str, Dict[str, Any]]:
        """场地地层厚度统计表 → 地层统计数据"""
        path = find_file(self.excel_dir, "地层厚度统计") or find_file(
            self.excel_dir, "厚度层底深度"
        )
        if not path:
            logger.warning("  [!] 未找到地层厚度统计表，跳过")
            return {}

        rows = self._read_sheet(path)
        layers: Dict[str, Dict[str, Any]] = {}

        for r in range(4, len(rows)):
            lid = safe_str(rows[r][0])
            if not lid:
                continue
            if len(rows[r]) < 17:
                continue
            layers[lid] = {
                "thick_min": safe_float(rows[r][1]),
                "thick_max": safe_float(rows[r][2]),
                "thick_avg": safe_float(rows[r][3]),
                "depth_min": safe_float(rows[r][4]),
                "depth_max": safe_float(rows[r][5]),
                "depth_avg": safe_float(rows[r][6]),
                "elv_min": safe_float(rows[r][7]),
                "elv_max": safe_float(rows[r][8]),
                "elv_avg": safe_float(rows[r][9]),
                "n": safe_float(rows[r][10]),
                "top_d_min": safe_float(rows[r][11]),
                "top_d_max": safe_float(rows[r][12]),
                "top_d_avg": safe_float(rows[r][13]),
                "top_e_min": safe_float(rows[r][14]),
                "top_e_max": safe_float(rows[r][15]),
                "top_e_avg": safe_float(rows[r][16]),
            }

        logger.info(f"  地层统计: {len(layers)} 层")
        return layers

    # ---- 物理力学性质指标统计 ----

    def load_physical_stats(self) -> Dict[str, Dict[str, Any]]:
        """物理力学性质指标统计表"""
        path = find_file(self.excel_dir, "物理力学")
        if not path:
            logger.warning("  [!] 未找到物理力学统计表，跳过")
            return {}

        wb = xlrd.open_workbook(path)
        ws = wb.sheet_by_index(0)
        phys: Dict[str, Dict[str, Any]] = {}
        current: Optional[str] = None

        for r in range(8, ws.nrows):
            lid = safe_str(ws.cell(r, 0).value)
            lname = safe_str(ws.cell(r, 1).value)
            stat = safe_str(ws.cell(r, 2).value).replace(" ", "")

            if lid:
                current = lid
            if not current:
                continue
            if current not in phys:
                phys[current] = {"name": lname, "stats": {}}

            if stat in STAT_KEYWORD_MAP:
                sk = STAT_KEYWORD_MAP[stat]
                phys[current]["stats"][sk] = {
                    "W": safe_float(ws.cell(r, 6).value),
                    "gamma": safe_float(ws.cell(r, 8).value),
                    "e0": safe_float(ws.cell(r, 10).value),
                    "WL": safe_float(ws.cell(r, 12).value),
                    "WP": safe_float(ws.cell(r, 13).value),
                    "IP": safe_float(ws.cell(r, 14).value),
                    "IL": safe_float(ws.cell(r, 15).value),
                }

        return phys

    # ---- 标准贯入试验 ----

    def load_spt_stats(self) -> Dict[str, Dict[str, Any]]:
        """标准贯入试验成果统计表"""
        path = find_file(self.excel_dir, "标准贯入")
        if not path:
            logger.warning("  [!] 未找到标贯统计表，跳过")
            return {}

        wb = xlrd.open_workbook(path)
        ws = wb.sheet_by_index(0)
        spt: Dict[str, Dict[str, Any]] = {}
        current: Optional[str] = None

        for r in range(6, ws.nrows):
            lid = safe_str(ws.cell(r, 0).value)
            stat = safe_str(ws.cell(r, 1).value).replace(" ", "")

            if lid and (
                lid.replace("-", "").replace(".", "").isdigit() or "-" in lid
            ):
                current = lid
            if not current:
                continue
            if stat in STAT_KEYWORD_MAP:
                spt.setdefault(current, {})[STAT_KEYWORD_MAP[stat]] = {
                    "raw": safe_float(ws.cell(r, 7).value),
                    "corr": safe_float(ws.cell(r, 8).value),
                }

        return spt

    # ---- 动力触探 ----

    def load_cpt_stats(self) -> Dict[str, Dict[str, Any]]:
        """动力触探 N63.5 试验成果统计表"""
        path = find_file(self.excel_dir, "动力触探")
        if not path:
            logger.warning("  [!] 未找到动探统计表，跳过")
            return {}

        wb = xlrd.open_workbook(path)
        ws = wb.sheet_by_index(0)
        cpt: Dict[str, Dict[str, Any]] = {}
        current: Optional[str] = None

        for r in range(ws.nrows):
            lid = safe_str(ws.cell(r, 0).value)
            stat = safe_str(ws.cell(r, 1).value).replace(" ", "")

            if lid and (lid.isdigit() or "-" in lid):
                current = lid
            if not current:
                continue
            if stat in STAT_KEYWORD_MAP:
                cpt.setdefault(current, {})[STAT_KEYWORD_MAP[stat]] = {
                    "raw": safe_float(ws.cell(r, 5).value),
                    "corr": safe_float(ws.cell(r, 6).value),
                }

        return cpt

    # ---- 液化判别 ----

    def load_liquefaction(self) -> Tuple[List[List[Any]], int, int]:
        """液化判别及液化指数计算成果表"""
        path = find_file(self.excel_dir, "液化")
        if not path:
            return [], 0, 0

        wb = xlrd.open_workbook(path)
        ws = wb.sheet_by_index(0)
        data: List[List[Any]] = []

        for r in range(6, ws.nrows):
            if not safe_str(ws.cell(r, 0).value):
                continue
            data.append([ws.cell(r, c).value for c in range(min(20, ws.ncols))])

        liq_count = sum(
            1 for row in data if len(row) > 10 and "液" in safe_str(row[10])
        )
        return data, liq_count, len(data) - liq_count

    # ---- 岩石试验 ----

    def load_rock_stats(self) -> Dict[str, Dict[str, Any]]:
        """岩石试验指标分层统计表"""
        path = find_file(self.excel_dir, "岩石试验")
        if not path:
            return {}

        wb = xlrd.open_workbook(path)
        ws = wb.sheet_by_index(0)
        rock: Dict[str, Dict[str, Any]] = {}
        current: Optional[str] = None

        for r in range(7, ws.nrows):
            lid = safe_str(ws.cell(r, 0).value)
            stat = safe_str(ws.cell(r, 1).value).replace(" ", "")

            if lid:
                current = lid
            if not current:
                continue
            if stat in STAT_KEYWORD_MAP:
                rock.setdefault(current, {})[STAT_KEYWORD_MAP[stat]] = safe_float(
                    ws.cell(r, 6).value
                )

        return rock

    # ---- 建筑物特征 ----

    def load_buildings(self) -> List[Dict[str, str]]:
        """建筑物特征一览表"""
        path = find_file(self.excel_dir, "建筑物特征")
        if not path:
            return []

        try:
            wb = openpyxl.load_workbook(path, data_only=True)
            ws = wb.worksheets[0]
            buildings: List[Dict[str, str]] = []

            for row in ws.iter_rows(values_only=True):
                row_list = list(row)
                name = safe_str(row_list[0])
                if not name or name in ("合计", "说明", "备注"):
                    continue
                if "名称" in name or "序号" in name:
                    continue
                buildings.append({
                    "name": name,
                    "floors": safe_str(row_list[2]) if len(row_list) > 2 else "",
                    "height": safe_str(row_list[3]) if len(row_list) > 3 else "",
                    "size": safe_str(row_list[4]) if len(row_list) > 4 else "",
                    "span": safe_str(row_list[5]) if len(row_list) > 5 else "",
                    "indoor_elv": safe_str(row_list[6]) if len(row_list) > 6 else "",
                    "structure": safe_str(row_list[7]) if len(row_list) > 7 else "",
                })

            wb.close()
            logger.info(f"  建筑物: {len(buildings)} 栋")
            return buildings
        except Exception as e:
            logger.error(f"  [!] 建筑物特征读取异常: {e}")
            return []

    # ---- 水样 (支持 A3 双面板) ----

    def load_water_samples(self) -> List[Dict[str, Any]]:
        """水样 xlsx 解析 (支持 A3 双面板格式)"""
        path = find_file(self.excel_dir, "水样")
        if not path:
            return []

        def parse_panel(
            label_col: int, name_col: int, val_col: int, extra_col: int,
            rows: List[List[Any]],
        ) -> Dict[str, Any]:
            result: Dict[str, Any] = {}
            # 从 label 列提取野编号
            for r in rows[:7]:
                txt = safe_str(r[label_col]) if len(r) > label_col else ""
                if "野编号" in txt:
                    m = re.search(r"(\d+)", txt)
                    if m:
                        result["field_id"] = m.group(1)
            # 解析离子含量
            for r in rows:
                if len(r) <= max(name_col, val_col, extra_col):
                    continue
                name = safe_str(r[name_col])
                val = safe_float(r[val_col]) if len(r) > val_col else None
                extra_lbl = (
                    safe_str(r[extra_col - 1]) if len(r) > extra_col - 1 else ""
                )
                extra_val = safe_float(r[extra_col]) if len(r) > extra_col else None

                if name == "PH值":
                    result["pH"] = val
                elif name == "Mg2+":
                    result["Mg"] = val
                elif name == "NH4+":
                    result["NH4"] = val or 0
                elif name == "Cl-":
                    result["Cl"] = val
                elif name == "SO42-":
                    result["SO4"] = val
                elif name == "HCO3-":
                    result["HCO3"] = val
                elif name == "CO32-":
                    result["CO3"] = val or 0
                elif name == "OH-":
                    result["OH"] = val or 0

                if safe_str(r[label_col]) == "PH值":
                    result["pH"] = safe_float(r[val_col])
                if "总矿化度" in extra_lbl:
                    result["TDS"] = extra_val
                if "侵蚀" in extra_lbl:
                    result["CO2_agg"] = extra_val or 0
            return result

        try:
            wb = openpyxl.load_workbook(path, data_only=True)
            ws = wb.worksheets[0]
            rows = [list(r) for r in ws.iter_rows(values_only=True)]
            wb.close()

            samples: List[Dict[str, Any]] = []
            # 左面板: label=0, name=1, val=2, extra=6
            # 右面板: label=9, name=10, val=11, extra=15
            panel_configs = [(0, 1, 2, 6), (9, 10, 11, 15)]
            for cfg in panel_configs:
                w = parse_panel(*cfg, rows)
                if w.get("SO4"):
                    samples.append(w)

            # 补充野编号
            if len(rows) > 4:
                for cfg_idx, label_col in [(0, 0), (1, 9)]:
                    if cfg_idx < len(samples):
                        txt = (
                            safe_str(rows[4][label_col])
                            if len(rows[4]) > label_col
                            else ""
                        )
                        m = re.search(r"(\d+)", txt)
                        if m:
                            samples[cfg_idx]["field_id"] = m.group(1)

            logger.info(f"  水样: {len(samples)} 件")
            return samples
        except Exception as e:
            logger.error(f"  [!] 水样读取异常: {e}")
            return []

    # ---- 易溶盐 (支持 A3 双面板) ----

    def load_salt_samples(self) -> List[Dict[str, Any]]:
        """易溶盐土样 xlsx 解析 (支持 A3 双面板格式)"""
        path = find_file(self.excel_dir, "易溶盐")
        if not path:
            return []

        def parse_panel(
            name_a: int, val_a: int, name_b: int, val_b: int,
            rows: List[List[Any]],
        ) -> Dict[str, Any]:
            result: Dict[str, Any] = {}
            for r in rows:
                if len(r) <= val_b:
                    continue
                for nc, vc in [(name_a, val_a), (name_b, val_b)]:
                    if len(r) <= nc:
                        continue
                    ion = safe_str(r[nc])
                    v = safe_float(r[vc]) if len(r) > vc else None
                    if ion == "K+":
                        result["K"] = v
                    elif ion == "Na+":
                        result["Na"] = v
                    elif ion == "Ca2+":
                        result["Ca"] = v
                    elif ion == "Mg2+":
                        result["Mg"] = v
                    elif ion == "Cl-":
                        result["Cl"] = v
                    elif ion == "SO42-":
                        result["SO4"] = v
                    elif ion == "HCO3-":
                        result["HCO3"] = v
                    elif ion == "CO32-":
                        result["CO3"] = v or 0
                    elif ion == "OH-":
                        result["OH"] = v or 0

                # 检查总含盐量和 pH (左侧)
                txt0 = safe_str(r[name_a - 1]) if name_a > 0 and len(r) > name_a - 1 else ""
                if "总含盐量" in txt0:
                    result["TDS"] = safe_float(r[name_a]) if len(r) > name_a else None
                if "PH值" in txt0:
                    result["pH"] = safe_float(r[name_a]) if len(r) > name_a else None
                # 检查右侧
                txt_r = safe_str(r[name_a + 5]) if len(r) > name_a + 5 else ""
                if "总含盐量" in txt_r:
                    result["TDS"] = (
                        safe_float(r[name_a + 6]) if len(r) > name_a + 6 else None
                    )
                if "PH值" in txt_r:
                    result["pH"] = (
                        safe_float(r[name_a + 6]) if len(r) > name_a + 6 else None
                    )
            return result

        try:
            wb = openpyxl.load_workbook(path, data_only=True)
            ws = wb.worksheets[0]
            rows = [list(r) for r in ws.iter_rows(values_only=True)]
            wb.close()

            samples: List[Dict[str, Any]] = []
            for cfg in [(0, 1, 3, 4), (7, 8, 10, 11)]:
                s = parse_panel(*cfg, rows)
                if s.get("SO4"):
                    samples.append(s)

            logger.info(f"  易溶盐: {len(samples)} 件")
            return samples
        except Exception as e:
            logger.error(f"  [!] 易溶盐读取异常: {e}")
            return []


# ============================================================
# 华宁勘察数据库读取器
# ============================================================

class HuaNingDBReader:
    """读取华宁 HNCAD 勘察软件原始数据库文件

    华宁数据库文件命名: FileType.ProjectCode (如 ZHMS.17, DCSH.17)
    编码: GB2312
    """

    # 成因代号 → 中文名
    ORIGIN_NAMES = {
        "ml": "人工堆积层",
        "al": "冲积层",
        "pl": "洪积层",
        "al+pl": "冲洪积层",
        "dl": "坡积层",
        "el": "残积层",
        "dl+el": "坡残积层",
        "eol": "风积层",
        "l": "湖积层",
        "m": "海积层",
        "mc": "海陆交互相沉积层",
        "fgl": "冰水沉积层",
    }

    # 年代代号 → 中文前缀
    AGE_PREFIXES = {
        "Q4": "第四系全新统",
        "Q3": "第四系上更新统",
        "Q2": "第四系中更新统",
        "Q1": "第四系下更新统",
        "N": "新近系",
        "Nh": "新元古界",
        "Pt": "元古界",
        "Ar": "太古界",
    }

    # 岩石代码 → 岩石名称 (5位岩性码 = 风化程度2位 + 岩石类型3位)
    ROCK_TYPE_MAP = {
        "150": "花岗片麻岩",
        "151": "花岗岩",
        "100": "砂岩",
        "101": "泥岩",
        "102": "页岩",
        "103": "石灰岩",
        "110": "片麻岩",
        "120": "大理岩",
        "130": "板岩",
        "140": "石英岩",
    }

    # 取样类型代号
    SAMPLE_TYPE_NAMES = {
        0: "扰动样",
        1: "原状样",
        2: "标贯试样",
        3: "水样",
        4: "岩样",
    }

    def __init__(self, db_dir: str, project_code: str = ""):
        self.db_dir = db_dir
        self.project_code = project_code
        if not project_code and db_dir:
            self.project_code = self._detect_project_code()

    # ---- 内部工具方法 ----

    def _detect_project_code(self) -> str:
        """自动检测项目代码"""
        for f in os.listdir(self.db_dir):
            if "." in f and f.upper().startswith("ZHMS"):
                return f.split(".")[-1]
        for f in os.listdir(self.db_dir):
            if "." in f and not f.startswith("_") and not f.startswith("."):
                parts = f.split(".")
                if len(parts) == 2 and parts[1].isdigit():
                    return parts[1]
        return ""

    def _find_file(self, file_type: str) -> Optional[str]:
        """查找数据库文件 (大小写不敏感)"""
        target_upper = f"{file_type.upper()}.{self.project_code}"
        for f in os.listdir(self.db_dir):
            if f.upper() == target_upper:
                return os.path.join(self.db_dir, f)
        # 不区分扩展名
        for f in os.listdir(self.db_dir):
            base = f.split(".")[0].upper()
            if base == file_type.upper() and os.path.isfile(os.path.join(self.db_dir, f)):
                return os.path.join(self.db_dir, f)
        return None

    def _read_lines(self, file_type: str) -> List[List[str]]:
        """读取逗号分隔的数据库文件，自动检测编码"""
        path = self._find_file(file_type)
        if not path:
            return []
        content = None
        for enc in ("gb2312", "gbk", "utf-8"):
            try:
                with open(path, "r", encoding=enc) as f:
                    content = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        if content is None:
            logger.warning(f"    无法读取 {file_type}: 编码不支持")
            return []
        lines: List[List[str]] = []
        for line in content.strip().split("\n"):
            line = line.strip()
            if line:
                lines.append(line.split(","))
        return lines

    # ---- ZHMS: 地层描述 ----

    def read_layer_descriptions(self) -> Dict[str, Dict[str, str]]:
        """读取 ZHMS 文件 → {layer_id: {name, description}}

        格式: 层号,岩土名:描述内容
        """
        lines = self._read_lines("ZHMS")
        result: Dict[str, Dict[str, str]] = {}
        for fields in lines:
            if len(fields) < 2:
                continue
            layer_id = fields[0].strip()
            rest = ",".join(fields[1:])  # 合并可能因逗号分割的内容
            if layer_id.upper() == "END":
                break
            colon_idx = rest.find(":")
            if colon_idx > 0:
                name = rest[:colon_idx].strip()
                desc = rest[colon_idx + 1:].strip()
            else:
                name = rest.strip()
                desc = ""
            if layer_id and name:
                result[layer_id] = {"name": name, "description": desc}
        return result

    # ---- DCSH: 地层序列与地质年代 ----

    def read_layer_sequence(self) -> List[Dict[str, str]]:
        """读取 DCSH 文件 → 地层序列 (含地质年代和成因)

        格式: 层号,岩性代码,,,成因,年代,
        """
        lines = self._read_lines("DCSH")
        result: List[Dict[str, str]] = []
        for fields in lines:
            if len(fields) < 2:
                continue
            layer_id = fields[0].strip()
            if layer_id.upper() == "END":
                break
            rock_code = fields[1].strip() if len(fields) > 1 else ""
            origin = fields[4].strip() if len(fields) > 4 else ""
            age = fields[5].strip() if len(fields) > 5 else ""
            if layer_id:
                result.append({
                    "layer_id": layer_id,
                    "rock_code": rock_code,
                    "origin": origin,
                    "age": age,
                })
        return result

    # ---- DCSJ: 逐孔地层数据 ----

    def read_borehole_layers(self) -> Tuple[
        Dict[str, List[Dict[str, Any]]],
        set,
        Dict[str, float],
    ]:
        """读取 DCSJ 文件 → (borehole_layers, unpentrated_ids, max_exposed)

        格式: 孔号,层号,层底深度,
        最后一层无深度 → 未穿透
        """
        lines = self._read_lines("DCSJ")
        bh_layers: Dict[str, List[Dict[str, Any]]] = {}
        max_exposed: Dict[str, float] = {}
        current_bh: Optional[str] = None

        for fields in lines:
            if not fields:
                continue
            bh_id = fields[0].strip() if fields[0].strip() else ""
            if bh_id and bh_id.upper() != "END":
                current_bh = bh_id
            elif bh_id.upper() == "END":
                break
            if not current_bh or len(fields) < 2:
                continue

            layer_id = fields[1].strip()
            depth_str = fields[2].strip() if len(fields) > 2 else ""

            if current_bh not in bh_layers:
                bh_layers[current_bh] = []

            depth: Optional[float] = None
            if depth_str:
                try:
                    depth = float(depth_str)
                except ValueError:
                    depth = None

            bh_layers[current_bh].append({
                "layer_id": layer_id,
                "depth": depth,
            })

            # 跟踪每层的最大揭露深度
            if depth is not None:
                prev = max_exposed.get(layer_id, 0)
                if depth > prev:
                    max_exposed[layer_id] = depth

        # 检测未穿透层: 某钻孔最后一层无深度
        unpentrated: set = set()
        for bh_id, layers in bh_layers.items():
            if layers and layers[-1]["depth"] is None:
                unpentrated.add(layers[-1]["layer_id"])

        return bh_layers, unpentrated, max_exposed

    # ---- BG: 标贯数据 ----

    def read_spt_data(self) -> List[Dict[str, Any]]:
        """读取 BG 文件 → SPT 数据列表

        格式: 孔号,起始深度,终止深度,N值
        (孔号为空表示续上一行孔号)
        """
        lines = self._read_lines("BG")
        result: List[Dict[str, Any]] = []
        current_bh: Optional[str] = None

        for fields in lines:
            if not fields:
                continue
            bh_id = fields[0].strip() if fields[0].strip() else ""
            if bh_id and bh_id.upper() != "END":
                current_bh = bh_id
            elif bh_id.upper() == "END":
                break
            if not current_bh or len(fields) < 4:
                continue

            try:
                depth_start = float(fields[1])
                depth_end = float(fields[2])
                n_value = float(fields[3])
            except (ValueError, IndexError):
                continue

            result.append({
                "borehole_id": current_bh,
                "depth_start": depth_start,
                "depth_end": depth_end,
                "n_value": n_value,
            })

        return result

    # ---- TY: 取样数据 ----

    def read_sample_data(self) -> List[Dict[str, Any]]:
        """读取 TY 文件 → 取样数据列表

        格式: 孔号,深度,类型代码
        类型: 0=扰动样, 1=原状样, 2=标贯试样
        """
        lines = self._read_lines("TY")
        result: List[Dict[str, Any]] = []
        current_bh: Optional[str] = None

        for fields in lines:
            if not fields:
                continue
            bh_id = fields[0].strip() if fields[0].strip() else ""
            if bh_id and bh_id.upper() != "END":
                current_bh = bh_id
            elif bh_id.upper() == "END":
                break
            if not current_bh or len(fields) < 3:
                continue

            try:
                depth = float(fields[1])
                sample_type = int(fields[2])
            except (ValueError, IndexError):
                continue

            result.append({
                "borehole_id": current_bh,
                "depth": depth,
                "sample_type": sample_type,
            })

        return result

    # ---- 液化判别表 (XLS) ----

    # 抗震设防烈度 → 标贯击数基准值 N0 (GB/T50011-2010 表4.3.4)
    INTENSITY_TO_N0 = {6: 5, 7: 7, 8: 10, 9: 16}

    # 抗震设防烈度 → 液化土特征深度 d0 (m) (GB/T50011-2010 表4.3.3)
    INTENSITY_TO_D0 = {6: 6, 7: 7, 8: 8, 9: 9}

    def read_liquefaction_table(self) -> Dict[str, Any]:
        """读取 sj 目录下的液化判别表 XLS

        返回:
        {
            "available": True/False,
            "intensity": 7,             # 抗震设防烈度
            "n0": 7,                    # 标贯击数基准值
            "max_depth": 20,            # 判别最大深度
            "design_group": "第一组",
            "beta": 0.8,               # 调整系数
            "points": [...],           # 逐点数据
            "total_points": 23,
            "liq_points": 0,
            "non_liq_points": 23,
            "target_layers": ["2"],     # 被判别的层号集合
            "target_names": ["含黏性土砂"],
            "dw_values": [1.0],        # 水位深度 (去重)
            "ile_max": 0.0,            # 最大液化指数
        }
        """
        result: Dict[str, Any] = {"available": False}

        # 查找液化判别表 XLS
        liq_path = find_file(self.db_dir, "液化判别")
        if not liq_path:
            return result

        try:
            wb = xlrd.open_workbook(liq_path)
        except Exception as e:
            logger.warning(f"    液化判别表读取失败: {e}")
            return result

        ws = wb.sheet_by_index(0)
        if ws.nrows < 10:
            return result

        # ---- 解析表头参数 (第3行) ----
        header_text = str(ws.cell(3, 0).value) if ws.nrows > 3 else ""
        for c in range(1, min(20, ws.ncols)):
            v = ws.cell(3, c).value
            if v:
                header_text += " " + str(v)

        intensity = 7
        m = re.search(r"(\d+)\s*度", header_text)
        if m:
            intensity = int(m.group(1))

        n0 = self.INTENSITY_TO_N0.get(intensity, 7)
        m = re.search(r"N0[:\s]*(\d+)", header_text)
        if m:
            n0 = int(m.group(1))

        max_depth = 20
        m = re.search(r"(\d+)\s*米", header_text)
        if m:
            max_depth = int(m.group(1))

        design_group = ""
        m = re.search(r"(第[一二三]组)", header_text)
        if m:
            design_group = m.group(1)

        beta = 0.8
        m = re.search(r"[βΒ]\s*[:\s]*\s*([\d.]+)", header_text)
        if m:
            beta = float(m.group(1))

        # ---- 解析数据行 ----
        points: List[Dict[str, Any]] = []
        liq_count = 0
        # 用有序字典保持 layer_id → soil_name 的对应关系
        target_layer_map: Dict[str, str] = {}
        dw_values: set = set()
        current_bh = ""

        for r in range(10, ws.nrows):
            # 孔号
            bh_cell = str(ws.cell(r, 0).value).strip() if ws.cell(r, 0).value else ""
            if bh_cell:
                current_bh = bh_cell

            # 层号
            layer_cell = str(ws.cell(r, 1).value).strip() if ws.cell(r, 1).value else ""
            if not layer_cell:
                continue  # 跳过空行

            # 试验深度
            depth_str = str(ws.cell(r, 2).value).strip() if ws.cell(r, 2).value else ""

            # 岩土名称
            soil_name = str(ws.cell(r, 3).value).strip() if ws.cell(r, 3).value else ""

            # 地下水位 dW
            try:
                dw = float(ws.cell(r, 4).value) if ws.cell(r, 4).value else None
            except (ValueError, TypeError):
                dw = None
            if dw is not None:
                dw_values.add(dw)

            # 黏粒含量
            try:
                clay = float(ws.cell(r, 5).value) if ws.cell(r, 5).value else 3.0
            except (ValueError, TypeError):
                clay = 3.0

            # 实测N
            try:
                n_val = float(ws.cell(r, 6).value) if ws.cell(r, 6).value else None
            except (ValueError, TypeError):
                n_val = None

            # 修正N1
            try:
                n1_val = float(ws.cell(r, 7).value) if ws.cell(r, 7).value else None
            except (ValueError, TypeError):
                n1_val = None

            # 临界Ncr
            try:
                ncr = float(ws.cell(r, 8).value) if ws.cell(r, 8).value else None
            except (ValueError, TypeError):
                ncr = None

            # 判别结果
            verdict = str(ws.cell(r, 9).value).strip() if ws.cell(r, 9).value else ""

            is_liq = "液" in verdict and "不" not in verdict
            if is_liq:
                liq_count += 1

            # N/Ncr 比值
            try:
                n_ncr = float(ws.cell(r, 16).value) if ws.ncols > 16 and ws.cell(r, 16).value else None
            except (ValueError, TypeError):
                n_ncr = None

            # 液化指数 ILEi (单点) 和 ILE (钻孔累计)
            try:
                ile_i = float(ws.cell(r, 13).value) if ws.ncols > 13 and ws.cell(r, 13).value else None
            except (ValueError, TypeError):
                ile_i = None
            try:
                ile_cum = float(ws.cell(r, 14).value) if ws.ncols > 14 and ws.cell(r, 14).value else None
            except (ValueError, TypeError):
                ile_cum = None

            # 液化等级 (轻微/中等/严重, 仅出现在钻孔最后一个液化点)
            liq_grade = str(ws.cell(r, 15).value).strip() if ws.ncols > 15 and ws.cell(r, 15).value else ""

            if layer_cell and layer_cell not in target_layer_map and soil_name:
                target_layer_map[layer_cell] = soil_name

            points.append({
                "borehole_id": current_bh,
                "layer_id": layer_cell,
                "depth_range": depth_str,
                "soil_name": soil_name,
                "dw": dw,
                "clay_content": clay,
                "n_measured": n_val,
                "n_corrected": n1_val,
                "n_critical": ncr,
                "is_liquefied": is_liq,
                "verdict": verdict,
                "n_ncr_ratio": n_ncr,
                "ile_i": ile_i,
                "ile_cum": ile_cum,
                "liq_grade": liq_grade,
            })

        if not points:
            return result

        # ---- 汇总统计 ----

        # N/Ncr 范围 (全部判别点)
        ncr_ratios = [p["n_ncr_ratio"] for p in points if p["n_ncr_ratio"] is not None]
        ncr_min = min(ncr_ratios) if ncr_ratios else None
        ncr_max = max(ncr_ratios) if ncr_ratios else None

        # 每孔累计液化指数 ILE (取 col[14] 最大值作为该孔 ILE)
        bh_ile: Dict[str, float] = {}
        for p in points:
            if p["ile_cum"] is not None and p["borehole_id"]:
                bh_id = p["borehole_id"]
                bh_ile[bh_id] = max(bh_ile.get(bh_id, 0.0), p["ile_cum"])

        # 全场最大 ILE
        ile_max = max(bh_ile.values()) if bh_ile else 0.0

        # 液化等级 (按最大 ILE 判定: ≤6轻微, 6~18中等, >18严重)
        if ile_max <= 0:
            liq_grade_overall = ""
        elif ile_max <= 6:
            liq_grade_overall = "轻微"
        elif ile_max <= 18:
            liq_grade_overall = "中等"
        else:
            liq_grade_overall = "严重"

        # 各孔液化等级分布
        grade_dist: Dict[str, int] = {}
        for bh_id, ile_val in bh_ile.items():
            if ile_val <= 0:
                continue
            if ile_val <= 6:
                g = "轻微"
            elif ile_val <= 18:
                g = "中等"
            else:
                g = "严重"
            grade_dist[g] = grade_dist.get(g, 0) + 1

        logger.info(
            f"    液化判别表: {len(points)} 点, "
            f"液化 {liq_count}, 不液化 {len(points) - liq_count}, "
            f"目标层 {list(target_layer_map.keys())}, "
            f"ILE_max={ile_max:.2f}({liq_grade_overall}), "
            f"N/Ncr={ncr_min}~{ncr_max}"
        )

        # 按层号排序 (保持 id↔name 对应)
        sorted_ids = sorted(target_layer_map.keys(), key=lambda x: int(x) if x.isdigit() else x)
        sorted_names = [target_layer_map[lid] for lid in sorted_ids]

        return {
            "available": True,
            "intensity": intensity,
            "n0": n0,
            "max_depth": max_depth,
            "design_group": design_group,
            "beta": beta,
            "points": points,
            "total_points": len(points),
            "liq_points": liq_count,
            "non_liq_points": len(points) - liq_count,
            "target_layers": sorted_ids,
            "target_names": sorted_names,
            "dw_values": sorted(dw_values),
            "ncr_ratio_min": ncr_min,
            "ncr_ratio_max": ncr_max,
            "ile_max": ile_max,
            "liq_grade": liq_grade_overall,
            "bh_ile": bh_ile,
            "grade_dist": grade_dist,
        }

    # ---- 数据汇总 ----

    # ---- 覆盖层厚度 (du) 计算 ----

    # 砂土关键字: 层名包含这些则为砂层 (不计算为覆盖层)
    _SAND_KEYWORDS = ("砂", "砾", "碎石")

    # 淤泥质土关键字: 层名包含这些则单独扣除
    _SILT_MUCK_KEYWORDS = ("淤泥", )

    @classmethod
    def _is_sand_layer(cls, name: str) -> bool:
        """判断是否为砂土层"""
        return any(kw in name for kw in cls._SAND_KEYWORDS)

    @classmethod
    def _is_muck_layer(cls, name: str) -> bool:
        """判断是否为淤泥质土层"""
        return any(kw in name for kw in cls._SILT_MUCK_KEYWORDS)

    def compute_overburden_thickness(
        self,
        bh_layers: Dict[str, List[Dict[str, Any]]],
        descriptions: Dict[str, Dict[str, str]],
        water_depth: float,
    ) -> float:
        """计算场地上覆非液化土层厚度 du (m)

        规则:
            du = 地下水位以上的非砂层总厚度 − 其中淤泥质土的厚度
            对所有钻孔取平均值

        Args:
            bh_layers:    DCSJ 各孔地层 {bh_id: [{layer_id, depth}, ...]}
            descriptions: ZHMS 层描述 {layer_id: {name, ...}}
            water_depth:  地下水位埋深 dw (m)
        """
        du_values: List[float] = []

        for bh_id, layers in bh_layers.items():
            prev_depth = 0.0
            non_sand_thickness = 0.0
            muck_thickness = 0.0

            for layer in layers:
                bottom = layer.get("depth")
                if bottom is None:
                    continue
                lid = layer["layer_id"]
                name = descriptions.get(lid, {}).get("name", "")
                top = prev_depth

                # 只考虑水位以上的部分
                if top >= water_depth:
                    break
                effective_bottom = min(bottom, water_depth)
                thickness = effective_bottom - top

                if not self._is_sand_layer(name):
                    non_sand_thickness += thickness
                    if self._is_muck_layer(name):
                        muck_thickness += thickness

                prev_depth = bottom

            du = non_sand_thickness - muck_thickness
            if du > 0:
                du_values.append(du)

        if not du_values:
            return 0.0
        avg = sum(du_values) / len(du_values)
        logger.info(
            f"    覆盖层厚度 du: {len(du_values)} 孔参与计算, "
            f"平均 {avg:.2f}m, 范围 {min(du_values):.2f}~{max(du_values):.2f}m"
        )
        return round(avg, 1)

    def _assign_to_layer(
        self,
        depth: float,
        layer_boundaries: List[Tuple[str, Optional[float]]],
    ) -> Optional[str]:
        """根据深度确定所属地层

        layer_boundaries: [(layer_id, bottom_depth), ...]
        """
        prev_depth = 0.0
        for lid, bdepth in layer_boundaries:
            if bdepth is None:
                # 未穿透层, depth 落在该层
                return lid
            if depth <= bdepth:
                return lid
            prev_depth = bdepth
        return None

    def compute_spt_counts(
        self,
        spt_records: List[Dict[str, Any]],
        bh_layers: Dict[str, List[Dict[str, Any]]],
    ) -> Dict[str, int]:
        """统计每层标贯次数"""
        counts: Dict[str, int] = defaultdict(int)
        for rec in spt_records:
            bh_id = rec["borehole_id"]
            midpoint = (rec["depth_start"] + rec["depth_end"]) / 2.0
            boundaries = [
                (l["layer_id"], l["depth"])
                for l in bh_layers.get(bh_id, [])
            ]
            lid = self._assign_to_layer(midpoint, boundaries)
            if lid:
                counts[lid] += 1
        return dict(counts)

    def compute_sample_counts(
        self,
        sample_records: List[Dict[str, Any]],
        bh_layers: Dict[str, List[Dict[str, Any]]],
    ) -> Dict[str, Dict[int, int]]:
        """统计每层各类取样数量

        返回: {layer_id: {sample_type_code: count}}
        """
        counts: Dict[str, Dict[int, int]] = defaultdict(lambda: defaultdict(int))
        for rec in sample_records:
            bh_id = rec["borehole_id"]
            depth = rec["depth"]
            stype = rec["sample_type"]
            boundaries = [
                (l["layer_id"], l["depth"])
                for l in bh_layers.get(bh_id, [])
            ]
            lid = self._assign_to_layer(depth, boundaries)
            if lid:
                counts[lid][stype] += 1
        return {lid: dict(types) for lid, types in counts.items()}

    def build_age_groups(
        self,
        layer_sequence: List[Dict[str, str]],
    ) -> List[Tuple[str, List[str]]]:
        """构建地质年代分组 (保持 DCSH 顺序)

        返回: [(group_key, [layer_ids]), ...]
        group_key 格式: "age|origin" 或 "age" (无成因时)
        """
        groups: List[Tuple[str, List[str]]] = []
        current_key: Optional[str] = None

        for rec in layer_sequence:
            age = rec.get("age", "")
            origin = rec.get("origin", "")
            layer_id = rec["layer_id"]

            # 基岩层 origin 为纯数字 (如 "1") 时视为无成因
            if origin and not origin.isdigit():
                key = f"{age}|{origin}"
            else:
                key = age

            if key != current_key:
                groups.append((key, []))
                current_key = key

            groups[-1][1].append(layer_id)

        return groups

    def build_age_display_name(
        self,
        group_key: str,
        layer_names: Dict[str, str],
        config_overrides: Optional[Dict[str, str]] = None,
    ) -> str:
        """构建地质年代显示名 (如 '第四系人工堆积层（Q4ml）')"""
        # 先检查配置覆盖
        if config_overrides and group_key in config_overrides:
            return config_overrides[group_key]

        parts = group_key.split("|")
        age = parts[0]
        origin = parts[1] if len(parts) > 1 else ""

        # 基岩 (无成因)
        if not origin:
            rock_name = ""
            for lid in layer_names:
                name = layer_names.get(lid, "")
                for prefix in ("强风化", "中风化", "微风化", "碎块状强风化",
                               "全风化", "弱风化"):
                    if name.startswith(prefix):
                        rock_name = name[len(prefix):]
                        break
                if rock_name:
                    break
            # 尝试从岩石代码映射
            if not rock_name:
                rock_name = self.ROCK_TYPE_MAP.get("150", "花岗片麻岩")
            return f"{self.AGE_PREFIXES.get(age, age)}{rock_name}（{age}）"

        # 第四系沉积层
        age_prefix = self.AGE_PREFIXES.get(age, age)
        origin_name = self.ORIGIN_NAMES.get(origin, origin + "层")
        return f"{age_prefix}{origin_name}（{age}{origin}）"

    # ---- 主读取入口 ----

    def read(self) -> Dict[str, Any]:
        """读取华宁数据库全部数据

        返回:
        {
            "available": True/False,
            "descriptions": {layer_id: {name, description}},
            "layer_sequence": [{layer_id, rock_code, origin, age}],
            "age_groups": [(group_key, [layer_ids])],
            "borehole_layers": {bh_id: [{layer_id, depth}]},
            "unpentrated": set(),
            "max_exposed": {layer_id: max_depth},
            "spt_records": [...],
            "sample_records": [...],
            "spt_counts": {layer_id: count},
            "sample_counts": {layer_id: {type: count}},
        }
        """
        descriptions = self.read_layer_descriptions()
        if not descriptions:
            logger.info("    华宁数据库: ZHMS 文件不存在或为空")
            return {"available": False}

        layer_sequence = self.read_layer_sequence()
        bh_layers, unpentrated, max_exposed = self.read_borehole_layers()
        spt_records = self.read_spt_data()
        sample_records = self.read_sample_data()
        liq_table = self.read_liquefaction_table()

        age_groups = self.build_age_groups(layer_sequence)
        spt_counts = self.compute_spt_counts(spt_records, bh_layers)
        sample_counts = self.compute_sample_counts(sample_records, bh_layers)

        logger.info(
            f"    华宁数据库: {len(descriptions)} 层描述, "
            f"{len(layer_sequence)} 层序列, "
            f"{len(bh_layers)} 孔, "
            f"{len(spt_records)} 条标贯, "
            f"{len(sample_records)} 件取样"
        )

        return {
            "available": True,
            "descriptions": descriptions,
            "layer_sequence": layer_sequence,
            "age_groups": age_groups,
            "borehole_layers": bh_layers,
            "unpentrated": unpentrated,
            "max_exposed": max_exposed,
            "spt_records": spt_records,
            "sample_records": sample_records,
            "spt_counts": spt_counts,
            "sample_counts": sample_counts,
            "liquefaction_table": liq_table,
            "elevation": self.read_elevation(),
            "borehole_coords": self.read_borehole_coords(),
        }

    # ---- DK 文件: 钻孔高程 ----

    def read_elevation(self) -> Dict[str, Any]:
        """读取 DK 文件 → 孔口高程统计

        DK 文件格式: 孔号,地面高程,孔深,,水位深度,日期1,日期2,X,Y,,N
        返回: {"elv_min": float, "elv_max": float, "elv_range": float}
        """
        rows = self._read_lines("DK")
        if not rows:
            return {}

        elevations: List[float] = []
        for fields in rows:
            if len(fields) < 3:
                continue
            try:
                elv = float(fields[1])
                elevations.append(elv)
            except (ValueError, TypeError):
                continue

        if not elevations:
            return {}

        elv_min = min(elevations)
        elv_max = max(elevations)
        return {
            "elv_min": elv_min,
            "elv_max": elv_max,
            "elv_range": elv_max - elv_min,
            "count": len(elevations),
        }

    def read_borehole_coords(self) -> Dict[str, Tuple[float, float]]:
        """读取 DK 文件 → 钻孔坐标 {bh_id: (X, Y)}

        DK 文件格式: 孔号,地面高程,孔深,,水位深度,日期1,日期2,X,Y,,N
        字段索引:     0    1      2   3  4      5     6    7 8  9  10
        """
        rows = self._read_lines("DK")
        if not rows:
            return {}

        coords: Dict[str, Tuple[float, float]] = {}
        for fields in rows:
            if len(fields) < 9:
                continue
            try:
                bh_id = fields[0].strip()
                x = float(fields[7])
                y = float(fields[8])
                if x > 0 and y > 0:
                    coords[bh_id] = (x, y)
            except (ValueError, TypeError):
                continue

        return coords

    # ---- 地貌类型自动判定 ----

    @staticmethod
    def determine_terrain(
        elv_range: float,
        origins: set,
        layer_names: List[str],
    ) -> str:
        """根据高程差、成因类型、地层组合自动判定地貌类型

        规则:
        1. 剥蚀丘陵: 高差>10m 且无 mc/m; 或地层仅有填土+残积土+基岩
        2. 海岸平原: 相对平缓, 只有 m 无 mc/al/pl
        3. 山区平原: 有陆相冲洪积层 (al/pl/al+pl) 无海相
        4. 山前海积、冲洪积小平原交界地带: 陆相和海相均有
        """
        has_mc = "mc" in origins
        has_m = "m" in origins
        has_marine = has_mc or has_m
        has_al = "al" in origins or "al+pl" in origins
        has_pl = "pl" in origins
        has_continental = has_al or has_pl

        # 检查地层是否仅有 填土+残积土+基岩
        fill_keywords = ("填", "素填", "杂填", "冲填")
        residual_keywords = ("残积",)
        rock_keywords = ("风化", "岩", "基岩")
        only_fill_residual_rock = True
        for name in layer_names:
            is_fill = any(k in name for k in fill_keywords)
            is_residual = any(k in name for k in residual_keywords)
            is_rock = any(k in name for k in rock_keywords)
            if not (is_fill or is_residual or is_rock):
                only_fill_residual_rock = False
                break

        # 规则1: 剥蚀丘陵
        if elv_range > 10 and not has_marine:
            return "剥蚀丘陵"
        if layer_names and only_fill_residual_rock:
            return "剥蚀丘陵"

        # 规则4: 陆相和海相均有 → 交界地带
        if has_continental and has_marine:
            return "山前海积、冲洪积小平原交界地带"

        # 规则2: 只有海相 m, 无 mc/al/pl → 海岸平原
        if has_m and not has_mc and not has_continental:
            return "海岸平原"

        # 规则3: 有陆相冲洪积层, 无海相 → 山区平原
        if has_continental and not has_marine:
            return "山区平原"

        # 兜底: 有海陆交互 mc 但无陆相冲洪积
        if has_mc and not has_continental:
            return "海岸平原"

        return ""

    @staticmethod
    def infer_origins_from_names(
        layers: Dict[str, Any],
    ) -> set:
        """从简报地层名称和描述文本推断成因类型代码

        简报中不含 DCSH 成因代码, 需从地层特征推断:
            淤泥质/淤泥 → m (海积)
            冲洪积 → al+pl
            含黏性土砂 (含冲洪积描述) → al+pl
            残积 → el
            填土 → ml

        Parameters
        ----------
        layers : dict
            简报 ``layers`` 字段, 每层含 name 和 full_desc
        """
        origins: set = set()
        for lid, info in layers.items():
            if not isinstance(info, dict):
                continue
            name = info.get("name", "")
            desc = info.get("full_desc", "")
            combined = name + desc

            # 海相标志: 淤泥质/淤泥 → m
            if "淤泥" in combined or "淤泥质" in combined:
                origins.add("m")

            # 海陆交互: 含贝壳/腥臭 → mc
            if "贝壳" in combined or "腥臭" in combined:
                origins.add("mc")

            # 冲洪积: 描述中明确提到
            if "冲洪积" in combined:
                origins.add("al+pl")
            elif "冲积" in combined:
                origins.add("al")
            elif "洪积" in combined:
                origins.add("pl")

            # 残积
            if "残积" in combined:
                origins.add("el")

            # 填土 → ml (人工)
            if "填土" in combined or "杂填" in combined or "素填" in combined:
                origins.add("ml")

        return origins


# ============================================================
# 简报 (华宁生成 .docx) 读取
# ============================================================

def _bf(v: str) -> Optional[float]:
    """安全转 float, 空串返回 None"""
    if not v or not v.strip():
        return None
    try:
        return float(v)
    except (ValueError, TypeError):
        return None


def _bf_str(v: str) -> str:
    """安全转 float 并保留两位小数, 空串返回空串"""
    f = _bf(v)
    return f"{f:.2f}" if f is not None else ""


def _cell(tbl: Any, r: int, c: int) -> str:
    """安全读取表格单元格"""
    try:
        return tbl.rows[r].cells[c].text.strip().replace("\n", "")
    except (IndexError, AttributeError):
        return ""


def read_briefing(sj_dir: str) -> Dict[str, Any]:
    """从华宁生成的简报 docx 提取项目数据

    搜索 ``岩土工程勘察报告.docx`` (优先 sj_dir 内, 回退上级目录)。
    提取高程、地下水、地层描述、物理力学指标等。

    Parameters
    ----------
    sj_dir : str
        项目 sj/SJ 目录路径 (通常与 ``hn_db_dir`` 相同)

    Returns
    -------
    dict  包含 available, elevation, groundwater, layers, physmech,
          bearing, layer_count, project_id 等键
    """
    if not sj_dir or not os.path.isdir(sj_dir):
        return {}

    # ---- 查找文件 ----
    briefing_path = ""
    search_dirs = [sj_dir, os.path.dirname(sj_dir)]
    for d in search_dirs:
        for root, _dirs, files in os.walk(d):
            for f in files:
                if f == "岩土工程勘察报告.docx":
                    briefing_path = os.path.join(root, f)
                    break
            if briefing_path:
                break
        if briefing_path:
            break

    if not briefing_path:
        logger.info(f"  简报: 在 {sj_dir} 及上级目录未找到 岩土工程勘察报告.docx")
        return {}

    logger.info(f"  简报文件: {briefing_path}")
    doc = Document(briefing_path)

    # ---- 按文档顺序收集段落和表格 ----
    paragraphs: List[Tuple[int, str]] = []          # (body_idx, text)
    tables: List[Tuple[int, Any]] = []               # (body_idx, table)

    for idx, el in enumerate(doc.element.body):
        tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
        if tag == "p":
            for p in doc.paragraphs:
                if p._element is el:
                    t = p.text.strip()
                    if t:
                        paragraphs.append((idx, t))
                    break
        elif tag == "tbl":
            for tbl in doc.tables:
                if tbl._element is el:
                    tables.append((idx, tbl))
                    break

    # ---- 关键段落索引 ----
    terrain_idx = -1       # "2.1 地形地貌" 段落序号
    layer_start_idx = -1   # "3.1 地层结构" 段落序号
    bearing_heading = -1   # "各层建议值" 段落序号
    section4_idx = -1      # "4 岩土工程分析评价" 段落序号

    for i, (_, text) in enumerate(paragraphs):
        if "2.1" in text and "地形地貌" in text:
            terrain_idx = i
        elif "3.1" in text and "地层结构" in text:
            layer_start_idx = i
        elif "各层建议值" in text:
            bearing_heading = i
        elif text.startswith("4 ") and "岩土工程分析" in text:
            section4_idx = i

    result: Dict[str, Any] = {"available": True, "path": briefing_path}

    # ================================================================
    # 1) 高程数据  (从 "2.1 地形地貌" 段落提取)
    # ================================================================
    elv: Dict[str, Any] = {}
    if terrain_idx >= 0:
        txt = paragraphs[terrain_idx + 1][1] if terrain_idx + 1 < len(paragraphs) else ""
        m = re.search(r"标高最大值\s*([\d.]+)\s*m", txt)
        if m:
            elv["max"] = float(m.group(1))
        m = re.search(r"最小值\s*([\d.]+)\s*m", txt)
        if m:
            elv["min"] = float(m.group(1))
        m = re.search(r"高差\s*([\d.]+)\s*m", txt)
        if m:
            elv["range"] = float(m.group(1))
        elif elv.get("max") is not None and elv.get("min") is not None:
            elv["range"] = elv["max"] - elv["min"]
        # 地貌类型 (如果简报中已填写)
        m = re.search(r"地貌类型为\s*(\S+?)[。\s,，]", txt)
        if m:
            elv["terrain_type"] = m.group(1)
    result["elevation"] = elv

    # ================================================================
    # 2) 地下水  (Table 0: 7 列水位统计表)
    # ================================================================
    gw: Dict[str, Any] = {}
    for _, tbl in tables:
        if len(tbl.rows) < 2 or len(tbl.columns) < 7:
            continue
        h0 = tbl.rows[0].cells[0].text
        if "数据" not in h0 and "个数" not in h0:
            continue
        h1 = tbl.rows[0].cells[1].text
        if "稳定水位" not in h1:
            continue
        r1 = tbl.rows[1]
        gw = {
            "count": int(r1.cells[0].text.strip()) if r1.cells[0].text.strip().isdigit() else 0,
            "depth_min": _bf(r1.cells[1].text),
            "depth_max": _bf(r1.cells[2].text),
            "depth_avg": _bf(r1.cells[3].text),
            "elv_min": _bf(r1.cells[4].text),
            "elv_max": _bf(r1.cells[5].text),
            "elv_avg": _bf(r1.cells[6].text),
        }
        break
    result["groundwater"] = gw

    # ================================================================
    # 3) 地层描述
    # ================================================================
    layers: "collections.OrderedDict[str, Dict]" = __import__("collections").OrderedDict()
    layer_order: List[str] = []

    if layer_start_idx >= 0:
        end_i = section4_idx if section4_idx > 0 else len(paragraphs)
        for i in range(layer_start_idx + 1, end_i):
            _, text = paragraphs[i]
            if "物理力学指标统计表" in text:
                continue
            m = re.match(r"(\d+(?:-\d+)?)层(\S+?)[：:]", text)
            if not m:
                continue

            layer_id = m.group(1)
            layer_name = m.group(2)
            # 去除名称尾部的逗号
            layer_name = layer_name.rstrip("，,")

            d: Dict[str, Any] = {"id": layer_id, "name": layer_name, "full_desc": text}

            # 厚度
            tm = re.search(
                r"厚度[:：]\s*([\d.]+)\s*[～~]\s*([\d.]+)\s*m\s*[,，]?\s*平均\s*([\d.]+)\s*m",
                text,
            )
            if tm:
                d["thickness"] = {
                    "min": float(tm.group(1)),
                    "max": float(tm.group(2)),
                    "avg": float(tm.group(3)),
                }

            # 层底标高
            em = re.search(
                r"层底标高[:：]\s*(-?[\d.]+)\s*[～~]\s*(-?[\d.]+)\s*m\s*[,，]?\s*平均\s*(-?[\d.]+)\s*m",
                text,
            )
            if em:
                d["bottom_elevation"] = {
                    "min": float(em.group(1)),
                    "max": float(em.group(2)),
                    "avg": float(em.group(3)),
                }

            # 层底埋深
            dm = re.search(
                r"层底埋深[:：]\s*([\d.]+)\s*[～~]\s*([\d.]+)\s*m\s*[,，]?\s*平均\s*([\d.]+)\s*m",
                text,
            )
            if dm:
                d["burial_depth"] = {
                    "min": float(dm.group(1)),
                    "max": float(dm.group(2)),
                    "avg": float(dm.group(3)),
                }

            # 是否穿透
            if "未穿透" in text:
                d["penetrated"] = False

            layers[layer_id] = d
            layer_order.append(layer_id)

    result["layers"] = layers
    result["layer_order"] = layer_order
    result["layer_count"] = len(layer_order)

    # ================================================================
    # 4) 物理力学指标表  (12 行 × 8 列, 紧跟 "物理力学指标统计表" 段落)
    # ================================================================
    _PARAM_KEYS = [
        "W", "gamma", "e", "WL", "WP", "IP", "IL", "C", "phi", "a1_2", "Es",
    ]

    physmech: Dict[str, Dict] = {}
    used_tbl_set: set = set()

    # 建立 "物理力学指标统计表" 段落在 body 中的位置
    stat_positions: List[int] = []
    for i, (_, text) in enumerate(paragraphs):
        if "物理力学指标统计表" in text:
            stat_positions.append(i)

    for si, para_i in enumerate(stat_positions):
        if si >= len(layer_order):
            break
        lid = layer_order[si]

        # 找到该段落之后的第一个未使用 phys-mech 表格
        body_idx = paragraphs[para_i][0]
        found_tbl = None
        found_id = None
        for ti, (tbl_body_idx, tbl) in enumerate(tables):
            if ti in used_tbl_set:
                continue
            if tbl_body_idx <= body_idx:
                continue
            if len(tbl.rows) < 12 or len(tbl.columns) < 8:
                continue
            # 确认是物理力学表 (第一行 W(%))
            r1c0 = _cell(tbl, 1, 0)
            if "W" not in r1c0:
                continue
            found_tbl = tbl
            found_id = ti
            break

        if found_tbl is None:
            continue

        used_tbl_set.add(found_id)
        props: Dict[str, Dict] = {}
        for row_i in range(1, min(12, len(found_tbl.rows))):
            param = _cell(found_tbl, row_i, 0)
            key_idx = row_i - 1
            key = _PARAM_KEYS[key_idx] if key_idx < len(_PARAM_KEYS) else param
            props[key] = {
                "name": param,
                "min": _bf(_cell(found_tbl, row_i, 1)),
                "max": _bf(_cell(found_tbl, row_i, 2)),
                "avg": _bf(_cell(found_tbl, row_i, 3)),
                "n": _bf(_cell(found_tbl, row_i, 4)),
                "stddev": _bf(_cell(found_tbl, row_i, 5)),
                "cv": _bf(_cell(found_tbl, row_i, 6)),
                "xk": _bf(_cell(found_tbl, row_i, 7)),
            }
        physmech[lid] = props

    result["physmech"] = physmech

    # ================================================================
    # 5) 承载力建议值表  (4 列: 层号, 名称, fak, Es)
    # ================================================================
    bearing: Dict[str, Dict] = {}
    if bearing_heading >= 0:
        bh_body_idx = paragraphs[bearing_heading][0]
        for _, tbl in tables:
            # 找 "各层建议值" 之后的第一个 4 列表格
            # (简化: 直接找 4 列且含 fak 的表)
            if len(tbl.columns) != 4:
                continue
            h = _cell(tbl, 0, 2).lower()
            if "fak" not in h and "承载力" not in _cell(tbl, 0, 2):
                continue
            for ri in range(1, len(tbl.rows)):
                lid = _cell(tbl, ri, 0)
                lname = _cell(tbl, ri, 1)
                fak_v = _bf(_cell(tbl, ri, 2))
                es_v = _bf(_cell(tbl, ri, 3))
                if lid:
                    bearing[lid] = {
                        "name": lname,
                        "fak": fak_v,
                        "es": es_v,
                    }
            break  # 只取第一个匹配的表
    result["bearing"] = bearing

    # ---- 项目基本信息 ----
    for _, text in paragraphs[:30]:
        if text.startswith("工程编号"):
            result["project_id"] = text.replace("工程编号", "").strip(":： ")
            break

    logger.info(
        f"  简报提取: 高程{elv}, 地下水{gw.get('count', 0)}个, "
        f"{len(layer_order)}层, 物理力学{len(physmech)}层有数据"
    )
    return result


# ============================================================
# 建筑-钻孔映射 (DXF 总平面图空间匹配)
# ============================================================

def _edge_length(p1: Tuple, p2: Tuple) -> float:
    """计算两点间距离"""
    return math.sqrt((p1[0] - p2[0]) ** 2 + (p1[1] - p2[1]) ** 2)


def _point_in_rect(px: float, py: float,
                   xmin: float, xmax: float,
                   ymin: float, ymax: float,
                   margin: float = 5) -> bool:
    """判断点是否在矩形范围内 (含 margin 容差)"""
    return xmin - margin <= px <= xmax + margin and \
           ymin - margin <= py <= ymax + margin


def _extract_outlines(msp: Any, layer_name: str,
                      x_min: float = 0, min_pts: int = 4,
                      div: float = 1.0) -> List[Dict]:
    """从 DXF 模型空间提取建筑物闭合轮廓

    Adapted from cad-building-extract skill:
    - 使用 LWPOLYLINE 闭合多边形
    - 实际边长 (非 bounding box), 支持旋转建筑
    - Shoelace 面积计算
    - div: 坐标单位转换 (1=米, 1000=毫米→米)

    Returns: list of outline dicts with cx, cy, length, width, area, bbox, pts
    """
    outlines = []
    for e in msp:
        if e.dxftype() != 'LWPOLYLINE':
            continue
        if getattr(e.dxf, 'layer', '') != layer_name:
            continue
        try:
            pts = list(e.get_points(format='xy'))
        except Exception:
            continue
        n = len(pts)
        if n < min_pts:
            continue

        # 中心点
        cx = sum(p[0] for p in pts) / n
        cy = sum(p[1] for p in pts) / n
        if cx < x_min:
            continue

        xs = [p[0] for p in pts]
        ys = [p[1] for p in pts]

        # Shoelace 面积
        area = 0
        for j in range(n):
            j1 = (j + 1) % n
            area += pts[j][0] * pts[j1][1] - pts[j1][0] * pts[j1][1]
        area = abs(area) / 2 / (div * div)

        # 实际边长 (4点矩形取对边平均, 多边形用 bbox)
        if n == 4:
            side1 = (_edge_length(pts[0], pts[1]) + _edge_length(pts[2], pts[3])) / 2 / div
            side2 = (_edge_length(pts[1], pts[2]) + _edge_length(pts[3], pts[0])) / 2 / div
            length = max(side1, side2)
            width = min(side1, side2)
        else:
            length = max(max(xs) - min(xs), max(ys) - min(ys)) / div
            width = min(max(xs) - min(xs), max(ys) - min(ys)) / div

        outlines.append({
            'cx': cx / div, 'cy': cy / div,
            'length': length, 'width': width,
            'area': area, 'npts': n,
            'xmin': min(xs) / div, 'xmax': max(xs) / div,
            'ymin': min(ys) / div, 'ymax': max(ys) / div,
            'pts': [(p[0] / div, p[1] / div) for p in pts],
            # 保留原始坐标 (未缩放), 用于与钻孔坐标匹配
            'raw_xmin': min(xs), 'raw_xmax': max(xs),
            'raw_ymin': min(ys), 'raw_ymax': max(ys),
            'raw_cx': cx, 'raw_cy': cy,
        })
    return outlines


def _extract_building_texts(msp: Any, x_min: float = 0) -> List[Dict]:
    """从 DXF 提取建筑编号标注 (TEXT + MTEXT)

    建筑编号识别规则: 以数字+#开头, 如 1#, 2#, A-1# 等
    """
    texts = []
    for e in msp:
        etype = e.dxftype()
        if etype not in ('TEXT', 'MTEXT'):
            continue
        try:
            if etype == 'TEXT':
                txt = e.dxf.text.strip()
                x, y = round(e.dxf.insert.x, 2), round(e.dxf.insert.y, 2)
            else:
                txt = e.text.strip()
                # MTEXT: 使用插入点
                x = round(e.dxf.insert.x, 2)
                y = round(e.dxf.insert.y, 2)
        except Exception:
            continue

        if x < x_min or not txt:
            continue

        # 建筑编号: 数字+# 或 含字母前缀的编号
        if re.match(r'^[A-Za-z]*-?\d+#', txt):
            texts.append({'text': txt, 'x': x, 'y': y})
    return texts


def _auto_detect_dxf_layers(msp: Any) -> Dict[str, str]:
    """自动检测 DXF 中建筑标注/轮廓图层

    Returns: {"text_layer": str, "outline_layer": str}
    """
    layer_info: Dict[str, Dict] = {}
    for e in msp:
        layer = getattr(e.dxf, 'layer', '')
        if not layer:
            continue
        if layer not in layer_info:
            layer_info[layer] = {'count': 0, 'types': set()}
        layer_info[layer]['count'] += 1
        layer_info[layer]['types'].add(e.dxftype())

    text_layer = ""
    outline_layer = ""

    # 文字图层: 含建筑编号标注的图层 (优先含关键词)
    text_layers = [
        (n, i) for n, i in layer_info.items()
        if 'TEXT' in i['types'] or 'MTEXT' in i['types'] and i['count'] > 10
    ]
    if text_layers:
        preferred = [
            x for x in text_layers
            if any(kw in x[0] for kw in ['住宅', '建筑', '标注', '指标'])
        ]
        text_layer = (max(preferred, key=lambda x: x[1]['count'])
                      if preferred
                      else max(text_layers, key=lambda x: x[1]['count']))[0]

    # 轮廓图层: 含 LWPOLYLINE 的图层 (优先含关键词)
    outline_layers = [
        (n, i) for n, i in layer_info.items()
        if 'LWPOLYLINE' in i['types']
    ]
    if outline_layers:
        preferred = [
            x for x in outline_layers
            if any(kw in x[0] for kw in ['地上', '轮廓', '建筑', '住宅'])
        ]
        outline_layer = (max(preferred, key=lambda x: x[1]['count'])
                         if preferred
                         else max(outline_layers, key=lambda x: x[1]['count']))[0]

    return {"text_layer": text_layer, "outline_layer": outline_layer}


def _detect_coord_div(msp: Any) -> float:
    """检测坐标单位转换因子

    规则: 建筑编号标注 X 坐标中位数 <1M → 米 (div=1), >1M → 毫米 (div=1000)
    """
    building_xs = []
    for e in msp:
        try:
            if e.dxftype() == 'TEXT' and '#' in (e.dxf.text or ''):
                building_xs.append(e.dxf.insert.x)
            elif e.dxftype() == 'MTEXT' and '#' in (e.text or ''):
                building_xs.append(e.dxf.insert.x)
        except Exception:
            continue

    if not building_xs:
        return 1.0

    median_x = sorted(building_xs)[len(building_xs) // 2]
    return 1000.0 if median_x > 1_000_000 else 1.0


def _match_outline(bx: float, by: float, outlines: List[Dict],
                   used: set, max_dist: float = 50,
                   min_area: float = 60) -> Optional[Dict]:
    """为建筑编号标注匹配最佳轮廓

    优先级: 标注点在轮廓内 > 面积最大 > 距离最近
    """
    candidates = []
    for idx, o in enumerate(outlines):
        if idx in used:
            continue
        if o['area'] < min_area:
            continue
        d = math.sqrt((bx - o['raw_cx']) ** 2 + (by - o['raw_cy']) ** 2)
        inside = _point_in_rect(bx, by,
                                o['raw_xmin'], o['raw_xmax'],
                                o['raw_ymin'], o['raw_ymax'], margin=10)
        if d < max_dist:
            candidates.append((idx, d, o, inside))

    if not candidates:
        return None

    # 优先: inside > larger area > closer distance
    candidates.sort(key=lambda x: (not x[3], -x[2]['area'], x[1]))
    best_idx = candidates[0][0]
    used.add(best_idx)
    return candidates[0][2]


def extract_building_borehole_mapping(
    dxf_path: str,
    borehole_coords: Dict[str, Tuple[float, float]],
    building_info_path: str = "",
    text_layer: str = "",
    outline_layer: str = "",
) -> Dict[str, Any]:
    """从 DXF 总平面图提取建筑-钻孔空间映射

    工作流程 (adapted from cad-building-extract skill):
    1. 读取 DXF, 自动检测图层和坐标单位
    2. 提取建筑编号标注和轮廓
    3. 匹配编号→轮廓 (标注点在轮廓内优先)
    4. 加载建筑信息 JSON (补充层数/尺寸/标高)
    5. 将钻孔坐标与建筑轮廓矩形做空间匹配

    Parameters
    ----------
    dxf_path : str
        DXF 总平面图文件路径
    borehole_coords : dict
        {bh_id: (X, Y)} 钻孔坐标, 来自 DK 文件
    building_info_path : str, optional
        建筑信息 JSON 文件路径 (含层数/尺寸/标高等)
    text_layer : str, optional
        手动指定文字标注图层名
    outline_layer : str, optional
        手动指定建筑轮廓图层名

    Returns
    -------
    dict
        {
            "available": True,
            "buildings": {
                "1#": {
                    "name": "1#",
                    "json_name": "A-1#",
                    "floors": 17,
                    "elev": 25.3,
                    "dim": "62.93×50.80",
                    "outline": {"xmin": ..., "xmax": ..., "ymin": ..., "ymax": ...},
                    "boreholes": [{"bh_id": "22", "distance": 8.4}, ...],
                    "bh_count": 14,
                },
                ...
            },
            "unmatched_boreholes": ["bh1", "bh2", ...],
            "total_buildings": 21,
            "total_matched_boreholes": 162,
        }
    """
    if not dxf_path or not os.path.isfile(dxf_path):
        return {"available": False}

    # 导入 ezdxf
    try:
        import ezdxf
    except ImportError:
        logger.warning("  DXF 提取需要 ezdxf 库: pip install ezdxf")
        return {"available": False}

    # 读取 DXF (多编码尝试)
    doc = None
    for enc in ('gbk', 'utf-8', 'gb2312', 'gb18030'):
        try:
            doc = ezdxf.readfile(dxf_path, encoding=enc)
            break
        except Exception:
            continue

    if doc is None:
        logger.warning(f"  无法读取 DXF 文件: {dxf_path}")
        return {"available": False}

    msp = doc.modelspace()

    # 自动检测图层和坐标单位
    detected = _auto_detect_dxf_layers(msp)
    if not text_layer:
        text_layer = detected.get("text_layer", "")
    if not outline_layer:
        outline_layer = detected.get("outline_layer", "")

    div = _detect_coord_div(msp)

    logger.info(f"  DXF 图层: text={text_layer}, outline={outline_layer}, div={div}")

    # 提取建筑编号标注
    x_min_filter = 0
    bldg_texts = _extract_building_texts(msp, x_min_filter)
    logger.info(f"  建筑编号标注: {len(bldg_texts)} 个")

    # 提取建筑轮廓
    outlines = _extract_outlines(msp, outline_layer, x_min_filter, min_pts=4, div=div) \
               if outline_layer else []
    logger.info(f"  建筑轮廓: {len(outlines)} 个")

    # 加载建筑信息 JSON (补充层数/尺寸等)
    building_info: Dict[str, Dict] = {}
    if building_info_path and os.path.isfile(building_info_path):
        try:
            with open(building_info_path, 'r', encoding='utf-8') as f:
                binfo_raw = json.load(f)
            # buildings 列表 → dict keyed by name
            for b in binfo_raw.get("buildings", []):
                building_info[b.get("name", "")] = b
            logger.info(f"  建筑信息 JSON: {len(building_info)} 栋")
        except Exception as e:
            logger.warning(f"  建筑信息 JSON 读取失败: {e}")

    # 匹配编号 → 轮廓
    used_outlines: set = set()
    buildings: Dict[str, Dict] = {}

    # 按编号排序
    bldg_texts.sort(
        key=lambda b: int(re.match(r'(\d+)', b['text']).group(1))
        if re.match(r'(\d+)', b['text']) else 999
    )

    for bn in bldg_texts:
        bname = bn['text']
        # 去掉字母前缀 (如 A-1# → 1#) 用于内部 key
        inner_key = re.sub(r'^[A-Za-z]-', '', bname)

        # 匹配轮廓
        outline = _match_outline(bn['x'], bn['y'], outlines, used_outlines,
                                 max_dist=100, min_area=30)

        # 构建建筑数据
        bdata: Dict[str, Any] = {
            "name": bname,
            "label_x": bn['x'],
            "label_y": bn['y'],
        }

        # 轮廓信息
        if outline:
            bdata["outline"] = {
                "xmin": outline['raw_xmin'],
                "xmax": outline['raw_xmax'],
                "ymin": outline['raw_ymin'],
                "ymax": outline['raw_ymax'],
                "cx": outline['raw_cx'],
                "cy": outline['raw_cy'],
                "length": outline['length'],
                "width": outline['width'],
                "area": outline['area'],
            }
        else:
            bdata["outline"] = None

        # 补充建筑信息 JSON 数据
        # JSON name 可能带字母前缀 (A-1#), 而 DXF 标注可能是 1#
        json_match = building_info.get(bname) or building_info.get(inner_key) or \
                     building_info.get(f"A-{inner_key}") or building_info.get(f"B-{inner_key}")
        if json_match:
            bdata["json_name"] = json_match.get("name", "")
            bdata["floors"] = json_match.get("floors", 0)
            bdata["elev"] = json_match.get("elev", 0)
            bdata["dim"] = json_match.get("dim", "")
        elif outline:
            bdata["floors"] = 0
            bdata["elev"] = 0
            bdata["dim"] = f"{outline['length']:.2f}×{outline['width']:.2f}"

        # 空间匹配钻孔
        #   策略: 使用建筑标注坐标 + 建筑信息尺寸构建虚拟矩形
        #         DXF 轮廓仅在面积合理时作为补充参考
        bh_list: List[Dict] = []
        if borehole_coords:
            # 确定匹配范围:
            #   优先用建筑信息尺寸建虚拟矩形 (标注坐标为中心)
            #   DXF 轮廓仅在面积 >= 建筑信息面积 50% 时使用
            ocx = bn['x']
            ocy = bn['y']

            # 从 dim 字符串中解析长宽 (来自建筑信息 JSON)
            half_w = 0
            half_h = 0
            dim_str = bdata.get("dim", "")
            if dim_str:
                try:
                    parts = dim_str.replace("x", "×").split("×")
                    half_w = float(parts[0]) / 2
                    half_h = float(parts[1]) / 2
                except Exception:
                    pass

            # 评估 DXF 轮廓是否可用 (面积应 >= 建筑信息面积 50%)
            outline = bdata.get("outline")
            use_outline = False
            if outline and half_w > 0 and half_h > 0:
                # 计算建筑信息面积
                expected_area = (half_w * 2) * (half_h * 2)
                outline_area = outline.get("area", 0)
                # DXF 轮廓面积 (需考虑 div 转换)
                # 如果 outline 面积太小 (<50%), 说明轮廓不是建筑主体
                if outline_area >= expected_area * 0.5:
                    use_outline = True
                    ocx = outline["cx"]
                    ocy = outline["cy"]

            # 默认尺寸 (无建筑信息时)
            if not half_w:
                half_w = 20
            if not half_h:
                half_h = 20

            # 构建匹配矩形 (DXF 原始坐标)
            oxmin = ocx - half_w * div
            oxmax = ocx + half_w * div
            oymin = ocy - half_h * div
            oymax = ocy + half_h * div
            outline_max_dim = max(half_w, half_h) * 2

            # 钻孔匹配: 在矩形内 (含 margin)
            # margin = 轮廓半宽 * 15% + 3m (确保边角钻孔也被捕获)
            margin = outline_max_dim * 0.15 + 3

            for bh_id, (bx, by) in borehole_coords.items():
                d = math.sqrt((bx - ocx) ** 2 + (by - ocy) ** 2)
                if _point_in_rect(bx, by, oxmin, oxmax, oymin, oymax, margin=margin):
                    bh_list.append({"bh_id": bh_id, "distance": round(d, 1)})

            # 按距离排序
            bh_list.sort(key=lambda x: x["distance"])

        bdata["boreholes"] = bh_list
        bdata["bh_count"] = len(bh_list)
        buildings[inner_key] = bdata

    # 统计未匹配钻孔
    matched_ids: set = set()
    for bdata in buildings.values():
        for bh in bdata["boreholes"]:
            matched_ids.add(bh["bh_id"])

    unmatched = [bh_id for bh_id in borehole_coords if bh_id not in matched_ids]

    result = {
        "available": True,
        "buildings": buildings,
        "unmatched_boreholes": unmatched,
        "total_buildings": len(buildings),
        "total_matched_boreholes": len(matched_ids),
        "dxf_path": dxf_path,
    }

    logger.info(
        f"  建筑映射: {len(buildings)} 栋建筑, "
        f"{len(matched_ids)} 个钻孔已分配, "
        f"{len(unmatched)} 个未分配"
    )
    return result


def extract_borehole_coords_from_dxf(
    dxf_path: str,
    hn_coords: Dict[str, Tuple[float, float]] = None,
) -> Dict[str, Tuple[float, float]]:
    """从 DXF 平面图直接提取钻孔坐标

    从 DXF 的勘探点图层提取 INSERT 块的插入点坐标作为钻孔位置。
    INSERT 块通常表示钻孔符号（圆环、十字等），其插入点即为钻孔中心。

    如果提供 hn_coords (华宁 DB 坐标)，则按坐标距离匹配编号；
    否则用 DXF 上的 TEXT 标注作为钻孔编号。

    Parameters
    ----------
    dxf_path : str
        DXF 平面图文件路径
    hn_coords : dict, optional
        华宁 DB 的钻孔坐标 {bh_id: (X, Y)}, 用于编号匹配

    Returns
    -------
    dict
        {bh_id: (X, Y)} 钻孔编号→坐标映射
    """
    if not dxf_path or not os.path.isfile(dxf_path):
        return {}

    try:
        import ezdxf
    except ImportError:
        logger.warning("  DXF 提取需要 ezdxf 库: pip install ezdxf")
        return {}

    # 读取 DXF
    doc = None
    for enc in ('gbk', 'utf-8', 'gb2312', 'gb18030'):
        try:
            doc = ezdxf.readfile(dxf_path, encoding=enc)
            break
        except Exception:
            continue

    if doc is None:
        logger.warning(f"  无法读取 DXF 文件: {dxf_path}")
        return {}

    msp = doc.modelspace()

    # Step 1: 找勘探点图层 (含最多纯数字 TEXT 标注的图层)
    layer_digits: Dict[str, int] = {}
    for e in msp:
        if e.dxftype() != 'TEXT':
            continue
        try:
            txt = e.dxf.text.strip()
            layer = getattr(e.dxf, 'layer', '')
            if txt.isdigit() and 1 <= int(txt) <= 999:
                layer_digits[layer] = layer_digits.get(layer, 0) + 1
        except Exception:
            continue

    if not layer_digits:
        return {}

    bh_layer = max(layer_digits.items(), key=lambda x: x[1])[0]

    # Step 2: 提取该图层上的 INSERT 块 (钻孔符号)
    inserts: List[Tuple[float, float]] = []
    for e in msp:
        if e.dxftype() != 'INSERT':
            continue
        if getattr(e.dxf, 'layer', '') != bh_layer:
            continue
        try:
            x = round(e.dxf.insert.x, 2)
            y = round(e.dxf.insert.y, 2)
            inserts.append((x, y))
        except Exception:
            continue

    # Step 3: 提取该图层上的 TEXT 标注 (钻孔编号)
    texts: List[Tuple[str, float, float]] = []
    for e in msp:
        if e.dxftype() != 'TEXT':
            continue
        if getattr(e.dxf, 'layer', '') != bh_layer:
            continue
        try:
            txt = e.dxf.text.strip()
            if not txt.isdigit():
                continue
            x = round(e.dxf.insert.x, 2)
            y = round(e.dxf.insert.y, 2)
            texts.append((txt, x, y))
        except Exception:
            continue

    # Step 4: 确定钻孔坐标和编号
    coords: Dict[str, Tuple[float, float]] = {}

    if hn_coords and inserts:
        # 有华宁 DB 坐标: 按坐标距离匹配 INSERT → 华宁编号
        used_hn: set = set()
        for ix, iy in inserts:
            min_dist = float('inf')
            best_id = None
            for hn_id, (hx, hy) in hn_coords.items():
                if hn_id in used_hn:
                    continue
                d = math.sqrt((ix - hx) ** 2 + (iy - hy) ** 2)
                if d < min_dist:
                    min_dist = d
                    best_id = hn_id
            if best_id and min_dist < 5.0:  # 5m 容差
                coords[best_id] = (ix, iy)
                used_hn.add(best_id)
    elif inserts and texts:
        # 无华宁 DB: 用 TEXT 标注编号匹配最近的 INSERT
        used_ins: set = set()
        for txt, tx, ty in texts:
            min_dist = float('inf')
            best_idx = -1
            for i, (ix, iy) in enumerate(inserts):
                if i in used_ins:
                    continue
                d = math.sqrt((tx - ix) ** 2 + (ty - iy) ** 2)
                if d < min_dist:
                    min_dist = d
                    best_idx = i
            if best_idx >= 0 and min_dist < 30:  # 30m 容差 (文字可能离得远)
                coords[txt] = inserts[best_idx]
                used_ins.add(best_idx)
    elif inserts:
        # 只有 INSERT 没有 TEXT: 自动编号 1, 2, 3...
        for i, (ix, iy) in enumerate(inserts, 1):
            coords[str(i)] = (ix, iy)

    logger.info(
        f"  DXF 钻孔: 图层 '{bh_layer}', "
        f"{len(inserts)} 个 INSERT, {len(texts)} 个 TEXT, "
        f"提取 {len(coords)} 个坐标"
    )
    return coords



# ============================================================
# 报告填充器
# ============================================================

class ReportFiller:
    """将勘察数据填充到 .docx 报告模板"""

    def __init__(
        self,
        template_path: str,
        output_path: str,
        data: Dict[str, Any],
        layer_names: Dict[str, str],
        layer_ids: List[str],
        config: ProjectConfig,
    ):
        self.template = template_path
        self.output = output_path
        self.data = data
        self.layer_names = layer_names
        self.layer_ids = layer_ids
        self.config = config
        self.ti = config.table_indices  # 表格索引快捷引用
        self.doc = Document(template_path)

    def save(self) -> None:
        """保存生成的报告"""
        self.doc.save(self.output)

    def fill_all(self) -> None:
        """执行全部填充步骤"""
        logger.info("")
        logger.info("=" * 60)
        logger.info("  开始填充报告（只改文字，不动格式）")
        logger.info("=" * 60)

        self._global_replace()
        self._fill_project_overview()       # 1. 第一章: 拟建工程概况
        self._fill_survey_purpose()         # 2. 第二章: 勘察目的条件过滤
        self._fill_buildings_table()        # 3. 建筑物特征表
        self._fill_workload()
        self._fill_water_level()
        self._fill_layer_descriptions()
        self._fill_phys_spt_tables()
        self._fill_bearing_capacity()
        self._fill_water_salt_tables()
        self._fill_corrosion_eval()
        self._fill_liquefaction()
        self._fill_foundation_tables()
        self._fill_conclusion()
        self._fill_site_conditions()
        self._fill_wave_velocity_table()    # 波速估算 & 场地类别判定
        self._fill_site_evaluation()
        self._fill_foundation_evaluation()
        self._fill_analysis_evaluation()    # 16. 第六章: 分析评价各子节
        self._fill_standards()
        self._apply_date_replacements()

        logger.info(f"\n  保存: {self.output}")

    # ---- 全局文本替换 ----

    def _global_replace(self) -> None:
        """按配置执行全局文本替换"""
        replacements = self.config.replacements
        if not replacements:
            return

        logger.info("  全局文本替换...")
        for old, new in replacements:
            for p in self.doc.paragraphs:
                replace_in_para(p, old, new)
            for t in self.doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for cp in cell.paragraphs:
                            replace_in_para(cp, old, new)

    # ---- 第一章: 拟建工程概况 ----

    def _fill_project_overview(self) -> None:
        """填充拟建工程概况章节（第一章）

        配置字段:
            commission_text:    委托段落全文（支持 {client}/{project_name}/{survey_stage} 占位符）
            site_location:      场地位置描述段落
            building_desc:      建筑物描述段落（含面积、层数等）
            importance_level:   重要性等级（一级/二级/三级）
            site_complexity:    场地复杂程度等级（一级/二级/三级）
            foundation_grade:   地基等级（一级/二级/三级）
            survey_grade:       岩土工程勘察等级（甲级/乙级/丙级）
            survey_grade_text:  勘察等级完整描述（覆盖自动拼装）
        """
        overview = self.config.get_project_overview()
        if not overview:
            return

        # 准备占位符
        fmt_vars: Dict[str, str] = {
            "client": overview.get("client", ""),
            "project_name": overview.get("project_name", self.config.project_name),
            "survey_stage": overview.get("survey_stage", "详细勘察"),
        }

        # 自动拼装勘察等级文字
        grade_parts = []
        for key, label in [
            ("importance_level", "重要性等级"),
            ("site_complexity", "场地复杂程度等级"),
            ("foundation_grade", "地基等级"),
        ]:
            val = overview.get(key, "")
            if val:
                grade_parts.append(f"{label}{val}")
        survey_grade = overview.get("survey_grade", "")
        if survey_grade:
            grade_parts.append(f"综合确定岩土工程勘察等级为{survey_grade}")
        auto_grade_text = "，".join(grade_parts) + "。" if grade_parts else ""
        grade_text = overview.get("survey_grade_text", auto_grade_text)

        # 在模板中定位并替换
        commission_text = overview.get("commission_text", "")
        site_location = overview.get("site_location", "")
        building_desc = overview.get("building_desc", "")

        replaced_count = 0
        for p in self.doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue

            # 1. 委托段落: 匹配 "受XX的委托" 或 "承担了" 开头
            if commission_text and ("的委托" in txt or "承担了" in txt):
                try:
                    full_text = commission_text.format(**fmt_vars)
                except KeyError:
                    full_text = commission_text
                set_para_text(p, full_text)
                replaced_count += 1
                continue

            # 2. 场地位置: 匹配 "勘察场地位于" 开头
            if site_location and txt.startswith("勘察场地位于"):
                set_para_text(p, site_location)
                replaced_count += 1
                continue

            # 3. 建筑物描述: 匹配 "本次勘察建筑物" 或 "总建筑面积"
            if building_desc and (
                "本次勘察建筑物" in txt or "总建筑面积" in txt
            ):
                set_para_text(p, building_desc)
                replaced_count += 1
                continue

            # 4. 勘察等级: 匹配 "岩土工程勘察等级"
            if grade_text and "岩土工程勘察等级" in txt:
                set_para_text(p, grade_text)
                replaced_count += 1
                continue

            # 5. 标准冻结深度 (第四章): 匹配 "标准冻结深度"
            frozen_depth = overview.get("frozen_depth", "")
            if frozen_depth and "标准冻结深度" in txt:
                set_para_text(p, f"场地土标准冻结深度{frozen_depth}。")
                replaced_count += 1
                continue

        if replaced_count:
            logger.info(f"  工程概况: 替换 {replaced_count} 段")
        else:
            logger.debug("  工程概况: 未找到匹配段落")

    # ---- 勘察目的 / 任务要求条件过滤 (第二章) ----

    def _fill_survey_purpose(self) -> None:
        """根据条件过滤勘察目的/任务要求条目（第二章）

        当无地下室 (has_basement=False) 时，删除含"基坑开挖"的条目并重新编号。
        可通过配置 project_overview.has_basement 显式控制。
        """
        overview = self.config.get_project_overview()
        if not overview:
            return

        # has_basement 判断: 优先取配置显式值, 否则从建筑物数据自动检测
        conditions = self._evaluate_standard_conditions()
        has_basement = overview.get("has_basement", conditions.get("has_basement", False))

        if has_basement:
            return  # 有地下室/基坑, 全部保留

        # 条件关键词: 包含此关键词的条目在条件不满足时删除
        cond_keyword = "基坑开挖"

        # 收集编号条目: [(para_index, item_number)]
        numbered_items: List[Tuple[int, int]] = []
        removed_indices: List[int] = []
        num_idx = 1
        for i, p in enumerate(self.doc.paragraphs):
            txt = p.text.strip()
            if not txt:
                continue
            prefix = f"{num_idx}."
            if txt.startswith(prefix):
                if cond_keyword in txt:
                    removed_indices.append(i)
                    # 仍加入列表以便后续重编号 (最终会被清空)
                    numbered_items.append((i, num_idx))
                else:
                    numbered_items.append((i, num_idx))
                num_idx += 1

        if not removed_indices:
            return

        # 清空需要删除的段落
        for idx in removed_indices:
            set_para_text(self.doc.paragraphs[idx], "")

        # 重新编号: 剩余条目按顺序 1,2,3...
        remaining = [(idx, n) for idx, n in numbered_items if idx not in removed_indices]
        for seq, (idx, old_num) in enumerate(remaining, 1):
            p = self.doc.paragraphs[idx]
            old_prefix = f"{old_num}."
            new_prefix = f"{seq}."
            new_text = new_prefix + p.text.strip()[len(old_prefix):]
            set_para_text(p, new_text)

        logger.info(f"  勘察目的: 删除基坑相关条目 {len(removed_indices)} 条, 重编号 {len(remaining)} 条")

    # ---- 建筑物特征表 ----

    def _fill_buildings_table(self) -> None:
        buildings = self.data.get("buildings", [])
        if not buildings:
            return

        logger.info(f"  建筑物特征表 ({len(buildings)} 栋)...")
        idx = self.ti["buildings"]
        t = self.doc.tables[idx] if len(self.doc.tables) > idx else None
        if not t:
            return

        building_keys = ["name", "floors", "height", "size", "span", "indoor_elv", "structure"]
        for bi, b in enumerate(buildings):
            if bi + 2 >= len(t.rows):
                break
            for ci, key in enumerate(building_keys):
                set_cell(t, bi + 2, ci, b.get(key, ""))

    # ---- 工作量表 + 段落 ----

    def _fill_workload(self) -> None:
        logger.info("  工作量表 + 段落...")
        bh_info = self.data["borehole_info"]

        # 工作量统计表
        idx = self.ti["workload"]
        t = self.doc.tables[idx] if len(self.doc.tables) > idx else None
        if t:
            set_cell(t, 1, 3, str(bh_info["total"]))
            set_cell(t, 2, 3, f"{bh_info['total_depth']:.0f}/{bh_info['total']}")
            set_cell(t, 3, 3, str(bh_info["total"]))
            set_cell(t, 4, 3, str(bh_info["undisturbed"]))
            set_cell(t, 5, 3, str(bh_info["disturbed"]))
            set_cell(t, 6, 3, str(bh_info["rock"]))
            set_cell(t, 7, 3, str(bh_info["spt_total"]))
            set_cell(t, 8, 3, str(bh_info["n63_total"]))
            if len(t.rows) > 9:
                set_cell(t, 9, 3, str(len(self.data.get("water_samples", []))))
            if len(t.rows) > 10:
                set_cell(t, 10, 3, str(len(self.data.get("salt_samples", []))))
            if len(t.rows) > 11:
                set_cell(t, 11, 3, str(bh_info.get("bosk", 0)))

        # 工作量段落
        for p in self.doc.paragraphs:
            txt = p.text.strip()
            if txt.startswith("布孔") and "钻孔" in txt:
                set_para_text(p, (
                    f"布孔：沿拟建物周边及角点并结合网度布设钻孔{bh_info['total']}个，"
                    f"间距16.0～20.00m；其中控制性钻孔{bh_info['ctrl']}个、"
                    f"一般性钻孔{bh_info['general']}个。"
                ))
            elif "钻孔类型" in txt and "取土孔" in txt:
                set_para_text(p, (
                    f"钻孔类型：取土孔{bh_info['qutu']}个、标贯孔{bh_info['biaoguan']}个，"
                    f"鉴别孔{bh_info['yiban']}个，波速孔{bh_info.get('bosk', 0)}个。"
                ))
            elif "钻孔孔深" in txt and "强风化" in txt:
                set_para_text(p, (
                    "钻孔孔深：控制性钻孔进入强风化岩面以下7-9m、"
                    "一般性钻孔进入强风化岩面以下不小于6m。"
                ))
            elif "水样及易溶盐" in txt:
                n_water = len(self.data.get("water_samples", []))
                n_salt = len(self.data.get("salt_samples", []))
                set_para_text(p, f"水样及易溶盐土样：地下水水样{n_water}件、易溶盐土样{n_salt}件。")
            elif "本次勘察实际完成" in txt:
                set_para_text(p, (
                    f"本次勘察实际完成钻孔{bh_info['total']}个，"
                    f"其中控制孔{bh_info['ctrl']}个，一般孔{bh_info['general']}个。"
                    f"其中取土孔{bh_info['qutu']}个、标贯孔{bh_info['biaoguan']}个，"
                    f"鉴别孔{bh_info['yiban']}个，波速孔{bh_info.get('bosk', 0)}个，"
                    f"勘察工作量统计表见表3-1。"
                ))

        # --- 勘察方法条件过滤 (第三章(一)) ---
        # N63.5 动力触探: 无数据时删除
        has_n63 = bh_info.get("n63_total", 0) > 0

        # 波速测试: 无高层建筑 (>24m) 时删除
        buildings = self.data.get("buildings", [])
        has_highrise = False
        for b in buildings:
            h_str = str(b.get("height", ""))
            # 提取数字部分 (兼容 "36m", "36.5", "18F" 等格式)
            h_match = re.match(r"(\d+\.?\d*)", h_str)
            h_val = float(h_match.group(1)) if h_match else 0
            if h_val > 24:
                has_highrise = True
                break

        if not has_n63 or not has_highrise:
            cleared = 0
            for p in self.doc.paragraphs:
                txt = p.text.strip()
                if not txt:
                    continue
                if not has_n63 and ("N63.5" in txt or "动力触探" in txt):
                    set_para_text(p, "")
                    cleared += 1
                elif not has_highrise and "波速测试" in txt:
                    set_para_text(p, "")
                    cleared += 1
            if cleared:
                logger.info(f"    勘察方法: 条件删除 {cleared} 段 (N63.5={has_n63}, 高层={has_highrise})")

        # --- 质量评述: 勘察等级注入 (第三章(五)) ---
        overview = self.config.get_project_overview()
        survey_grade = (overview or {}).get("survey_grade", "")
        if survey_grade:
            for p in self.doc.paragraphs:
                txt = p.text.strip()
                if "勘察等级为" in txt:
                    for grade in ("甲级", "乙级", "丙级"):
                        if grade in txt:
                            replace_in_para(p, grade, survey_grade)
                            break

    # ---- 水位表 + 段落 ----

    def _fill_water_level(self) -> None:
        logger.info("  水位表 + 段落...")
        bh_info = self.data["borehole_info"]
        if "wt_depth_min" not in bh_info:
            return

        idx = self.ti["water_level"]
        t = self.doc.tables[idx] if len(self.doc.tables) > idx else None

        if t:
            wtd = [
                bh["wt_depth"]
                for bh in self.data.get("boreholes", [])
                if bh.get("wt_depth")
            ]
            wte = [
                bh["wt_elv"]
                for bh in self.data.get("boreholes", [])
                if bh.get("wt_elv")
            ]
            if wtd:
                set_cell(t, 1, 0, str(len(wtd)))
                set_cell(t, 1, 1, f"{min(wtd):.2f}")
                set_cell(t, 1, 2, f"{max(wtd):.2f}")
                set_cell(t, 1, 3, f"{sum(wtd) / len(wtd):.2f}")
                set_cell(t, 1, 4, f"{min(wte):.2f}")
                set_cell(t, 1, 5, f"{max(wte):.2f}")
                set_cell(t, 1, 6, f"{sum(wte) / len(wte):.2f}")

        for p in self.doc.paragraphs:
            if "勘察期间测得钻孔内水位埋深" in p.text and "wt_depth_min" in bh_info:
                set_para_text(p, (
                    f"勘察期间测得钻孔内水位埋深"
                    f"{bh_info['wt_depth_min']:.2f}~{bh_info['wt_depth_max']:.2f}m，"
                    f"水位标高"
                    f"{bh_info['wt_elv_min']:.2f}~{bh_info['wt_elv_max']:.2f}m，"
                    f"详见表5-13。"
                ))
                break

    # ---- 地层描述段落 (第五章(二) 岩土结构及工程特性) ----

    def _build_layer_display_name(
        self, lid: str, name: str, geo_code: str,
    ) -> str:
        """构建地层段落的显示名称

        格式: {lid}{name}（{geo_code}）
        例: 1杂填土（Q4ml）, 5-1强风化花岗片麻岩（NhηγRw）
        """
        if geo_code:
            return f"{lid}{name}（{geo_code}）"
        return f"{lid}层{name}"

    def _format_layer_stats_text(
        self,
        lid: str,
        layer_stats: Dict[str, Any],
        hn_data: Dict[str, Any],
        bh_layers: Dict[str, List[Dict[str, Any]]],
    ) -> str:
        """构建地层段落的统计文本 (厚度/标高/埋深 或 未穿透)"""
        unpentrated = hn_data.get("unpentrated", set())
        max_exposed = hn_data.get("max_exposed", {})
        ldata = layer_stats.get(lid, {})

        if lid in unpentrated:
            # 未穿透: 使用最大揭露厚度
            max_d = max_exposed.get(lid, 0)
            # 计算分布
            bh_count = sum(
                1 for layers in bh_layers.values()
                if any(l["layer_id"] == lid for l in layers)
            )
            total_bh = len(bh_layers) or 1
            ratio = bh_count / total_bh
            dist = "普遍" if ratio > 0.5 else "较普遍" if ratio > 0.2 else "局部"
            return (
                f"该层在场区{dist}分布，未穿透，"
                f"最大揭露厚度{max_d:.2f}m。"
            )

        # 正常穿透层: 厚度/标高/埋深统计
        if not ldata:
            return ""

        # 计算分布
        bh_count = sum(
            1 for layers in bh_layers.values()
            if any(l["layer_id"] == lid for l in layers)
        )
        total_bh = len(bh_layers) or 1
        ratio = bh_count / total_bh
        dist = "普遍" if ratio > 0.5 else "较普遍" if ratio > 0.2 else "局部"

        return (
            f"场区{dist}分布，"
            f"厚度:{fmt_val(ldata.get('thick_min'))}~"
            f"{fmt_val(ldata.get('thick_max'))}m,"
            f"平均{fmt_val(ldata.get('thick_avg'))}m;"
            f"层底标高:{fmt_val(ldata.get('elv_min'))}~"
            f"{fmt_val(ldata.get('elv_max'))}m,"
            f"平均{fmt_val(ldata.get('elv_avg'))}m;"
            f"层底埋深:{fmt_val(ldata.get('depth_min'))}~"
            f"{fmt_val(ldata.get('depth_max'))}m,"
            f"平均{fmt_val(ldata.get('depth_avg'))}m。"
        )

    def _format_test_info_text(
        self,
        lid: str,
        table_seq: int,
        hn_data: Dict[str, Any],
    ) -> str:
        """构建试验信息段落文字"""
        spt_count = hn_data.get("spt_counts", {}).get(lid, 0)
        sample_types = hn_data.get("sample_counts", {}).get(lid, {})

        type_0 = sample_types.get(0, 0)  # 扰动样
        type_1 = sample_types.get(1, 0)  # 原状样

        parts: List[str] = []
        if type_1 > 0:
            parts.append(f"取原状样{type_1}件")
        if type_0 > 0:
            parts.append(f"取扰动样{type_0}件")
        if spt_count > 0:
            parts.append(f"进行标准贯入试验{spt_count}次")

        if not parts:
            return ""

        table_ref = f"，有关工程特性指标见表5-{table_seq}" if table_seq > 0 else ""
        return "该层" + "，".join(parts) + table_ref + "。"

    def _make_bold_paragraph(self, text: str, template_para: Any) -> Any:
        """创建一个加粗段落 (复制模板段落的格式)"""
        new_p = OxmlElement("w:p")

        # 复制段落属性 (pPr) 并设置加粗
        if template_para._element.find(qn("w:pPr")) is not None:
            pPr = copy.deepcopy(template_para._element.find(qn("w:pPr")))
            new_p.append(pPr)

        # 创建 run
        new_r = OxmlElement("w:r")

        # 复制 run 属性 (rPr) 并强制加粗
        if template_para.runs:
            src_rPr = template_para.runs[0]._element.find(qn("w:rPr"))
            if src_rPr is not None:
                rPr = copy.deepcopy(src_rPr)
                # 确保有 w:b 加粗标记
                if rPr.find(qn("w:b")) is None:
                    b_elem = OxmlElement("w:b")
                    rPr.insert(0, b_elem)
                new_r.append(rPr)
            else:
                rPr = OxmlElement("w:rPr")
                b_elem = OxmlElement("w:b")
                rPr.append(b_elem)
                new_r.append(rPr)
        else:
            rPr = OxmlElement("w:rPr")
            b_elem = OxmlElement("w:b")
            rPr.append(b_elem)
            new_r.append(rPr)

        # 设置文字
        new_t = OxmlElement("w:t")
        new_t.text = text
        new_t.set(qn("xml:space"), "preserve")
        new_r.append(new_t)
        new_p.append(new_r)

        return new_p

    def _make_normal_paragraph(self, text: str, template_para: Any) -> Any:
        """创建一个普通段落 (复制模板段落的格式, 不加粗)"""
        new_p = OxmlElement("w:p")

        # 复制段落属性 (pPr)
        if template_para._element.find(qn("w:pPr")) is not None:
            pPr = copy.deepcopy(template_para._element.find(qn("w:pPr")))
            # 移除 pPr 中的加粗 (pPr/rPr/w:b)
            pPr_rPr = pPr.find(qn("w:rPr"))
            if pPr_rPr is not None:
                b_elem = pPr_rPr.find(qn("w:b"))
                if b_elem is not None:
                    pPr_rPr.remove(b_elem)
            new_p.append(pPr)

        # 创建 run
        new_r = OxmlElement("w:r")

        # 复制 run 属性 (rPr), 但确保不加粗
        if template_para.runs:
            src_rPr = template_para.runs[0]._element.find(qn("w:rPr"))
            if src_rPr is not None:
                rPr = copy.deepcopy(src_rPr)
                # 确保没有 w:b 加粗标记
                b_elem = rPr.find(qn("w:b"))
                if b_elem is not None:
                    rPr.remove(b_elem)
                new_r.append(rPr)

        # 设置文字
        new_t = OxmlElement("w:t")
        new_t.text = text
        new_t.set(qn("xml:space"), "preserve")
        new_r.append(new_t)
        new_p.append(new_r)

        return new_p

    def _fill_layer_descriptions(self) -> None:
        """填充第五章(二) 岩土结构及工程特性

        优先使用华宁数据库数据 (ZHMS/DCSH/BG/TY/DCSJ) 动态生成;
        若无数据库则回退到模板匹配模式。
        """
        hn_data = self.data.get("hn_data", {})
        layer_stats = self.data.get("layers", {})

        if not hn_data or not hn_data.get("available"):
            logger.info("  地层描述段落... (模板模式)")
            self._fill_layer_descriptions_fallback(layer_stats)
            return

        logger.info("  地层描述段落... (华宁数据库模式)")

        descriptions = hn_data.get("descriptions", {})
        age_groups = hn_data.get("age_groups", [])
        bh_layers = hn_data.get("borehole_layers", {})
        unpentrated = hn_data.get("unpentrated", set())

        # 从 DCSH 层序列构建层号→地质代号映射
        geo_codes: Dict[str, str] = {}
        for rec in hn_data.get("layer_sequence", []):
            lid = rec["layer_id"]
            age = rec.get("age", "")
            origin = rec.get("origin", "")
            # 基岩层 origin 为纯数字时不计入代号
            if origin and not origin.isdigit():
                geo_codes[lid] = f"{age}{origin}"
            elif age:
                geo_codes[lid] = age

        # 获取地质年代名称配置覆盖
        geo_age_overrides = self.config.get_geo_age_names()

        # 定位 (二) 和 (三) 节 (仅匹配 Heading 2, 跳过目录条目)
        section_start = None
        section_end = None
        intro_para = None
        closing_para = None
        intro_idx = None
        closing_idx = None

        for i, p in enumerate(self.doc.paragraphs):
            txt = p.text.strip()
            if not txt:
                continue
            style = p.style.name
            if style != "Heading 2":
                continue
            if re.match(r"[（(]二[）)]", txt):
                section_start = i
            elif section_start is not None and re.match(r"[（(]三[）)]", txt):
                section_end = i
                break

        if section_start is None:
            logger.warning("    未找到(二)节标记, 回退到模板模式")
            self._fill_layer_descriptions_fallback(layer_stats)
            return

        # 查找 intro 和 closing
        if section_start + 1 < len(self.doc.paragraphs):
            intro_para = self.doc.paragraphs[section_start + 1]
            intro_idx = section_start + 1

        if section_end is not None:
            closing_idx = section_end - 1
            if closing_idx > (intro_idx or section_start):
                closing_para = self.doc.paragraphs[closing_idx]

        # 确定插入锚点和需要清除的范围
        body = self.doc.element.body
        anchor = intro_para._element if intro_para else \
            self.doc.paragraphs[section_start]._element

        # 清除 intro 与 (三) 之间的所有内容 (段落+表格)
        if intro_para and section_end is not None:
            start_elem = intro_para._element
            end_elem = self.doc.paragraphs[section_end]._element
            elems_to_remove = []
            clearing = False
            for child in body:
                if child is start_elem:
                    clearing = True
                    continue
                if child is end_elem:
                    break
                if clearing:
                    elems_to_remove.append(child)
            for elem in elems_to_remove:
                body.remove(elem)

        # 保留 closing 段落 (如果有)
        # 否则在末尾添加
        if closing_para is not None and closing_para._element.getparent() is not None:
            # closing 段落还在, 移到 section_end 之前
            pass
        else:
            closing_para = None

        # 生成新段落内容
        hn_reader = HuaNingDBReader("", "")  # 仅用于静态方法
        table_seq = 0
        inserted_elements: List[Any] = []

        for group_key, layer_ids_in_group in age_groups:
            # 1) 地质年代标题 (加粗)
            # 获取层名映射 (合并 HN 描述和现有 layer_names)
            merged_names = dict(self.layer_names)
            for lid, info in descriptions.items():
                if lid not in merged_names:
                    merged_names[lid] = info["name"]

            age_display = hn_reader.build_age_display_name(
                group_key, merged_names, geo_age_overrides,
            )
            bold_p = self._make_bold_paragraph(age_display, intro_para)
            inserted_elements.append(bold_p)

            # 2) 每层段落
            for lid in layer_ids_in_group:
                desc_info = descriptions.get(lid, {})
                name = desc_info.get("name", self.layer_names.get(lid, f"第{lid}层"))
                desc_text = desc_info.get("description", "")
                geo_code = geo_codes.get(lid, "")

                # 使用配置文件的描述覆盖 (如果有)
                config_desc = self.config.get_layer_description(lid)
                if config_desc:
                    # 配置描述是完整模板, 可含 {thick_*} 占位符
                    ldata = layer_stats.get(lid, {})
                    try:
                        desc_text = config_desc.format(
                            thick_min=fmt_val(ldata.get("thick_min")),
                            thick_max=fmt_val(ldata.get("thick_max")),
                            thick_avg=fmt_val(ldata.get("thick_avg")),
                            depth_min=fmt_val(ldata.get("depth_min")),
                            depth_max=fmt_val(ldata.get("depth_max")),
                            depth_avg=fmt_val(ldata.get("depth_avg")),
                            elv_min=fmt_val(ldata.get("elv_min")),
                            elv_max=fmt_val(ldata.get("elv_max")),
                            elv_avg=fmt_val(ldata.get("elv_avg")),
                        )
                    except (KeyError, IndexError):
                        pass

                # 构建完整段落
                display_name = self._build_layer_display_name(lid, name, geo_code)
                stats_text = self._format_layer_stats_text(
                    lid, layer_stats, hn_data, bh_layers,
                )

                para_text = f"{display_name}：{desc_text}"
                if desc_text and not desc_text.endswith(("。", "，")):
                    para_text += "。"
                if stats_text:
                    if not para_text.endswith("。"):
                        para_text += "。"
                    para_text += stats_text

                layer_p = self._make_normal_paragraph(para_text, intro_para)
                inserted_elements.append(layer_p)

                # 3) 试验信息段落
                table_seq += 1
                test_text = self._format_test_info_text(lid, table_seq, hn_data)
                if test_text:
                    test_p = self._make_normal_paragraph(test_text, intro_para)
                    inserted_elements.append(test_p)

                # 4) 表标题段落
                table_title = f"第{lid}层{name}工程特性指标统计表         表5-{table_seq}"
                title_p = self._make_normal_paragraph(table_title, intro_para)
                inserted_elements.append(title_p)

        # 5) 收尾段落
        closing_text = "以上各层的埋藏与分布见工程地质剖面图。"
        closing_p = self._make_normal_paragraph(closing_text, intro_para)
        inserted_elements.append(closing_p)

        # 在 anchor 之后批量插入
        for elem in reversed(inserted_elements):
            anchor.addnext(elem)

        # 如果有原有的 closing 段落 (被保留), 移到 inserted 之后
        # (已在上方处理)

        logger.info(
            f"    华宁模式: {len(age_groups)} 个年代组, "
            f"{sum(len(ids) for _, ids in age_groups)} 层, "
            f"{table_seq} 张表"
        )

    def _fill_layer_descriptions_fallback(self, layer_stats: Dict) -> None:
        """回退模式: 基于模板段落匹配填充地层描述 (原始逻辑)"""
        default_desc_templates: Dict[str, str] = {
            "1": (
                "黄褐色、灰褐色，松散、局部中密，强度不均，主要成分为风化岩渣土、"
                "碎石及建筑垃圾，块石含量约20~30%，径多10~30cm。回填时间约8年，"
                "尚未完成自重固结。场区普遍分布，"
                "厚度:{thick_min}~{thick_max}m,平均{thick_avg}m;"
                "层底标高:{elv_min}~{elv_max}m,平均{elv_avg}m;"
                "层底埋深:{depth_min}~{depth_max}m,平均{depth_avg}m。"
            ),
            "2": (
                "灰黄色、灰色，稍密~中密，饱和，主要成分为长石、石英，"
                "颗粒级配一般，含少量云母碎片。场区普遍分布，"
                "厚度:{thick_min}~{thick_max}m,平均{thick_avg}m;"
                "层底标高:{elv_min}~{elv_max}m,平均{elv_avg}m;"
                "层底埋深:{depth_min}~{depth_max}m,平均{depth_avg}m。"
            ),
        }

        generic_desc = (
            "场区{dist}分布，"
            "厚度:{thick_min}~{thick_max}m,平均{thick_avg}m;"
            "层底标高:{elv_min}~{elv_max}m,平均{elv_avg}m;"
            "层底埋深:{depth_min}~{depth_max}m,平均{depth_avg}m。"
        )

        filled: set = set()
        for p in self.doc.paragraphs:
            txt = p.text.strip()
            for lid in self.layer_ids:
                if lid in filled:
                    continue
                name = self.layer_names.get(lid, "")
                ldata = layer_stats.get(lid, {})
                if not ldata:
                    continue
                if not (
                    txt.startswith(f"{lid}层")
                    or txt.startswith(f"{lid}、")
                    or (txt and txt[0].isdigit() and name in txt[:15])
                ):
                    continue

                desc_tpl = (
                    self.config.get_layer_description(lid)
                    or default_desc_templates.get(lid)
                )

                fmt_data = {
                    "thick_min": fmt_val(ldata.get("thick_min")),
                    "thick_max": fmt_val(ldata.get("thick_max")),
                    "thick_avg": fmt_val(ldata.get("thick_avg")),
                    "depth_min": fmt_val(ldata.get("depth_min")),
                    "depth_max": fmt_val(ldata.get("depth_max")),
                    "depth_avg": fmt_val(ldata.get("depth_avg")),
                    "elv_min": fmt_val(ldata.get("elv_min")),
                    "elv_max": fmt_val(ldata.get("elv_max")),
                    "elv_avg": fmt_val(ldata.get("elv_avg")),
                }

                try:
                    if desc_tpl:
                        text = desc_tpl.format(**fmt_data)
                    else:
                        n_num = ldata.get("n", 0) or 0
                        dist = (
                            "普遍" if n_num > 50
                            else "较普遍" if n_num > 20
                            else "局部"
                        )
                        text = generic_desc.format(dist=dist, **fmt_data)
                except (KeyError, IndexError):
                    text = f"{lid}层{name}"

                set_para_text(p, text)
                filled.add(lid)
                break

        logger.info(f"    回退模式已更新: {sorted(filled)}")

    # ---- 物理力学 & 原位测试表 ----

    def _fill_phys_spt_tables(self) -> None:
        logger.info("  物理力学 & 原位测试表...")
        phys = self.data.get("phys", {})
        spt = self.data.get("spt", {})
        cpt = self.data.get("cpt", {})

        start = self.ti["phys_spt_start"]
        end = self.ti["phys_spt_end"]

        for ti in range(start, end + 1):
            if ti >= len(self.doc.tables):
                break
            t = self.doc.tables[ti]

            # 通过表头匹配地层
            hdr = "".join(
                t.rows[0].cells[c].text for c in range(min(6, len(t.rows[0].cells)))
            )
            matched = None
            for lid in self.layer_ids:
                if self.layer_names.get(lid, "") in hdr:
                    matched = lid
                    break
            if not matched:
                continue

            pd = phys.get(matched, {}).get("stats", {})
            sd = spt.get(matched, {})
            cd = cpt.get(matched, {})

            # 填充物理力学指标
            if pd:
                for ri, indicator in PHYS_ROW_MAP.items():
                    if ri >= len(t.rows):
                        continue
                    for sk, ci in STAT_COL_MAP.items():
                        v = pd.get(sk, {}).get(indicator)
                        if v is not None:
                            set_cell(t, ri, ci, fmt_val_int(v) if sk == "n" else fmt_val(v))

            # 填充标贯/动探数据
            nr = len(t.rows)
            ts = sd if sd else cd
            if ts:
                for sk, ci in STAT_COL_MAP.items():
                    raw = ts.get(sk, {}).get("raw")
                    if raw is not None and nr > 2:
                        set_cell(t, nr - 2, ci, fmt_val_int(raw))
                    corr = ts.get(sk, {}).get("corr")
                    if corr is not None and nr > 1:
                        set_cell(t, nr - 1, ci, fmt_val_int(corr))

            logger.debug(f"    T{ti}: {matched}层 {self.layer_names.get(matched, '')}")

    # ---- 承载力建议值表 ----

    def _fill_bearing_capacity(self) -> None:
        logger.info("  承载力建议值表...")
        idx = self.ti["bearing_capacity"]
        t = self.doc.tables[idx] if len(self.doc.tables) > idx else None
        if not t:
            return

        for li, lid in enumerate(self.layer_ids):
            if li + 2 >= len(t.rows):
                break
            name = self.layer_names.get(lid, "")

            # 优先从配置读取，否则留空让工程师手动填写
            bearing = self.config.get_bearing_values(lid)
            if bearing:
                fak = bearing.get("fak", "")
                es12 = bearing.get("es", "")
                es = bearing.get("e0", "")  # 注意: 列含义需匹配模板
                e0 = ""
            else:
                fak, es12, es, e0 = "", "", "", ""
                logger.debug(f"    [!] {lid}层{name} 无承载力配置")

            set_cell(t, li + 2, 0, lid)
            set_cell(t, li + 2, 1, name)
            set_cell(t, li + 2, 2, fak)
            set_cell(t, li + 2, 3, es12)
            set_cell(t, li + 2, 4, es)
            set_cell(t, li + 2, 5, e0)

    # ---- 水样 / 盐样表 ----

    def _fill_water_salt_tables(self) -> None:
        logger.info("  水样 / 盐样表...")
        water_samples = self.data.get("water_samples", [])
        salt_samples = self.data.get("salt_samples", [])

        # T19: 水样分析表
        idx_w = self.ti["water_sample"]
        t19 = self.doc.tables[idx_w] if len(self.doc.tables) > idx_w else None
        if t19 and water_samples:
            water_cols = [
                (2, "SO4"), (3, "Mg"), (4, "NH4"), (5, "OH"),
                (6, "TDS"), (7, "pH"), (8, "CO2_agg"), (9, "HCO3"), (10, "Cl"),
            ]
            for wi, w in enumerate(water_samples[:3]):
                if wi + 1 >= len(t19.rows):
                    break
                set_cell(t19, wi + 1, 0, str(w.get("field_id", "")))
                set_cell(t19, wi + 1, 1, "水")
                for ci, key in water_cols:
                    if ci < len(t19.rows[wi + 1].cells):
                        set_cell(t19, wi + 1, ci, fmt_val(w.get(key)))

        # T21: 易溶盐分析表
        idx_s = self.ti["salt_sample"]
        t21 = self.doc.tables[idx_s] if len(self.doc.tables) > idx_s else None
        if t21 and salt_samples:
            salt_cols = [(3, "SO4"), (4, "Mg"), (5, "TDS"), (6, "Cl"), (7, "pH")]
            for si, s in enumerate(salt_samples[:2]):
                if si + 1 >= len(t21.rows):
                    break
                set_cell(t21, si + 1, 0, str(s.get("location", "")))
                set_cell(t21, si + 1, 1, "土")
                for ci, key in salt_cols:
                    if ci < len(t21.rows[si + 1].cells):
                        set_cell(t21, si + 1, ci, fmt_val(s.get(key)))

    # ---- 腐蚀性评价 ----

    def _fill_corrosion_eval(self) -> None:
        logger.info("  腐蚀性评价...")
        water_samples = self.data.get("water_samples", [])
        salt_samples = self.data.get("salt_samples", [])
        corr = evaluate_corrosion(water_samples, salt_samples)
        if not corr:
            return

        corrosion_levels = ["微", "弱", "中", "强"]

        # T20: 水腐蚀性评价
        idx_wc = self.ti["water_corrosion"]
        t20 = self.doc.tables[idx_wc] if len(self.doc.tables) > idx_wc else None
        if t20 and corr.get("water"):
            wc = corr["water"]
            for wi in range(len(water_samples[:3])):
                if wi + 1 >= len(t20.rows):
                    break
                set_cell(t20, wi + 1, 2, wc["II_conc"].get("SO4", "微"))
                set_cell(t20, wi + 1, 3, wc["II_conc"].get("SO4", "微"))
                set_cell(t20, wi + 1, 4, wc["II_conc"].get("Mg", "微"))
                set_cell(t20, wi + 1, 5, wc["II_conc"].get("NH4", "微"))
                set_cell(t20, wi + 1, 7, wc["perm"])
                if len(t20.rows[wi + 1].cells) > 8:
                    set_cell(t20, wi + 1, 8, wc["steel_wet"])
                if len(t20.rows[wi + 1].cells) > 9:
                    set_cell(t20, wi + 1, 9, wc["steel_dry"])

        # T22: 土腐蚀性评价
        idx_sc = self.ti["salt_corrosion"]
        t22 = self.doc.tables[idx_sc] if len(self.doc.tables) > idx_sc else None
        if t22 and corr.get("soil"):
            sc = corr["soil"]
            for si in range(len(salt_samples[:2])):
                if si + 1 >= len(t22.rows):
                    break
                set_cell(t22, si + 1, 2, sc["II_conc"].get("SO4", "微"))
                set_cell(t22, si + 1, 3, sc["II_conc"].get("Mg", "微"))
                set_cell(t22, si + 1, 5, sc["perm"])
                if len(t22.rows[si + 1].cells) > 6:
                    set_cell(t22, si + 1, 6, sc["steel"])

        # 腐蚀性评价段落
        if corr.get("water") and corr.get("soil"):
            wc = corr["water"]
            sc = corr["soil"]
            for p in self.doc.paragraphs:
                txt = p.text.strip()
                if "腐蚀性综合评价" in txt:
                    worst_level = max(
                        wc["II_conc"].values(),
                        key=lambda x: corrosion_levels.index(x),
                    )
                    set_para_text(p, (
                        f"腐蚀性综合评价：地下水对混凝土结构具{worst_level}腐蚀性；"
                        f"地下水在长期浸水条件下对钢筋混凝土结构中的钢筋具"
                        f"{wc['steel_wet']}腐蚀性(Cl⁻)，"
                        f"在干湿交替条件下对钢筋混凝土结构中的钢筋具"
                        f"{wc['steel_dry']}腐蚀性(Cl⁻)。"
                    ))
                elif "场地土对混凝土结构具" in txt and "对钢筋混凝土" in txt:
                    set_para_text(p, (
                        f"腐蚀性评价：场地土对混凝土结构具"
                        f"{sc['II_conc'].get('SO4', '微')}腐蚀性，"
                        f"对钢筋混凝土结构中的钢筋具{sc['steel']}腐蚀性(Cl⁻)。"
                    ))

    # ---- 液化判别段落 ----

    def _replace_liquefaction_content(self, new_text: str) -> None:
        """清空 "2、液化判别" 与 "3、" 之间的模板段落，插入单段新文本"""
        liq_heading_idx = None
        next_heading_idx = None

        for i, p in enumerate(self.doc.paragraphs):
            txt = p.text.strip()
            if not txt:
                continue
            if re.match(r"2、液化判别", txt):
                liq_heading_idx = i
            elif liq_heading_idx is not None and next_heading_idx is None:
                if re.match(r"3、", txt):
                    next_heading_idx = i
                    break

        if liq_heading_idx is None:
            logger.warning("    未找到'2、液化判别'标题, 跳过")
            return

        # 移除两个标题之间的所有段落
        body = self.doc.element.body
        to_remove = []
        for i in range(liq_heading_idx + 1, next_heading_idx if next_heading_idx else len(self.doc.paragraphs)):
            to_remove.append(self.doc.paragraphs[i]._element)
        for elem in to_remove:
            body.remove(elem)

        # 插入新段落
        anchor = self.doc.paragraphs[liq_heading_idx]
        new_para = self._make_normal_paragraph(new_text, anchor)
        anchor._element.addnext(new_para)

    def _fill_liquefaction(self) -> None:
        """填充第五章(四)2 液化判别

        优先使用华宁数据库 + 液化判别表 XLS 动态生成;
        若无砂土层则输出简短说明;
        若无数据则回退到原始模板匹配。
        """
        hn_data = self.data.get("hn_data", {})
        liq_table = hn_data.get("liquefaction_table", {}) if hn_data else {}

        # ---- 检查是否有可液化砂土层 ----
        liq_layers_cfg = self.config.raw.get("liquefaction_layers", {})
        has_sand = bool(liq_layers_cfg) or (
            liq_table and liq_table.get("available") and liq_table.get("target_layers")
        )

        if not has_sand:
            logger.info("  液化判别 (无砂土层)")
            no_sand_text = "拟建场区无饱和砂土和饱和粉土层，因此不考虑场区地基土的液化影响。"
            self._replace_liquefaction_content(no_sand_text)
            return

        if not liq_table or not liq_table.get("available"):
            # 回退: 使用 Excel 导出数据的旧逻辑
            liq_data, liq_liq, liq_non = self.data.get("liquefaction", ([], 0, 0))
            logger.info(f"  液化判别 (模板模式, {len(liq_data)} 点, 液化 {liq_liq})...")
            for p in self.doc.paragraphs:
                txt = p.text.strip()
                if "综合确定" in txt and "液化" in txt:
                    if liq_liq > 0:
                        set_para_text(p, "综合确定场地饱和砂土层存在液化，场地液化等级为轻微。")
                    else:
                        set_para_text(p, "综合确定场地饱和砂土层均不液化。")
                    break
                if "进行液化判别" in txt and ("个点" in txt or "个，" in txt):
                    if liq_data:
                        text = (
                            f"对饱和砂土层进行液化判别共{len(liq_data)}个点，"
                            f"其中液化{liq_liq}个点，不液化{liq_non}个点，"
                        )
                        text += "液化等级为轻微；" if liq_liq > 0 else "均不液化。"
                        set_para_text(p, text)
                    break
            return

        # ---- HN 数据库模式: 动态生成 ----
        logger.info(
            f"  液化判别 (HN模式, {liq_table['total_points']} 点, "
            f"液化 {liq_table['liq_points']})..."
        )

        # 基础参数
        dw_values = liq_table["dw_values"]
        # 用平均水位做初判 (逐孔 SPT 判别用各孔自己的水位)
        dw = sum(dw_values) / len(dw_values) if dw_values else 1.0
        intensity = liq_table["intensity"]
        n0 = liq_table["n0"]
        d0 = HuaNingDBReader.INTENSITY_TO_D0.get(intensity, 7)
        design_group = liq_table["design_group"]

        # db 从 config 读取, 默认 1.5
        overview = self.config.get_project_overview()
        db = float(overview.get("db_foundation_depth", 1.5))

        # du 计算
        bh_layers = hn_data.get("borehole_layers", {})
        descriptions = hn_data.get("descriptions", {})
        du = HuaNingDBReader(self.config.get_hn_db_dir(),
                             self.config.get_hn_project_code()
                             ).compute_overburden_thickness(bh_layers, descriptions, dw)

        # 目标层信息
        target_layer_ids = liq_table["target_layers"]
        target_names = liq_table["target_names"]
        target_display = "、".join(
            f"第{lid}层{name}"
            for lid, name in zip(target_layer_ids, target_names)
        )
        first_target_id = target_layer_ids[0] if target_layer_ids else ""
        first_target_name = target_names[0] if target_names else ""

        # 地质年代 (从 DCSH)
        layer_sequence = hn_data.get("layer_sequence", [])
        target_age = ""
        for rec in layer_sequence:
            if rec["layer_id"] == first_target_id:
                target_age = rec.get("age", "")
                break

        # 年代显示名
        age_display = HuaNingDBReader.AGE_PREFIXES.get(target_age, target_age)
        if target_age == "Q4":
            age_condition_text = "第四纪全新世"
        elif target_age == "Q3":
            age_condition_text = "第四纪晚更新世（Q3）"
        elif target_age in ("Q2", "Q1"):
            age_condition_text = f"{age_display}（{target_age}）及其以前"
        else:
            age_condition_text = age_display

        # SPT 统计
        total = liq_table["total_points"]
        liq_count = liq_table["liq_points"]
        non_liq_count = liq_table["non_liq_points"]
        ncr_min = liq_table.get("ncr_ratio_min")
        ncr_max = liq_table.get("ncr_ratio_max")
        ile_max = liq_table.get("ile_max", 0.0)
        liq_grade = liq_table.get("liq_grade", "")
        grade_dist = liq_table.get("grade_dist", {})

        # ---- 定位液化判别段落范围 ----
        liq_heading_idx = None
        next_heading_idx = None

        for i, p in enumerate(self.doc.paragraphs):
            txt = p.text.strip()
            if not txt:
                continue
            # "2、液化判别"
            if re.match(r"2、液化判别", txt):
                liq_heading_idx = i
            # "3、软土震陷" or "3、" 紧跟其后
            elif liq_heading_idx is not None and next_heading_idx is None:
                if re.match(r"3、", txt):
                    next_heading_idx = i
                    break

        if liq_heading_idx is None:
            logger.warning("    未找到'2、液化判别'标题, 跳过")
            return

        if next_heading_idx is None:
            # 如果没找到 "3、" 就取到 (五) 或下一个 Heading 2
            for i in range(liq_heading_idx + 1, len(self.doc.paragraphs)):
                style = self.doc.paragraphs[i].style.name
                if style == "Heading 2" or style.startswith("样式 标题"):
                    next_heading_idx = i
                    break
            if next_heading_idx is None:
                next_heading_idx = len(self.doc.paragraphs)

        # 获取模板段落 (用于格式克隆)
        template_para = self.doc.paragraphs[liq_heading_idx]

        # ---- 清空原有段落 (全部移除, 公式图片后续手动补回) ----
        to_remove = []
        body = self.doc.element.body
        for i in range(liq_heading_idx + 1, next_heading_idx):
            p = self.doc.paragraphs[i]
            to_remove.append(p._element)

        # 批量移除 (避免循环中删除导致索引偏移)
        for elem in to_remove:
            body.remove(elem)

        # ---- 生成新段落 ----
        dw_str = f"{dw:.2f}"
        du_str = f"{du:.1f}" if du != int(du) else f"{du:.0f}"
        db_str = f"{db:.1f}" if db != int(db) else f"{db:.0f}"

        # 条件 (3) 公式值
        cond3_a = f"du>{d0}+{db_str}-2"
        cond3_b = f"dw>{d0}+{db_str}-3"
        cond3_c = f"du+dw>1.5×{d0}+2×{db_str}-4.5"

        # 计算条件 (3) 的实际值
        du_val = du
        dw_val = dw
        rhs_a = d0 + db - 2
        rhs_b = d0 + db - 3
        rhs_c = 1.5 * d0 + 2 * db - 4.5

        cond3_a_met = du_val > rhs_a
        cond3_b_met = dw_val > rhs_b
        cond3_c_met = (du_val + dw_val) > rhs_c
        cond3_any = cond3_a_met or cond3_b_met or cond3_c_met

        # 初判结论
        cond1_met = target_age in ("Q3", "Q2", "Q1") and intensity <= 8
        # 条件2 需要黏粒含量, 砂土不满足
        cond2_met = False  # 砂土不满足条件2

        # 构建段落文本列表
        new_paras = []

        # 1) 开头段
        new_paras.append(
            f"地下水位深度按埋深{dw_str}m考虑，地面下20m深度内饱和砂（粉）土层为"
            f"{target_display}。根据《建筑抗震设计标准》(GB/T50011-2010  2024年版)"
            f"4.3.3进行初步判别，满足条件之一时可初步判别为不液化或可不考虑液化影响，"
            f"初步判别条件如下："
        )

        # 2) 三个初判条件
        new_paras.append(
            f"(1)地质年代为第四纪晚更新世（Q3）及其以前时，7、8度时可判为不液化土；"
        )
        new_paras.append(
            "(2)粉土的黏粒（粒径小于0.005mm的颗粒）含量百分率，"
            "7度、8度、9度分别不小于10、13和16时，可判为不液化土；"
        )
        new_paras.append(
            "(3)浅埋天然地基的建筑，当上覆非液化土层厚度和地下水位深度"
            "符合下列条件之一时，可不考虑液化影响："
        )

        # 3) 条件 (3) 公式
        new_paras.append(f"            du>d0+db-2")
        new_paras.append(f"            dw>d0+db-3")
        new_paras.append(f"            du+dw>1.5d0+2db-4.5")

        # 4) 初判结果
        cond1_text = (
            f"{first_target_name}为{age_condition_text}"
            if cond1_met else
            f"第{first_target_id}层{first_target_name}为{age_condition_text}，不满足条件（1）"
        )
        cond2_text = (
            f"满足条件（2）"
            if cond2_met else
            f"第{first_target_id}层{first_target_name}为砂土，不满足条件（2）"
        )

        # 条件3 详细判定
        cond3_detail = (
            f"地下水位深度dW按{dw_str}m、上覆非液化土层厚度du按{du_str}m、"
            f"基础埋置深度db按{db_str}m、液化土特征深度d0按{d0:.1f}m"
        )
        if cond3_any:
            cond3_text = f"{cond3_detail}，满足条件（3），可不考虑液化影响。"
        else:
            cond3_text = f"{cond3_detail}，经判别不满足条件（3），需进行进一步判别。"

        new_paras.append(
            f"初步判定情况：{cond1_text}；{cond2_text}；{cond3_text}"
        )

        # 5) SPT 复判 (始终执行)
        new_paras.append(
            f"根据《建筑抗震设计标准》（GB/T50011—2010  2024年版）4.3.4节"
            f"使用标准贯入试验判别法对{target_display}进一步进行液化判别，"
            f"判别深度20.0m，地下水水位埋深按{dw_str}m。标准贯入试验液化判别公式"
            f"和液化指数的计算公式如下："
        )

        # 6) 公式占位 (保留空行, 用户插入公式图片)
        new_paras.append("")
        new_paras.append("")
        new_paras.append("")

        # 7) SPT 结果 + N/Ncr
        ncr_range_text = ""
        if ncr_min is not None and ncr_max is not None:
            ncr_range_text = f"，N/Ncr范围{ncr_min:.2f}~{ncr_max:.2f}"
        new_paras.append(
            f"{target_display}进行液化判别{total}个点，"
            f"其中液化点{liq_count}个，不液化点{non_liq_count}个{ncr_range_text}，"
            f"详见标准贯入试验液化判别及液化指数计算成果表。"
        )

        # 8) 综合结论 (含液化指数和等级)
        if liq_count > 0:
            # 液化等级分布描述
            grade_parts = []
            for g_name in ("轻微", "中等", "严重"):
                cnt = grade_dist.get(g_name, 0)
                if cnt > 0:
                    grade_parts.append(f"{g_name}{cnt}个孔")
            grade_text = "、".join(grade_parts) if grade_parts else ""

            if grade_text:
                new_paras.append(
                    f"综合判定场地{target_display}存在液化，"
                    f"最大液化指数ILE={ile_max:.2f}，"
                    f"液化等级为{liq_grade}（{grade_text}）。"
                )
            else:
                new_paras.append(
                    f"综合判定场地{target_display}存在液化，"
                    f"最大液化指数ILE={ile_max:.2f}，液化等级为{liq_grade}。"
                )
        else:
            new_paras.append(
                f"综合判定场地{target_display}不液化。"
            )

        # ---- 插入段落 ----
        anchor = self.doc.paragraphs[liq_heading_idx]._element
        inserted = []
        for text in new_paras:
            p = self._make_normal_paragraph(text, template_para)
            inserted.append(p)

        for elem in reversed(inserted):
            anchor.addnext(elem)

        logger.info(
            f"    HN模式: du={du_str}m, dw={dw_str}m, db={db_str}m, "
            f"d0={d0}m, 初判{cond1_met=}/{cond2_met=}/{cond3_any=}, "
            f"SPT: {total}点(液化{liq_count}), "
            f"ILE={ile_max:.2f}({liq_grade}), N/Ncr={ncr_min}~{ncr_max}"
        )

    # ---- 基础建议表 + 桩基参数 ----

    def _fill_foundation_tables(self) -> None:
        logger.info("  基础建议表 + 桩基参数...")
        buildings = self.data.get("buildings", [])

        # T23: 基础类型选择
        idx_f = self.ti["foundation_type"]
        t23 = self.doc.tables[idx_f] if len(self.doc.tables) > idx_f else None
        if t23 and buildings:
            for bi, b in enumerate(buildings):
                if bi + 1 >= len(t23.rows):
                    break
                h = safe_float(b.get("height", 0))
                set_cell(t23, bi + 1, 0, b["name"])
                set_cell(t23, bi + 1, 1, b.get("height", ""))
                set_cell(t23, bi + 1, 2, "桩基础")

                # 持力层选择 (可通过配置自定义)
                if h and h > 20:
                    pl = "10-2强风化片麻岩"
                elif h and h > 10:
                    pl = "7层中粗砂或9层粗砂"
                else:
                    pl = "7层中粗砂"
                set_cell(t23, bi + 1, 3, pl)
                set_cell(t23, bi + 1, 4, f"{bi + 1}～{bi + 2}剖面")

        # T24: 桩基参数
        idx_p = self.ti["pile_params"]
        t24 = self.doc.tables[idx_p] if len(self.doc.tables) > idx_p else None
        if not t24:
            return

        for li, lid in enumerate(self.layer_ids):
            if li + 2 >= len(t24.rows):
                break
            name = self.layer_names.get(lid, "")
            pile = self.config.get_pile_values(lid)
            if pile:
                q1 = pile.get("qsik1", "")
                q2 = pile.get("qpk1", "")
                q3 = pile.get("qsik2", "")
                q4 = pile.get("qpk2", "")
            else:
                q1 = q2 = q3 = q4 = ""
                logger.debug(f"    [!] {lid}层{name} 无桩基参数配置")

            set_cell(t24, li + 2, 0, lid)
            set_cell(t24, li + 2, 1, name)
            set_cell(t24, li + 2, 2, q1)
            set_cell(t24, li + 2, 3, q2)
            set_cell(t24, li + 2, 4, q3)
            set_cell(t24, li + 2, 5, q4)

    # ---- 第七章: 结论与建议 ----

    def _fill_conclusion(self) -> None:
        """填充第七章结论与建议 (配置驱动 + 自动占位符注入)

        配置结构 conclusion_suggestions:
            conclusion:   (一)结论 — 段落数组, 支持占位符
            suggestions:  (二)建议 — 段落数组

        自动占位符:
            {layer_names}        地层名称列表
            {elv_min}            最低孔口高程
            {elv_max}            最高孔口高程
            {elv_range}          高程范围
            {frozen_depth}       标准冻结深度
            {liquefaction_text}  液化判别结论
            {stability_grade}    场地稳定性等级
            {suitability_grade}  适宜性等级
            {corrosion_water}    地下水腐蚀性结论
            {corrosion_soil}     场地土腐蚀性结论
        """
        cs = self.config.get_conclusion_suggestions()
        if not cs:
            return

        # --- 自动生成占位符数据 ---
        bh_info = self.data.get("borehole_info", {})
        overview = self.config.get_project_overview() or {}
        site_eval = self.config.get_site_evaluation() or {}

        # 地层名称
        layer_names = "、".join(
            self.layer_names[lid]
            for lid in self.layer_ids
            if lid in self.layer_names
        )

        # 高程统计
        elv_min = bh_info.get("elv_min")
        elv_max = bh_info.get("elv_max")
        elv_range = ""
        if elv_min is not None and elv_max is not None:
            elv_range = f"{elv_min:.2f}～{elv_max:.2f}m"

        # 液化判别
        liq_data, liq_liq, liq_non = self.data.get(
            "liquefaction", ([], 0, 0)
        )
        if liq_liq > 0:
            liq_text = "存在液化土层"
        elif liq_data:
            liq_text = "不液化"
        else:
            liq_text = "未揭露液化土层"

        # 腐蚀性评价 (从数据自动提取)
        water_samples = self.data.get("water_samples", [])
        salt_samples = self.data.get("salt_samples", [])
        corr = evaluate_corrosion(water_samples, salt_samples)
        corr_water = ""
        corr_soil = ""
        if corr.get("water"):
            wc = corr["water"]
            worst = max(
                wc["II_conc"].values(),
                key=lambda x: ["微", "弱", "中", "强"].index(x),
            )
            corr_water = f"地下水对混凝土结构具{worst}腐蚀性"
        if corr.get("salt"):
            sc = corr["salt"]
            worst_s = max(
                sc["II_conc"].values(),
                key=lambda x: ["微", "弱", "中", "强"].index(x),
            )
            corr_soil = f"场地土对混凝土结构具{worst_s}腐蚀性"

        fmt_vars: Dict[str, str] = {
            "layer_names": layer_names,
            "elv_min": fmt_val(elv_min),
            "elv_max": fmt_val(elv_max),
            "elv_range": elv_range,
            "frozen_depth": overview.get("frozen_depth", ""),
            "liquefaction_text": liq_text,
            "stability_grade": site_eval.get("stability_grade", "基本稳定"),
            "suitability_grade": site_eval.get("suitability_grade", "较适宜"),
            "corrosion_water": corr_water,
            "corrosion_soil": corr_soil,
        }

        # --- 处理 (一)结论 子节 ---
        conclusion_paras = cs.get("conclusion", [])
        if conclusion_paras:
            self._fill_conclusion_subsection(
                "(一)结论", conclusion_paras, fmt_vars,
            )

        # --- 处理 (二)建议 子节 ---
        suggestions_paras = cs.get("suggestions", [])
        if suggestions_paras:
            self._fill_conclusion_subsection(
                "(二)建议", suggestions_paras, fmt_vars,
            )

    def _fill_conclusion_subsection(
        self,
        heading: str,
        config_paras: List[str],
        fmt_vars: Dict[str, str],
    ) -> None:
        """填充结论/建议子节的段落"""
        # 定位子节标题
        start_idx: Optional[int] = None
        for i, p in enumerate(self.doc.paragraphs):
            if heading in p.text:
                start_idx = i + 1
                break
        if start_idx is None:
            return

        # 找到子节结束位置 (下一个同级标题或文档末尾)
        end_idx = len(self.doc.paragraphs)
        for j in range(start_idx, len(self.doc.paragraphs)):
            txt = self.doc.paragraphs[j].text.strip()
            if txt.startswith("(") and ")" in txt[:6] and heading not in txt:
                end_idx = j
                break
            if txt.startswith("七") or txt.startswith("八"):
                end_idx = j
                break

        # 格式化并替换
        replaced = 0
        for k, tpl in enumerate(config_paras):
            try:
                text = tpl.format(**fmt_vars)
            except KeyError:
                text = tpl
            idx = start_idx + k
            if idx < end_idx:
                set_para_text(self.doc.paragraphs[idx], text)
                replaced += 1

        # 多余模板段落清空
        for j in range(start_idx + len(config_paras), end_idx):
            if self.doc.paragraphs[j].text.strip():
                set_para_text(self.doc.paragraphs[j], "")

        if replaced:
            logger.info(f"  结论{heading}: 替换 {replaced} 段")

    # ---- 第五章: 场地条件 (地形地貌/地下水/地震/不良地质) ----

    def _fill_site_conditions(self) -> None:
        """填充第五章各子节的文字段落

        配置结构 site_conditions:
            terrain_text:       地貌描述 (触发: "地貌单元" 或 "地貌")
            topography_text:    地形地势 (触发: "地形地势" 或 "钻孔孔口高程")
            environment_text:   周边环境 (触发: "场区周边环境" 或 "建筑红线")
            surface_water_text: 地表水 (触发: "地表水" 或 "地表水体")
            seismic_params_text: 地震参数 (触发: "地震设计基本参数" 或 "设计基本地震动")
            site_class_text:    场地类别 (触发: "等效剪切波速" 或 "场地类别判定")
            seismic_stability_text: 地震稳定性 (触发: "地震稳定性")
            seismic_zone_text:  抗震地段 (触发: "抗震地段" 或 "抗震一般地段")
            soft_soil_text:     软土震陷 (触发: "软土震陷" 或 "震陷判别")
            adverse_text:       不良地质作用 (触发: "不良地质作用" 且含 "未发现" 或 "崩塌")
            buried_text:        不利埋藏物 (触发: "不利埋藏物" 或 "埋藏的河道")
        """
        sc = self.config.get_site_conditions()
        if not sc:
            return

        # 高程统计 (自动注入, 优先简报 → Excel → 华宁 DK)
        bh_info = self.data.get("borehole_info", {})
        elv_min = fmt_val(bh_info.get("elv_min"))
        elv_max = fmt_val(bh_info.get("elv_max"))

        # 简报高程 (优先)
        briefing = self.data.get("briefing", {})
        bf_elv = briefing.get("elevation", {}) if briefing.get("available") else {}
        if not elv_min and bf_elv:
            elv_min = fmt_val(bf_elv.get("min"))
            elv_max = fmt_val(bf_elv.get("max"))

        # 华宁 DK 文件高程回退
        hn_data = self.data.get("hn_data", {})
        hn_elv = hn_data.get("elevation", {}) if hn_data else {}
        if not elv_min and hn_elv:
            elv_min = fmt_val(hn_elv.get("elv_min"))
            elv_max = fmt_val(hn_elv.get("elv_max"))

        # 地貌类型自动判定 (优先简报高差, 回退 DK; 地层名称优先简报)
        auto_terrain = ""
        elv_range = 0.0
        if bf_elv and bf_elv.get("range"):
            elv_range = bf_elv["range"]
        elif hn_elv:
            elv_range = hn_elv.get("elv_range", 0.0)

        # 简报中已填写地貌类型则直接使用
        if bf_elv.get("terrain_type"):
            auto_terrain = bf_elv["terrain_type"]
        elif elv_range > 0 or (hn_data and hn_data.get("available")):
            origins = set()
            layer_names = []
            # 简报中的地层名称
            if briefing.get("available") and briefing.get("layers"):
                for lid, info in briefing["layers"].items():
                    layer_names.append(info.get("name", ""))
            # DCSH 中的成因 (优先, 精确)
            if hn_data and hn_data.get("available"):
                for rec in hn_data.get("layer_sequence", []):
                    o = rec.get("origin", "")
                    if o and not o.isdigit():
                        origins.add(o)
                if not layer_names:
                    for lid, info in hn_data.get("descriptions", {}).items():
                        if isinstance(info, dict):
                            layer_names.append(info.get("name", ""))
            # 无 DCSH 成因时, 从简报地层名推断
            if not origins and briefing.get("available") and briefing.get("layers"):
                origins = HuaNingDBReader.infer_origins_from_names(
                    briefing["layers"]
                )
            auto_terrain = HuaNingDBReader.determine_terrain(
                elv_range, origins, layer_names
            )
            if auto_terrain:
                logger.info(
                    f"  地貌自动判定: {auto_terrain} "
                    f"(高差{elv_range:.1f}m, 成因{origins})"
                )

        # 液化自动判断
        liq_data, liq_liq, liq_non = self.data.get("liquefaction", ([], 0, 0))
        if liq_liq > 0:
            auto_liq_short = "存在液化"
        elif liq_data:
            auto_liq_short = "不液化"
        else:
            auto_liq_short = "未揭露液化土层"

        # 占位符变量
        fmt_vars: Dict[str, str] = {
            "elv_min": elv_min,
            "elv_max": elv_max,
            "liquefaction_result": auto_liq_short,
        }

        # 场地位置 (从 project_overview 注入)
        overview = self.config.get_project_overview()
        fmt_vars["site_location"] = overview.get("site_location", "")

        # 地震参数 (自动查表 + config覆盖)
        seismic_params = self.config.get_seismic_params()
        if seismic_params:
            # 注入: {pga} {period} {intensity} {group}
            fmt_vars["pga"] = seismic_params.get("pga", "")
            fmt_vars["period"] = seismic_params.get("period", "")
            fmt_vars["intensity"] = seismic_params.get("intensity", "")
            fmt_vars["group"] = seismic_params.get("group", "")

        # 触发词 → 配置键映射
        triggers: List[Tuple[List[str], str]] = [
            (["地貌单元", "所处地貌"], "terrain_text"),
            (["地形地势", "钻孔孔口高程"], "topography_text"),
            (["建筑红线", "场区周边环境", "周边环境"], "environment_text"),
            (["地表水体", "地表水"], "surface_water_text"),
            (["地震设计基本参数", "设计基本地震动"], "seismic_params_text"),
            (["等效剪切波速", "场地类别判定"], "site_class_text"),
            (["地震稳定性"], "seismic_stability_text"),
            (["抗震地段"], "seismic_zone_text"),
            (["不良地质作用"], "adverse_text"),
            (["不利埋藏物", "埋藏的河道"], "buried_text"),
        ]

        replaced_count = 0
        for p in self.doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue

            # 跳过子节标题行 (如 "1、地震设计基本参数", "2、液化判别")
            if re.match(r"^\d+、", txt) and len(txt) < 20:
                continue

            for keywords, config_key in triggers:
                config_text = sc.get(config_key, "")

                # 地貌: config 为空时使用自动判定结果
                if config_key == "terrain_text" and not config_text and auto_terrain:
                    config_text = f"场区地貌类型属{auto_terrain}。"

                if not config_text:
                    continue
                if any(kw in txt for kw in keywords):
                    try:
                        filled = config_text.format(**fmt_vars)
                    except KeyError:
                        filled = config_text
                    set_para_text(p, filled)
                    replaced_count += 1
                    break

        # ---- 软土震陷判别: 位置定位 ----
        # 找到 "3、软土震陷判别" 标题，替换其后的内容段落
        _DEFAULT_SOFT_SOIL = (
            "根据《岩土工程勘察规范》（2009年版）（GB50021-2001）第5.7.11条"
            "及表5.5，场区不存在软弱层，因此不需要考虑震陷影响。"
        )
        soft_soil_cfg = sc.get("soft_soil_text", "")
        soft_soil_text = soft_soil_cfg if soft_soil_cfg else _DEFAULT_SOFT_SOIL
        try:
            soft_soil_text = soft_soil_text.format(**fmt_vars)
        except KeyError:
            pass

        for i, p in enumerate(self.doc.paragraphs):
            txt = p.text.strip()
            if re.match(r"3、软土震陷", txt) or re.match(r"3、.*震陷判别", txt):
                # 替换标题之后的第一个非空段落
                for j in range(i + 1, min(i + 5, len(self.doc.paragraphs))):
                    p2 = self.doc.paragraphs[j]
                    t2 = p2.text.strip()
                    if not t2:
                        continue
                    # 遇到下一个子节标题 "4、" 则停止
                    if re.match(r"4、", t2):
                        break
                    set_para_text(p2, soft_soil_text)
                    replaced_count += 1
                    break
                break

        if replaced_count:
            logger.info(f"  场地条件: 替换 {replaced_count} 段")
        else:
            logger.debug("  场地条件: 未找到匹配段落")

    # ---- 等效剪切波速估算 & 场地类别判定 (GB 50011-2010 §4.1) ----

    def _select_wave_velocity_boreholes(
        self, boreholes: List[Dict[str, Any]], count: int
    ) -> List[Dict[str, Any]]:
        """选取波速估算钻孔，按空间分布均匀选取

        Args:
            boreholes: 所有钻孔列表 (含 x, y, elevation, depth, type)
            count: 选取数量 (2 或 3)

        Returns:
            选取的钻孔列表
        """
        if len(boreholes) <= count:
            return list(boreholes)

        # 优先选有坐标的钻孔
        with_xy = [bh for bh in boreholes if bh.get("x") and bh.get("y")]
        if len(with_xy) >= count:
            # 按坐标空间分布选取: 取边界极点 + 中间点
            pool = sorted(with_xy, key=lambda bh: bh["x"])  # 按X排序
            if count == 2:
                # 一东一西 (X坐标最远端)
                selected = [pool[0], pool[-1]]
            else:  # count == 3
                # 东西两端 + Y方向中间
                mid_idx = len(pool) // 2
                # 在中间1/3范围内按Y取中间值
                third = len(pool) // 3
                mid_pool = sorted(pool[third: 2 * third], key=lambda bh: bh.get("y", 0) or 0)
                selected = [pool[0], mid_pool[len(mid_pool)//2] if mid_pool else pool[mid_idx], pool[-1]]
            # 去重
            seen = set()
            result = []
            for bh in selected:
                key = bh.get("id", "")
                if key not in seen:
                    seen.add(key)
                    result.append(bh)
            if len(result) >= count:
                return result[:count]

        # 回退: 按孔深降序，取最深的 count 个（深孔穿透更多地层，Vs估算更准确）
        deep = sorted(boreholes, key=lambda bh: bh.get("depth") or 0, reverse=True)
        # 尽量分散: 取最深的 + 中等 + 最浅的
        n = len(deep)
        indices = [0] if count == 1 else [0, n - 1] if count == 2 else [0, n // 2, n - 1]
        result = []
        seen = set()
        for i in indices:
            if i < n:
                key = deep[i].get("id", "")
                if key not in seen:
                    seen.add(key)
                    result.append(deep[i])
        return result[:count]

    def _get_borehole_layers_for_vs(
        self, bh_id: str,
    ) -> Optional[List[Dict[str, Any]]]:
        """获取指定钻孔的分层数据，用于波速估算

        优先从华宁 DCSJ 读取，回退到 Excel 地层统计
        """
        hn_data = self.data.get("hn_data", {})
        if hn_data and hn_data.get("available"):
            bh_layers = hn_data.get("borehole_layers", {})
            descriptions = hn_data.get("descriptions", {})
            if bh_id in bh_layers:
                layers = []
                prev_depth = 0.0
                for entry in bh_layers[bh_id]:
                    lid = entry["layer_id"]
                    depth = entry.get("depth")
                    name = descriptions.get(lid, {}).get("name", f"第{lid}层")
                    thickness = (depth - prev_depth) if depth is not None else 0
                    if thickness > 0:
                        layers.append({
                            "name": name,
                            "thickness": round(thickness, 1),
                            "depth_bottom": round(depth, 1) if depth else None,
                        })
                    prev_depth = depth if depth else prev_depth
                return layers if layers else None

        # 回退: 用 Excel 地层统计数据（按平均厚度，孔间差异忽略）
        layer_stats = self.data.get("layers", {})
        if layer_stats:
            return [
                {"name": info.get("name", ""), "thickness": info.get("thick_avg", 0) or 0}
                for lid, info in sorted(layer_stats.items(),
                                        key=lambda x: layer_sort_key(x[0]))
            ]
        return None

    def _fill_wave_velocity_table(self) -> None:
        """填充等效剪切波速计算及场地类别判定表

        优先方案二 (实测波速数据):
            自动检测项目 base_dir/波速/ 目录下的波速报告.docx 或波速数据.xlsx
        回退方案一 (估算):
            按土层名称估算 Vs → 计算 νse → 判定场地类别
        """
        if evaluate_site_class_from_layers is None:
            logger.debug("  波速: 模块未加载，跳过")
            return

        # ---- 1. 查找波速表 (按表头关键词匹配) ----
        wave_table = None
        for ti, t in enumerate(self.doc.tables):
            header_text = ""
            for ri in range(min(3, len(t.rows))):
                cells_text = [cell.text.strip() for cell in t.rows[ri].cells]
                header_text += " ".join(cells_text)
            if any(kw in header_text for kw in ["等效剪切波速", "场地类别判定"]):
                wave_table = t
                logger.info(f"  波速表: 表格 #{ti} ({len(t.rows)-1}行数据区)")
                break

        if wave_table is None:
            logger.debug("  波速表: 模板中未找到等效剪切波速表，跳过")
            return

        # ---- 2. 尝试加载实测波速数据 (方案二) ----
        wave_data = self._load_wave_velocity_from_project()
        if wave_data:
            self._fill_wave_table_from_real_data(wave_table, wave_data)
            return

        # ---- 回退方案一: 估算 ----
        self._fill_wave_table_by_estimation(wave_table)

    def _load_wave_velocity_from_project(self) -> Optional[List[Dict[str, Any]]]:
        """自动检测项目目录中的波速数据"""
        if load_wave_velocity_data is None:
            return None

        base_dir = self.config.base_dir
        if not base_dir:
            return None

        # 自动检测 波速/ 子目录
        wave_dir = os.path.join(base_dir, "波速")
        if not os.path.isdir(wave_dir):
            # 尝试 已有资料/波速/
            wave_dir = os.path.join(base_dir, "已有资料", "波速")
        if not os.path.isdir(wave_dir):
            # 尝试直接找文件
            for root, dirs, files in os.walk(base_dir):
                for d in dirs:
                    if "波速" in d:
                        wave_dir = os.path.join(root, d)
                        break
                if os.path.isdir(wave_dir):
                    break

        if not os.path.isdir(wave_dir):
            logger.debug("  波速: 未找到波速数据目录")
            return None

        docx_path = None; xlsx_path = None
        for f in os.listdir(wave_dir):
            fp = os.path.join(wave_dir, f)
            if f.endswith(".docx") and not f.startswith("~$") and "波速" in f:
                docx_path = fp
            elif f.endswith(".xlsx") and not f.startswith("~$") and "波速" in f:
                xlsx_path = fp

        if not docx_path and not xlsx_path:
            logger.debug(f"  波速: 目录存在但未找到波速文件: {wave_dir}")
            return None

        logger.info(f"  波速: 加载实测数据 — {wave_dir}")
        data = load_wave_velocity_data(xlsx_path=xlsx_path, docx_path=docx_path)
        logger.info(f"  波速: 共 {len(data)} 个钻孔实测数据")
        return data if data else None

    def _fill_wave_table_from_real_data(
        self, tbl, wave_data: List[Dict[str, Any]]
    ) -> None:
        """用实测波速数据填充表格 (方案二)"""
        # 判断表格格式: 8列(逐层) vs 6列(汇总)
        col_count = len(tbl.columns)
        is_detail_table = col_count >= 7

        # 找数据起始行
        data_start_row = 1
        for ri in range(min(3, len(tbl.rows))):
            row_text = " ".join(c.text.strip() for c in tbl.rows[ri].cells)
            if "钻孔编号" in row_text or "孔号" in row_text or "序号" in row_text:
                data_start_row = ri + 1
                break

        # 如果实测数据有 layers，填充逐层详情
        first = wave_data[0]
        has_layers = "layers" in first and first.get("layers")

        current_row = data_start_row
        for data in wave_data:
            bh_id = data.get("bh_id", "")
            vse = data.get("vse")
            cover = data.get("cover_thickness")
            sc = data.get("site_class", "")

            if has_layers:
                # 逐层填充 (8列格式)
                for li, layer in enumerate(data["layers"]):
                    if current_row >= len(tbl.rows):
                        break
                    if li == 0:
                        set_cell(tbl, current_row, 0, bh_id)
                        set_cell(tbl, current_row, 6,
                                 f"<{cover}m" if cover and cover > 50 else str(cover or ""))
                        set_cell(tbl, current_row, 7,
                                 f"{vse}m/s" if vse else "")
                    else:
                        set_cell(tbl, current_row, 0, "")
                        set_cell(tbl, current_row, 6, "")
                        set_cell(tbl, current_row, 7, "")

                    set_cell(tbl, current_row, 1, str(li + 1))
                    set_cell(tbl, current_row, 2, layer.get("name", ""))
                    set_cell(tbl, current_row, 3, str(layer.get("vs", "")))
                    set_cell(tbl, current_row, 4, str(layer.get("depth", "")))
                    set_cell(tbl, current_row, 5, str(layer.get("thickness", "")))
                    current_row += 1
            else:
                # 汇总格式 (6列): 序号 | 孔号 | 覆盖层 | d0 | νse | 场地类别
                if current_row >= len(tbl.rows):
                    break
                set_cell(tbl, current_row, 0, str(current_row - data_start_row + 1))
                set_cell(tbl, current_row, 1, f"ZK{bh_id}" if not bh_id.startswith("ZK") else bh_id)
                if col_count >= 3:
                    set_cell(tbl, current_row, 2, str(cover) if cover else "")
                if col_count >= 4:
                    d0_val = data.get("d0", min(cover or 20, 20))
                    set_cell(tbl, current_row, 3, str(d0_val) if d0_val else "")
                if col_count >= 5:
                    set_cell(tbl, current_row, 4, str(vse) if vse else "")
                if col_count >= 6:
                    set_cell(tbl, current_row, 5, sc)
                current_row += 1

        # 综合判定
        all_vse = [d.get("vse") for d in wave_data if d.get("vse")]
        all_sc = [d.get("site_class") for d in wave_data if d.get("site_class")]
        if all_vse and all_sc:
            avg_vse = round(sum(all_vse) / len(all_vse), 1)
            unique_sc = list(set(all_sc))
            sc_display = unique_sc[0] if len(unique_sc) == 1 else "~".join(sorted(
                unique_sc, key=lambda x: {"I₀": 0, "I₁": 1, "Ⅱ": 2, "Ⅲ": 3, "Ⅳ": 4}.get(x, 0)
            ))
            logger.info(
                f"  波速综合: νse={min(all_vse)}~{max(all_vse)}m/s, "
                f"场地类别={sc_display}"
            )

            # 更新段落: "本场地在0～XXm深度范围内等效剪切波速为XX～XXm/s"
            for p in self.doc.paragraphs:
                txt = p.text.strip()
                if "等效剪切波速为" in txt and ("深度范围内" in txt or "场地覆盖层" in txt):
                    cover_min = min(d.get("cover_thickness", 0) or 0 for d in wave_data)
                    cover_max = max(d.get("cover_thickness", 0) or 0 for d in wave_data)
                    new_text = (
                        f"据测试结果，本场地在0～{cover_max:.1f}m深度范围内"
                        f"等效剪切波速为{min(all_vse):.1f}～{max(all_vse):.1f}m/s，"
                        f"场地覆盖层厚度{cover_min:.1f}～{cover_max:.1f}m，"
                        f"依据《建筑与市政工程抗震通用规范》(GB55002)中规定，"
                        f"该建筑场地类别属于{sc_display}类。"
                    )
                    set_para_text(p, new_text)
                    break

        # 更新波速测试段落
        bh_ids = [d.get("bh_id", "") for d in wave_data]
        for p in self.doc.paragraphs:
            txt = p.text.strip()
            if "剪切波测试" in txt and ("钻孔内进行了" in txt or "号钻孔" in txt):
                new_text = (
                    f"在本场地{'、'.join(bh_ids[:6])}号等{len(bh_ids)}个钻孔内"
                    f"进行了岩土体剪切波测试，实测成果见表5-11。"
                )
                set_para_text(p, new_text)
                break

    def _fill_wave_table_by_estimation(self, tbl) -> None:
        """估算模式填充波速表 (方案一)"""
        boreholes = self.data.get("boreholes", [])
        buildings = self.data.get("buildings", [])
        total_bh = len(boreholes)

        is_large = len(buildings) > 3 or total_bh > 50
        bh_count = 3 if is_large else 2
        selected_bhs = self._select_wave_velocity_boreholes(boreholes, bh_count)

        if not selected_bhs:
            logger.warning("  波速表: 无可用钻孔")
            return

        logger.info(
            f"  波速估算: {'大' if is_large else '小'}场地, "
            f"选 {len(selected_bhs)} 孔: "
            f"{[bh['id'] for bh in selected_bhs]}"
        )

        # 逐孔估算
        results: List[Dict[str, Any]] = []
        for bh in selected_bhs:
            bh_id = bh.get("id", "")
            layers = self._get_borehole_layers_for_vs(bh_id)
            if not layers:
                continue
            r = evaluate_site_class_from_layers(layers)
            r["bh_id"] = bh_id
            results.append(r)
            logger.info(
                f"    钻孔 {bh_id}: νse={r['vse']}m/s, "
                f"d={r['cover_thickness']}m, {r['site_class']}类"
            )

        if not results:
            logger.warning("  波速表: 所有钻孔均无有效数据")
            return

        # 填充表格 (逐层详情)
        data_start_row = 1
        for ri in range(min(3, len(tbl.rows))):
            row_text = " ".join(c.text.strip() for c in tbl.rows[ri].cells)
            if "钻孔编号" in row_text or "土层编号" in row_text:
                data_start_row = ri + 1
                break

        current_row = data_start_row
        for r in results:
            for li, layer in enumerate(r["layers_detail"]):
                if current_row >= len(tbl.rows):
                    break
                if li == 0:
                    set_cell(tbl, current_row, 0, r["bh_id"])
                    set_cell(tbl, current_row, 7, f"{r['vse']}m/s")
                    set_cell(tbl, current_row, 6,
                             f"<{r['cover_thickness']}m"
                             if r['cover_thickness'] > 50
                             else str(r['cover_thickness']))
                else:
                    set_cell(tbl, current_row, 0, "")
                    set_cell(tbl, current_row, 6, "")
                    set_cell(tbl, current_row, 7, "")
                set_cell(tbl, current_row, 1, str(li + 1))
                set_cell(tbl, current_row, 2, layer["name"])
                set_cell(tbl, current_row, 3, str(layer["vs"]))
                set_cell(tbl, current_row, 4, str(layer["depth_bottom"]))
                set_cell(tbl, current_row, 5, str(layer["thickness"]))
                current_row += 1

        # 更新波速段落
        bh_ids_str = "、".join(r["bh_id"] for r in results)
        for p in self.doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue
            if "剪切波测试" in txt and ("钻孔内进行了" in txt or "号钻孔" in txt):
                set_para_text(p, (
                    f"在本场地{bh_ids_str}号等{len(results)}个钻孔内"
                    f"进行了岩土体剪切波测试，实测成果见表5-11。"
                ))
                break

    # ---- 场地稳定性及适宜性评价 (CJJ57-2012) ----

    def _fill_site_evaluation(self) -> None:
        """根据配置和实际数据生成场地稳定性及适宜性评价段落"""
        eval_config = self.config.get_site_evaluation()
        if not eval_config:
            return

        paragraphs = eval_config.get("paragraphs", [])
        if not paragraphs:
            return

        # 从数据自动生成液化文本
        liq_data, liq_liq, liq_non = self.data.get("liquefaction", ([], 0, 0))
        if liq_liq > 0:
            auto_liq = f"场地饱和砂土层存在液化（液化等级轻微）"
        elif liq_data:
            auto_liq = "场地饱和砂土层均不液化"
        else:
            auto_liq = "场地未揭露液化土层"

        # 从腐蚀性评价自动获取文本
        water_samples = self.data.get("water_samples", [])
        salt_samples = self.data.get("salt_samples", [])
        corr = evaluate_corrosion(water_samples, salt_samples)
        if corr.get("water"):
            wc = corr["water"]
            worst_level = max(
                wc["II_conc"].values(),
                key=lambda x: ["微", "弱", "中", "强"].index(x),
            )
            auto_corr = f"地下水对混凝土结构具{worst_level}腐蚀性"
        else:
            auto_corr = ""

        # 准备占位符变量: 优先使用配置中的值，其次自动生成
        placeholders = {
            "liquefaction_text": eval_config.get("liquefaction_text", auto_liq),
            "corrosion_text": eval_config.get("corrosion_text", auto_corr),
            "adverse_geology": eval_config.get("adverse_geology", "不良地质作用不发育"),
            "buried_objects": eval_config.get(
                "buried_objects",
                "未发现埋藏的河道、沟浜、墓穴、防空洞、孤石等对工程不利的埋藏物",
            ),
            "seismic_section": eval_config.get("seismic_section", "对建筑抗震一般地段"),
            "stability_grade": eval_config.get("stability_grade", "基本稳定"),
            "suitability_grade": eval_config.get("suitability_grade", "较适宜"),
            "suitability_text": eval_config.get("suitability_text", ""),
        }

        # 格式化各段落
        formatted: List[str] = []
        for tpl in paragraphs:
            try:
                text = tpl.format(**placeholders)
            except KeyError:
                text = tpl  # 未定义的占位符保留原文
            formatted.append(text)

        # 在模板中定位并替换: 找到包含触发关键词的首段，替换后续段
        trigger_keywords = [
            "场地稳定性", "稳定性评价", "适宜性评价",
            "场地稳定性及适宜性", "稳定性和适宜性",
        ]

        replaced = False
        for i, p in enumerate(self.doc.paragraphs):
            txt = p.text.strip()
            if any(kw in txt for kw in trigger_keywords):
                # 首段替换为格式化后的第一段
                set_para_text(p, formatted[0])
                # 后续配置段落: 尝试替换模板中的后续段落
                for j, ftext in enumerate(formatted[1:], start=1):
                    target_idx = i + j
                    if target_idx < len(self.doc.paragraphs):
                        next_txt = self.doc.paragraphs[target_idx].text.strip()
                        # 如果下一段是相关段落（非新章节），替换它
                        if next_txt and not any(
                            kw in next_txt
                            for kw in ("结论", "建议", "技术标准", "勘察依据", "地基基础")
                        ):
                            set_para_text(self.doc.paragraphs[target_idx], ftext)
                        else:
                            # 模板段落不够，剩余文本追加到上一段末尾
                            prev_p = self.doc.paragraphs[target_idx - 1]
                            old_text = prev_p.text.rstrip()
                            set_para_text(prev_p, old_text + "\n" + ftext)
                    else:
                        # 模板段落不足，追加到最后一段
                        last_p = self.doc.paragraphs[len(self.doc.paragraphs) - 1]
                        old_text = last_p.text.rstrip()
                        set_para_text(last_p, old_text + "\n" + ftext)
                replaced = True
                break

        if replaced:
            logger.info(
                f"  场地评价: {placeholders['stability_grade']} / "
                f"{placeholders['suitability_grade']}"
            )
        else:
            logger.debug("  场地评价: 未找到匹配段落，跳过")

    # ---- 地基评价 (4.5.7 §2 均匀性 / §5 软弱下卧层 / §6 变形参数) ----

    def _fill_foundation_evaluation(self) -> None:
        """根据配置和数据生成地基评价段落"""
        eval_config = self.config.get_foundation_evaluation()
        if not eval_config:
            return

        paragraphs = eval_config.get("paragraphs", [])
        if not paragraphs:
            return

        # --- 自动数据收集 ---
        layer_stats = self.data.get("layers", {})
        bearing_config = {
            lid: self.config.get_bearing_values(lid)
            for lid in self.layer_ids
        }

        # §2 地基均匀性: 根据各层厚度变异系数判断
        auto_uniformity = self._auto_uniformity_text(layer_stats)

        # §5 软弱下卧层: 找出低承载力层
        auto_weak_layer = self._auto_weak_layer_text(bearing_config)

        # §6 变形计算参数: 引用 T18 表 Es 值
        bearing_idx = self.ti.get("bearing_capacity", 18)
        table_ref = f"表{bearing_idx}" if bearing_idx else "承载力建议值表"
        auto_deformation = (
            f"各土层压缩模量（Es）详见{table_ref}，"
            f"建议按《建筑地基基础设计规范》(GB 50007-2011)进行地基变形计算。"
        )

        # 准备占位符: 优先使用配置值，否则使用自动生成值
        placeholders: Dict[str, str] = {
            "uniformity_text": eval_config.get("uniformity_text", auto_uniformity),
            "weak_layer_text": eval_config.get("weak_layer_text", auto_weak_layer),
            "deformation_text": eval_config.get("deformation_text", auto_deformation),
        }

        # 附加占位符: 持力层信息
        bearing_layer = eval_config.get("bearing_layer", "")
        bearing_fak = eval_config.get("bearing_fak", "")
        placeholders["bearing_layer"] = bearing_layer
        placeholders["bearing_fak"] = bearing_fak

        # 格式化段落
        formatted: List[str] = []
        for tpl in paragraphs:
            try:
                text = tpl.format(**placeholders)
            except KeyError:
                text = tpl
            formatted.append(text)

        # 在模板中定位触发关键词段落并替换
        trigger_keywords = [
            "地基均匀性", "天然地基评价", "软弱下卧层",
            "地基评价", "地基土评价", "地基均匀",
        ]

        replaced = False
        for i, p in enumerate(self.doc.paragraphs):
            txt = p.text.strip()
            if any(kw in txt for kw in trigger_keywords):
                # 首段替换
                set_para_text(p, formatted[0])
                # 后续段落
                for j, ftext in enumerate(formatted[1:], start=1):
                    target_idx = i + j
                    if target_idx < len(self.doc.paragraphs):
                        next_txt = self.doc.paragraphs[target_idx].text.strip()
                        # 如果下一段是相关段落（非新章节标题），替换
                        if next_txt and not any(
                            kw in next_txt for kw in (
                                "结论", "建议", "技术标准", "勘察依据",
                                "场地稳定性", "地下水", "基坑",
                            )
                        ):
                            set_para_text(
                                self.doc.paragraphs[target_idx], ftext
                            )
                        else:
                            # 模板段落不够，追加到当前段
                            prev_p = self.doc.paragraphs[target_idx - 1]
                            old = prev_p.text.rstrip()
                            set_para_text(prev_p, old + "\n" + ftext)
                    else:
                        # 模板段落不足，追加到最后一段
                        last_p = self.doc.paragraphs[len(self.doc.paragraphs) - 1]
                        old = last_p.text.rstrip()
                        set_para_text(last_p, old + "\n" + ftext)
                replaced = True
                break

        if replaced:
            logger.info("  地基评价: 已填充 %d 段", len(formatted))
        else:
            logger.debug("  地基评价: 未找到匹配段落，跳过")

    def _auto_uniformity_text(
        self, layer_stats: Dict[str, Dict[str, Any]]
    ) -> str:
        """根据各层厚度变异系数自动判断地基均匀性"""
        if not layer_stats or not self.layer_ids:
            return "地基土均匀性需根据实际勘探资料判定。"

        # 计算各层厚度变异系数 (CV of thickness)
        cvs: List[float] = []
        for lid in self.layer_ids:
            ls = layer_stats.get(lid, {})
            thk_avg = ls.get("thick_avg")
            thk_min = ls.get("thick_min")
            thk_max = ls.get("thick_max")
            if thk_avg and thk_min is not None and thk_max is not None and thk_avg > 0:
                # 近似变异系数 = (max - min) / avg
                cv = (thk_max - thk_min) / thk_avg
                cvs.append(cv)

        if not cvs:
            return "地基土均匀性需根据实际勘探资料判定。"

        avg_cv = sum(cvs) / len(cvs)
        max_cv = max(cvs)

        if max_cv < 0.3:
            return (
                "拟建场地地基土为均匀地基，各土层分布较稳定，"
                "厚度变化较小，地基土均匀性较好。"
            )
        elif max_cv < 0.5:
            return (
                "拟建场地地基土为较均匀地基，各土层分布有一定变化，"
                "厚度变化一般，地基土均匀性一般。"
            )
        else:
            return (
                "拟建场地地基土为不均匀地基，各土层分布变化较大，"
                "厚度变化较大，地基土均匀性较差。"
            )

    def _auto_weak_layer_text(
        self, bearing_config: Dict[str, Optional[Dict[str, str]]]
    ) -> str:
        """自动检测软弱下卧层（承载力 < 100 kPa 的土层）"""
        weak_layers: List[Tuple[str, str, str]] = []  # (层号, 名称, fak)

        for lid in self.layer_ids:
            bc = bearing_config.get(lid)
            if not bc:
                continue
            fak_str = bc.get("fak", "")
            fak = safe_float(fak_str)
            if fak is not None and fak < 100 and fak > 0:
                name = self.layer_names.get(lid, "")
                weak_layers.append((lid, name, str(int(fak))))

        if weak_layers:
            desc_parts = []
            for lid, name, fak in weak_layers:
                desc_parts.append(f"{lid}层{name}（fak={fak}kPa）")
            layers_desc = "、".join(desc_parts)
            return (
                f"场地内局部分布{layers_desc}，承载力较低，"
                f"作为软弱下卧层需进行验算，"
                f"建议按《建筑地基基础设计规范》(GB 50007-2011)第5.2.7条进行软弱下卧层验算。"
            )
        else:
            return (
                "场地内未发现明显的软弱下卧层，"
                "各土层承载力可满足一般建筑天然地基要求。"
            )

    # ---- 第六章: 岩土工程分析评价 (配置驱动段落填充) ----

    def _fill_analysis_evaluation(self) -> None:
        """填充第六章岩土工程分析评价各子节

        配置结构 analysis_evaluation:
            layer_eval:       (一)岩土层逐层工程评价 — key为层号(如"1","1-1")
            anti_float:       (四)3 地下水的力学作用/抗浮
            foundation_text:  (五)1 基础选型补充说明
            pile_eval:        (五)2 桩基评价综合文字
            special_soils:    (六)特殊性岩土工程分析评价
            stability:        (八)地基稳定性评价
            excavation:       (九)基坑开挖有关问题 (has_basement=False时清空)
            risk:             (十)地质条件可能造成的工程风险 (has_basement=False时清空)
            deformation:      (十一)建筑物变形分析
        """
        ae = self.config.get_analysis_evaluation()
        if not ae:
            return

        # 有地下室判断 (基坑/风险章节条件控制)
        overview = self.config.get_project_overview()
        conditions = self._evaluate_standard_conditions()
        has_basement = (overview or {}).get(
            "has_basement", conditions.get("has_basement", False)
        )

        # 章节映射: (触发关键词列表, 配置键, 条件类型)
        #   cond=None: 始终处理
        #   cond="basement": 仅 has_basement=True 时处理, 否则清空
        sections: List[Tuple[List[str], str, Any]] = [
            (
                ["岩土层(体)工程分析", "岩土层(体)工程评价"],
                "layer_eval", None,
            ),
            (
                ["地下水的力学作用"],
                "anti_float", None,
            ),
            (
                ["地基及基础方案", "基础选型"],
                "foundation_text", None,
            ),
            (
                ["桩基评价", "桩基形式"],
                "pile_eval", None,
            ),
            (
                ["特殊性岩土"],
                "special_soils", None,
            ),
            (
                ["地基稳定性评价"],
                "stability", None,
            ),
            (
                ["基坑开挖有关问题", "基坑开挖"],
                "excavation", "basement",
            ),
            (
                ["工程风险"],
                "risk", "basement",
            ),
            (
                ["建筑物变形分析"],
                "deformation", None,
            ),
        ]

        current_section: Optional[str] = None
        current_config_key: Optional[str] = None
        section_start: int = 0
        replaced_count = 0

        for i, p in enumerate(self.doc.paragraphs):
            txt = p.text.strip()

            # --- 检测章节标题 (合并专用 + 通用检测) ---
            matched_section_key: Optional[str] = None

            # 1) 专用关键词匹配 (sections 列表)
            for keywords, config_key, cond in sections:
                if any(kw in txt for kw in keywords):
                    matched_section_key = config_key
                    break

            # 2) 通用标题模式: 括号格式 (X) 或已知非括号标题
            if not matched_section_key:
                if (
                    (txt.startswith("(") and ")" in txt[:6])
                    or txt in ("基坑开挖有关问题",)
                ):
                    matched_section_key = "__boundary__"

            # 处理检测结果
            if matched_section_key:
                # 先 flush 上一个章节
                if current_section:
                    replaced_count += self._process_analysis_section(
                        current_section, current_config_key,
                        section_start, i, ae, has_basement,
                    )

                if matched_section_key == "__boundary__":
                    # 不关心的章节, 重置状态
                    current_section = None
                    current_config_key = None
                else:
                    current_section = txt
                    current_config_key = matched_section_key
                    section_start = i + 1

        # 处理最后一个章节
        if current_section:
            replaced_count += self._process_analysis_section(
                current_section, current_config_key,
                section_start, len(self.doc.paragraphs), ae, has_basement,
            )

        if replaced_count:
            logger.info(f"  分析评价: 处理 {replaced_count} 个子节")
        else:
            logger.debug("  分析评价: 未找到匹配章节")

    def _process_analysis_section(
        self,
        heading: str,
        config_key: Optional[str],
        start: int,
        end: int,
        ae: Dict[str, Any],
        has_basement: bool,
    ) -> int:
        """处理第六章的一个子节

        返回处理的段落数。
        """
        if not config_key:
            return 0

        config_data = ae.get(config_key, None)

        # === 条件清空: 无基坑时清空基坑/风险章节 ===
        # 匹配条件: 标题中包含 "基坑" 或 "工程风险"
        if not has_basement and (
            "基坑" in heading or "工程风险" in heading
        ):
            cleared = 0
            for j in range(start, end):
                if self.doc.paragraphs[j].text.strip():
                    set_para_text(self.doc.paragraphs[j], "")
                    cleared += 1
            if cleared:
                logger.info(
                    f"    {config_key}: 无基坑, 清空 {cleared} 段"
                )
            return cleared

        if not config_data:
            return 0

        # === (一) 岩土层逐层评价: 按层号匹配 ===
        if config_key == "layer_eval" and isinstance(config_data, dict):
            replaced = 0
            for j in range(start, end):
                txt_j = self.doc.paragraphs[j].text.strip()
                if not txt_j:
                    continue
                for layer_key, layer_text in config_data.items():
                    if layer_text and txt_j.startswith(f"第{layer_key}层"):
                        set_para_text(self.doc.paragraphs[j], layer_text)
                        replaced += 1
                        break
            return replaced

        # === 通用段落替换 ===
        if not isinstance(config_data, list):
            config_data = [str(config_data)]

        replaced = 0
        for k, text in enumerate(config_data):
            idx = start + k
            if idx < end:
                set_para_text(self.doc.paragraphs[idx], text)
                replaced += 1
        # 多余模板段落清空
        for j in range(start + len(config_data), end):
            if self.doc.paragraphs[j].text.strip():
                set_para_text(self.doc.paragraphs[j], "")

        return replaced

    # ---- 技术标准列表 (配置驱动 + 条件过滤) ----

    def _fill_standards(self) -> None:
        """根据配置和实际数据生成技术标准列表段落"""
        standards_config = self.config.get_technical_standards()
        if not standards_config:
            return

        # 始终包含的标准
        always = standards_config.get("always", [])

        # 按条件过滤的标准
        conditional = standards_config.get("conditional", [])
        conditions = self._evaluate_standard_conditions()

        # 合并: 始终项 + 满足条件的条件项
        all_standards = list(always)
        for s in conditional:
            cond = s.get("condition", "always")
            if cond == "always" or conditions.get(cond, False):
                all_standards.append(s)

        # 法律法规和其他依据
        laws = standards_config.get("laws", [])
        other = standards_config.get("other", [])

        # 组装格式化文本
        lines: List[str] = []
        idx = 1

        if all_standards:
            lines.append("1、国家标准：")
            for s in all_standards:
                lines.append(f"{idx})《{s['name']}》({s['code']})")
                idx += 1

        if laws:
            lines.append("2、法律法规：")
            for law in laws:
                lines.append(f"{idx})《{law}》" if "《" not in law else f"{idx}){law}")
                idx += 1

        if other:
            lines.append("3、其他：")
            for item in other:
                lines.append(f"{idx}){item}")
                idx += 1

        full_text = "\n".join(lines)

        # 在模板中定位技术标准段落并替换
        trigger_keywords = ["执行的主要技术标准", "技术标准", "勘察依据", "依据的技术标准"]
        replaced = False

        for i, p in enumerate(self.doc.paragraphs):
            txt = p.text.strip()
            if any(kw in txt for kw in trigger_keywords):
                set_para_text(p, full_text)
                replaced = True
                # 清除后续已有的标准条目段落
                for j in range(i + 1, min(i + 40, len(self.doc.paragraphs))):
                    next_txt = self.doc.paragraphs[j].text.strip()
                    if not next_txt:
                        continue
                    # 匹配标准条目格式: 数字)《...》(编号) 或 纯标准名称
                    if (
                        re.match(r"^\d+[)）、]", next_txt)
                        or "《" in next_txt and "(" in next_txt
                        or next_txt.startswith(("国家标准", "行业标准", "地方标准", "法律法规", "其他"))
                    ):
                        set_para_text(self.doc.paragraphs[j], "")
                    elif any(kw in next_txt for kw in ("工程概况", "勘察目的", "场地", "拟建")):
                        break  # 已到达下一个章节
                break

        if replaced:
            included = [s["name"] for s in all_standards]
            excluded = [
                s["name"] for s in conditional
                if s.get("condition", "always") != "always"
                and not conditions.get(s["condition"], False)
            ]
            logger.info(f"  技术标准: 纳入 {len(all_standards)} 条, 排除 {len(excluded)} 条")
            if excluded:
                logger.debug(f"    排除: {', '.join(excluded)}")
        else:
            logger.debug("  技术标准: 未找到匹配段落，跳过")

    def _evaluate_standard_conditions(self) -> Dict[str, bool]:
        """评估标准条件，返回各条件是否满足"""
        layer_ids = self.layer_ids
        bh_info = self.data.get("borehole_info", {})
        buildings = self.data.get("buildings", [])
        water_samples = self.data.get("water_samples", [])

        # 有岩层: 层号含 "10-" 等
        has_rock_layers = any("-" in str(lid) for lid in layer_ids)

        # 有岩石试验样品
        has_rock_samples = (bh_info.get("rock", 0) or 0) > 0

        # 有水样
        has_water_samples = len(water_samples) > 0

        # 配置了桩基参数
        has_pile = any(
            self.config.get_pile_values(lid) is not None
            for lid in layer_ids
        )

        # 有地下室/地下结构
        has_basement = False
        for b in buildings:
            floors_str = b.get("floors", "")
            height_str = b.get("height", "")
            # 检查是否有地下层数标注 (如 "地上18层/地下1层")
            if "地下" in floors_str:
                has_basement = True
                break
            # 检查备注或结构字段
            for field in ["size", "span", "indoor_elv", "structure"]:
                if "地下" in str(b.get(field, "")):
                    has_basement = True
                    break
            if has_basement:
                break

        return {
            "has_rock_layers": has_rock_layers,
            "has_rock_samples": has_rock_samples,
            "has_water_samples": has_water_samples,
            "has_pile_foundation": has_pile,
            "has_basement": has_basement,
        }

    # ---- 日期替换 (配置驱动) ----

    def _apply_date_replacements(self) -> None:
        """按配置执行日期替换 (不再硬编码)"""
        date_repls = self.config.get_date_replacements()
        if not date_repls:
            return

        logger.info("  日期替换...")
        for old, new in date_repls:
            for p in self.doc.paragraphs:
                replace_in_para(p, old, new)


# ============================================================
# 主程序
# ============================================================

def load_layer_names(
    config: ProjectConfig,
    loader: SurveyDataLoader,
    layer_stats: Dict[str, Any],
    args_layers: Optional[str],
) -> Tuple[Dict[str, str], List[str]]:
    """加载地层名称映射和排序后的地层 ID 列表"""

    layer_names: Dict[str, str] = {}

    # 优先使用 --layers 参数
    if args_layers and os.path.isfile(args_layers):
        with open(args_layers, "r", encoding="utf-8") as f:
            layer_names = json.load(f)
    # 其次使用配置中的地层定义
    elif config.layers:
        for layer in config.layers:
            lid = str(layer.get("id", ""))
            name = layer.get("name", "")
            if lid and name:
                layer_names[lid] = name
    # 最后从物理力学表推断
    else:
        phys = loader.load_physical_stats()
        for lid in layer_stats:
            if not (any(c.isdigit() for c in str(lid)) and len(str(lid)) < 10):
                continue
            if lid in phys and phys[lid].get("name"):
                layer_names[lid] = phys[lid]["name"]
            else:
                layer_names[lid] = f"第{lid}层"

    # 排序层号
    valid_ids = [
        k
        for k in layer_stats.keys()
        if any(c.isdigit() for c in str(k)) and len(str(k)) < 10
    ]
    layer_ids = sorted(valid_ids, key=layer_sort_key)

    return layer_names, layer_ids


def main() -> None:
    parser = argparse.ArgumentParser(
        description=f"岩土工程勘察报告自动生成工具 v{__version__}"
    )
    parser.add_argument("--config", "-c", help="JSON 配置文件路径 (推荐)")
    parser.add_argument("--project", "-p", help="项目目录路径 (向后兼容)")
    parser.add_argument("--template", "-t", help="模板文件路径")
    parser.add_argument("--output", "-o", help="输出文件路径")
    parser.add_argument("--layers", help="地层名称映射 JSON 文件")
    parser.add_argument("--dry-run", action="store_true", help="仅加载数据，不生成报告")
    parser.add_argument("--verbose", "-v", action="store_true", help="输出详细日志")
    args = parser.parse_args()

    setup_logging(args.verbose)

    if not args.config and not args.project:
        parser.error("必须指定 --config 或 --project")

    # 检查第三方依赖 (dry-run 也需要 xlrd/openpyxl)
    _check_dependencies()

    # 1. 加载配置
    logger.info("=" * 60)
    logger.info(f"  岩土工程勘察报告自动生成工具 v{__version__}")
    logger.info("=" * 60)

    config = ProjectConfig(
        config_path=args.config,
        project_dir=args.project,
        template_override=args.template,
        output_override=args.output,
        layers_override=args.layers,
    )

    # 2. 加载数据
    loader = SurveyDataLoader(config.excel_dir)
    logger.info(f"\n[数据加载] Excel 目录: {config.excel_dir}")

    boreholes = loader.load_boreholes()
    bh_info = loader.classify_boreholes(boreholes)
    logger.info(
        f"  钻孔: {bh_info['total']} 个, "
        f"进尺 {bh_info['total_depth']:.2f}m"
    )
    logger.info(
        f"  控制孔 {bh_info['ctrl']}, 一般孔 {bh_info['general']}"
    )
    logger.info(
        f"  取土 {bh_info['qutu']}, 标贯 {bh_info['biaoguan']}, "
        f"一般 {bh_info['yiban']}, 动探 {bh_info['zhongtan']}, "
        f"波速 {bh_info['bosk']}"
    )

    layer_stats = loader.load_layer_stats()
    phys = loader.load_physical_stats()
    spt = loader.load_spt_stats()
    cpt = loader.load_cpt_stats()
    liq_data, liq_liq, liq_non = loader.load_liquefaction()
    rock = loader.load_rock_stats()
    buildings = loader.load_buildings()
    water_samples = loader.load_water_samples()
    salt_samples = loader.load_salt_samples()

    logger.info(
        f"  物理力学: {len(phys)} 层, "
        f"标贯: {len(spt)} 层, "
        f"动探: {len(cpt)} 层"
    )
    logger.info(
        f"  液化: {len(liq_data)} 点, "
        f"水样: {len(water_samples)} 件, "
        f"盐样: {len(salt_samples)} 件"
    )

    # 地层名称
    layer_names, layer_ids = load_layer_names(config, loader, layer_stats, args.layers)
    logger.info(f"  地层 ID 序列: {layer_ids}")

    # 华宁数据库 (可选)
    hn_data: Dict[str, Any] = {"available": False}
    hn_db_dir = config.get_hn_db_dir()
    if hn_db_dir and os.path.isdir(hn_db_dir):
        logger.info(f"\n[华宁数据库] 目录: {hn_db_dir}")
        project_code = config.get_hn_project_code()
        reader = HuaNingDBReader(hn_db_dir, project_code)
        hn_data = reader.read()
    elif hn_db_dir:
        logger.warning(f"  华宁数据库目录不存在: {hn_db_dir}")

    # 简报 (华宁生成的 .docx, 提取高程/地下水/地层等)
    briefing_data: Dict[str, Any] = {}
    sj_dir = config.get_sj_dir()
    if sj_dir:
        logger.info(f"\n[简报] 搜索目录: {sj_dir}")
        briefing_data = read_briefing(sj_dir)

    # 建筑-钻孔映射 (DXF 总平面图空间匹配)
    building_bh_mapping: Dict[str, Any] = {"available": False}
    dxf_path = config.get_dxf_path()
    building_info_path = config.get_building_info_path()
    borehole_coords_hn = (hn_data.get("borehole_coords", {})
                          if hn_data.get("available") else {})
    # 优先从 DXF INSERT 块获取钻孔坐标 (图面精确位置)
    borehole_coords: Dict[str, Tuple[float, float]] = {}
    if dxf_path:
        borehole_coords = extract_borehole_coords_from_dxf(
            dxf_path, hn_coords=borehole_coords_hn or None)
        if borehole_coords:
            src = "INSERT→华宁匹配" if borehole_coords_hn else "INSERT/TEXT"
            logger.info(f"  钻孔坐标从 DXF 提取 ({src}, {len(borehole_coords)} 个)")
    # DXF 无结果时回退到华宁 DK 坐标
    if not borehole_coords and borehole_coords_hn:
        borehole_coords = borehole_coords_hn
        logger.info(f"  钻孔坐标使用华宁 DK ({len(borehole_coords)} 个)")
    if dxf_path and borehole_coords:
        logger.info(f"\n[建筑映射] DXF: {dxf_path}")
        if building_info_path:
            logger.info(f"  建筑信息: {building_info_path}")
        building_bh_mapping = extract_building_borehole_mapping(
            dxf_path, borehole_coords, building_info_path,
            text_layer=config.raw.get("dxf_text_layer", ""),
            outline_layer=config.raw.get("dxf_outline_layer", ""),
        )
    elif dxf_path:
        logger.info(f"\n[建筑映射] DXF 存在但无钻孔坐标数据")

    if args.dry_run:
        logger.info("\n[Dry-run] 数据加载完成，不生成报告。")
        return

    # 3. 填充报告
    logger.info(f"\n[模板] {config.template_path}")
    logger.info(f"[输出] {config.output_path}")

    data = {
        "boreholes": boreholes,
        "borehole_info": bh_info,
        "layers": layer_stats,
        "phys": phys,
        "spt": spt,
        "cpt": cpt,
        "liquefaction": (liq_data, liq_liq, liq_non),
        "rock": rock,
        "buildings": buildings,
        "water_samples": water_samples,
        "salt_samples": salt_samples,
        "hn_data": hn_data,
        "briefing": briefing_data,
        "building_bh_mapping": building_bh_mapping,
    }

    filler = ReportFiller(
        config.template_path,
        config.output_path,
        data,
        layer_names,
        layer_ids,
        config,
    )
    filler.fill_all()
    filler.save()

    logger.info(f"\n{'=' * 60}")
    logger.info(f"  生成完成: {config.output_path}")
    logger.info(f"{'=' * 60}")


if __name__ == "__main__":
    main()
